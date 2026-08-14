import datetime as dt
from datetime import date, datetime as dt_datetime
import time
from io import BytesIO
from pathlib import Path
import pandas as pd
import streamlit as st
import hashlib
import re
import textwrap
import html as html_sanitizer
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import plotly.express as px
import streamlit.components.v1 as components
from streamlit_autorefresh import st_autorefresh
import qrcode
import io
import urllib.parse
import zipfile
import google.generativeai as genai
import base64
import os
import uuid
import traceback
import logging

import requests
from icalendar import Calendar
import pytz
from datetime import datetime, date

@st.cache_data(ttl=1800) # Se actualiza cada 30 minutos para no saturar el servidor
def obtener_clima_vicuna():
    try:
        # Coordenadas exactas de Vicuña, Chile
        lat = -30.0319
        lon = -70.7081
        url = f"https://api.open-meteo.com/v1/forecast?latitude={lat}&longitude={lon}&current_weather=true"
        
        respuesta = requests.get(url, timeout=5).json()
        temp = respuesta["current_weather"]["temperature"]
        codigo_clima = respuesta["current_weather"]["weathercode"]
        
        # Elegir el ícono Phosphor correcto según el clima
        if codigo_clima in [0, 1]: # Despejado
            icon = "ph-sun"
            color = "#f59e0b" # Naranja sol
        elif codigo_clima in [2, 3]: # Nublado parcial
            icon = "ph-cloud-sun"
            color = "#94a3b8" # Gris azulado
        elif codigo_clima in [45, 48]: # Niebla
            icon = "ph-cloud-fog"
            color = "#cbd5e1"
        elif codigo_clima in [51, 53, 55, 61, 63, 65, 80, 81, 82]: # Lluvia
            icon = "ph-cloud-rain"
            color = "#3b82f6" # Azul agua
        else: # Otros (tormenta, nieve, etc)
            icon = "ph-cloud"
            color = "#64748b"

        # Devolvemos el HTML listo para insertar
        return f"<i class='ph-fill {icon}' style='color: {color}; font-size: 1.3rem; margin-right: 5px; vertical-align: middle;'></i> <span style='vertical-align: middle;'>{temp}°C</span>"
    except Exception as e:
        return "" # Si no hay internet, no mostramos nada para que no salga error

st.set_page_config(page_title="Sistema de Horarios CAV", page_icon="📅", layout="wide", initial_sidebar_state="expanded")

# ==============================================================================
# DISEÑO RESPONSIVO INTELIGENTE V18
# Se adapta automáticamente a móvil, tablet, notebook, escritorio y pantallas TV.
# ==============================================================================
def aplicar_diseno_responsivo_inteligente():
    css_responsivo = """
    <style>
    :root {
        --cav-touch-size: 46px;
        --cav-radius-responsive: clamp(10px, 1.2vw, 18px);
        --cav-gap-responsive: clamp(.55rem, 1.1vw, 1.15rem);
        --cav-page-padding-x: clamp(.65rem, 2.4vw, 2.4rem);
        --cav-page-padding-y: clamp(.75rem, 1.8vw, 1.8rem);
        --cav-content-width: 1760px;
    }

    html {
        text-size-adjust: 100%;
        -webkit-text-size-adjust: 100%;
        scroll-behavior: smooth;
    }

    body,
    [data-testid="stAppViewContainer"] {
        overflow-x: hidden;
    }

    .block-container {
        width: 100% !important;
        max-width: var(--cav-content-width) !important;
        padding:
            var(--cav-page-padding-y)
            var(--cav-page-padding-x)
            calc(var(--cav-page-padding-y) + env(safe-area-inset-bottom))
            var(--cav-page-padding-x)
            !important;
    }

    h1 {
        font-size: clamp(1.75rem, 3.3vw, 3rem) !important;
        line-height: 1.08 !important;
        text-wrap: balance;
    }

    h2 {
        font-size: clamp(1.4rem, 2.6vw, 2.25rem) !important;
        line-height: 1.12 !important;
        text-wrap: balance;
    }

    h3 {
        font-size: clamp(1.12rem, 2vw, 1.65rem) !important;
        line-height: 1.16 !important;
    }

    p,
    label,
    [data-testid="stMarkdownContainer"] {
        overflow-wrap: anywhere;
    }

    img,
    video,
    canvas,
    svg {
        max-width: 100% !important;
        height: auto;
    }

    iframe {
        max-width: 100% !important;
    }

    /* streamlit-autorefresh usa un iframe invisible.
       No se debe convertir en un bloque alto con height:auto. */
    iframe[title*="autorefresh"],
    iframe[src*="streamlit_autorefresh"] {
        position: absolute !important;
        width: 0 !important;
        min-width: 0 !important;
        height: 0 !important;
        min-height: 0 !important;
        border: 0 !important;
        opacity: 0 !important;
        pointer-events: none !important;
        overflow: hidden !important;
    }

    div[data-testid="stElementContainer"]:has(
        iframe[title*="autorefresh"]
    ),
    div[data-testid="stElementContainer"]:has(
        iframe[src*="streamlit_autorefresh"]
    ) {
        display: none !important;
        width: 0 !important;
        height: 0 !important;
        min-height: 0 !important;
        margin: 0 !important;
        padding: 0 !important;
        overflow: hidden !important;
    }

    [data-testid="stImage"] img {
        object-fit: contain;
    }

    [data-testid="stHorizontalBlock"] {
        gap: var(--cav-gap-responsive) !important;
        align-items: stretch;
    }

    [data-testid="column"] {
        min-width: 0 !important;
    }

    [data-testid="stForm"],
    [data-testid="stExpander"],
    [data-testid="stMetric"],
    [data-testid="stAlert"],
    [data-testid="stDataFrame"],
    [data-testid="stFileUploader"] {
        border-radius: var(--cav-radius-responsive) !important;
    }

    [data-testid="stForm"] {
        padding: clamp(.85rem, 1.8vw, 1.5rem) !important;
    }

    [data-testid="stMetric"] {
        height: 100%;
        padding: clamp(.7rem, 1.25vw, 1.05rem) !important;
    }

    [data-testid="stMetricValue"] {
        font-size: clamp(1.3rem, 2.2vw, 2rem) !important;
        overflow-wrap: anywhere;
    }

    [data-testid="stMetricLabel"] {
        white-space: normal !important;
    }

    .stButton > button,
    .stDownloadButton > button,
    [data-testid="stLinkButton"] a,
    button[kind] {
        min-height: var(--cav-touch-size) !important;
        border-radius: clamp(9px, 1vw, 14px) !important;
        white-space: normal !important;
        line-height: 1.2 !important;
    }

    input,
    textarea,
    [data-baseweb="select"] > div {
        min-height: var(--cav-touch-size) !important;
        font-size: max(16px, 1rem) !important;
        border-radius: clamp(8px, .9vw, 13px) !important;
    }

    textarea {
        min-height: 110px !important;
    }

    [data-testid="stTabs"] [role="tablist"] {
        overflow-x: auto;
        overflow-y: hidden;
        scrollbar-width: thin;
        flex-wrap: nowrap;
    }

    [data-testid="stTabs"] [role="tab"] {
        flex: 0 0 auto;
        min-height: var(--cav-touch-size);
        white-space: nowrap;
    }

    [data-testid="stDataFrame"],
    [data-testid="stTable"] {
        width: 100% !important;
        overflow-x: auto !important;
    }

    [data-testid="stDataFrame"] > div {
        max-width: 100% !important;
    }

    [data-testid="stSidebar"] {
        width: clamp(270px, 25vw, 370px) !important;
    }

    [data-testid="stSidebarContent"] {
        padding-bottom: env(safe-area-inset-bottom);
    }

    [data-testid="stDialog"] > div {
        width: min(94vw, 760px) !important;
        max-height: 92dvh !important;
        overflow-y: auto !important;
    }

    /* Tablets y notebooks pequeños */
    @media (max-width: 1180px) {
        :root {
            --cav-content-width: 100%;
            --cav-gap-responsive: .75rem;
        }

        [data-testid="stHorizontalBlock"] {
            flex-wrap: wrap !important;
        }

        [data-testid="column"] {
            flex: 1 1 min(320px, 100%) !important;
            width: auto !important;
        }
    }

    /* Celulares y tablets verticales */
    @media (max-width: 760px) {
        :root {
            --cav-touch-size: 48px;
            --cav-page-padding-x: .7rem;
            --cav-page-padding-y: .8rem;
        }

        .block-container {
            max-width: 100% !important;
        }

        h1,
        h2 {
            text-align: center;
        }

        [data-testid="stHorizontalBlock"] {
            display: flex !important;
            flex-direction: column !important;
            gap: .7rem !important;
        }

        [data-testid="column"] {
            flex: 1 1 100% !important;
            width: 100% !important;
            min-width: 100% !important;
        }

        .stButton,
        .stDownloadButton,
        [data-testid="stLinkButton"] {
            width: 100%;
        }

        .stButton > button,
        .stDownloadButton > button,
        [data-testid="stLinkButton"] a {
            width: 100% !important;
            justify-content: center !important;
            padding-left: .8rem !important;
            padding-right: .8rem !important;
        }

        [data-testid="stSidebar"] {
            width: min(88vw, 340px) !important;
        }

        [data-testid="stMetric"] {
            min-height: 88px;
        }

        [data-testid="stForm"] {
            padding: .85rem !important;
        }

        [data-testid="stFileUploader"] section {
            padding: .8rem !important;
        }

        [data-testid="stFileUploaderDropzoneInstructions"] {
            min-width: 0 !important;
        }

        [data-testid="stToolbar"],
        [data-testid="stDecoration"] {
            max-width: 100vw !important;
        }

        div[data-testid="stPlotlyChart"] {
            overflow-x: auto;
        }

        div[data-testid="stPlotlyChart"] > div {
            min-width: 100%;
        }
    }

    /* Celulares muy pequeños */
    @media (max-width: 390px) {
        :root {
            --cav-page-padding-x: .5rem;
        }

        h1 {
            font-size: 1.62rem !important;
        }

        h2 {
            font-size: 1.32rem !important;
        }

        [data-testid="stMetricValue"] {
            font-size: 1.22rem !important;
        }
    }

    /* Pantallas grandes y televisores */
    @media (min-width: 1700px) {
        :root {
            --cav-content-width: 1900px;
            --cav-page-padding-x: 2.8rem;
        }

        body {
            font-size: 1.05rem;
        }

        .stButton > button,
        .stDownloadButton > button,
        [data-testid="stLinkButton"] a {
            min-height: 50px !important;
        }
    }

    /* Dispositivos táctiles */
    @media (pointer: coarse) {
        button,
        a,
        input,
        textarea,
        [role="tab"],
        [role="option"] {
            touch-action: manipulation;
        }

        [data-baseweb="select"] [role="option"] {
            min-height: 46px;
        }
    }

    /* Pantallas horizontales de poca altura */
    @media (orientation: landscape) and (max-height: 650px) {
        :root {
            --cav-page-padding-y: .45rem;
        }

        h1 {
            margin-top: .2rem !important;
            margin-bottom: .45rem !important;
        }

        [data-testid="stForm"] {
            padding-top: .65rem !important;
            padding-bottom: .65rem !important;
        }
    }

    /* Accesibilidad: reducir animaciones cuando el dispositivo lo solicita */
    @media (prefers-reduced-motion: reduce) {
        *,
        *::before,
        *::after {
            scroll-behavior: auto !important;
            animation-duration: .01ms !important;
            animation-iteration-count: 1 !important;
            transition-duration: .01ms !important;
        }
    }
    </style>
    """

    if hasattr(st, "html"):
        st.html(css_responsivo)
    else:
        st.markdown(css_responsivo, unsafe_allow_html=True)


aplicar_diseno_responsivo_inteligente()

@st.cache_data(ttl=1800) # Se actualiza cada 30 minutos
def obtener_eventos_google_calendar(url_ics):
    if not url_ics:
        return []
    try:
        respuesta = requests.get(url_ics)
        respuesta.raise_for_status()
        
        # Leer el calendario con la nueva librería
        calendario = Calendar.from_ical(respuesta.text)
        zona_horaria = pytz.timezone('America/Santiago')
        ahora = dt_datetime.now(zona_horaria).date()
        
        eventos_hoy = []
        
        # Buscar todos los eventos (VEVENT)
        for componente in calendario.walk('VEVENT'):
            inicio = componente.get('dtstart').dt
            
            # Verificar si es un evento con hora específica
            if isinstance(inicio, datetime):
                inicio_dt = inicio.astimezone(zona_horaria)
                es_hoy = (inicio_dt.date() == ahora)
                hora_str = inicio_dt.strftime("%H:%M")
                hora_sort = hora_str
            # O si es un evento de "Todo el día" (solo fecha)
            else:
                es_hoy = (inicio == ahora)
                hora_str = "TODO EL DÍA"
                hora_sort = "00:00" # Para que salga primero en la lista

            if es_hoy:
                titulo = str(componente.get('summary', 'Evento Especial'))
                descripcion = str(componente.get('description', ''))
                
                eventos_hoy.append({
                    "hora_sort": hora_sort,
                    "display_hora": hora_str,
                    "titulo": titulo,
                    "descripcion": descripcion,
                    "categoria": "Evento Especial"
                })
                
        # Ordenamos los eventos por hora
        return sorted(eventos_hoy, key=lambda x: x['hora_sort'])
    except Exception as e:
        return []
        
# --- CONFIGURACIÓN DE GEMINI ---
try:
    GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
except KeyError:
    st.error(
        "🚨 Error: No se encontró la API Key en los secretos de Streamlit."
    )
    st.stop()


@st.cache_resource(show_spinner=False)
def obtener_modelo_gemini(api_key):
    """Reutiliza el modelo de IA y evita reconstruirlo en cada rerun."""
    genai.configure(api_key=api_key)
    return genai.GenerativeModel("gemini-2.5-flash")


model = obtener_modelo_gemini(GEMINI_API_KEY)

# --- TU FUNCIÓN PARA LLAMAR A LA IA ---
def consultar_gemini(prompt):
    try:
        response = model.generate_content(prompt) 
        return response.text
    except Exception as e:
        return f"Error con la IA: {e}"
        st.header("💬 Asistente Técnico IA")

# ------------------------------------------------------------------
# CONFIGURACIÓN SUPABASE (NUEVO MOTOR DE BASE DE DATOS)
# ------------------------------------------------------------------
from supabase import create_client, Client, ClientOptions

# Las credenciales se leen exclusivamente desde Streamlit Secrets.
try:
    URL_SUPABASE = st.secrets["SUPABASE_URL"]
    CLAVE_SUPABASE = st.secrets["SUPABASE_KEY"]
except KeyError as e:
    st.error(f"🚨 Falta configurar {e} en los Secrets de Streamlit.")
    st.stop()

@st.cache_resource(show_spinner=False)
def obtener_cliente_supabase(url_supabase, clave_supabase):
    """
    Reutiliza el cliente Supabase entre reruns.
    """
    opciones = ClientOptions(
        postgrest_client_timeout=45,
        storage_client_timeout=60,
    )
    return create_client(
        url_supabase,
        clave_supabase,
        options=opciones,
    )


supabase: Client = obtener_cliente_supabase(
    URL_SUPABASE,
    CLAVE_SUPABASE,
)


# Registro local de errores para diagnóstico administrativo.
logging.basicConfig(level=logging.INFO)
LOGGER = logging.getLogger("sistema_cav")

def registrar_error(contexto, error):
    LOGGER.error("%s: %s\n%s", contexto, error, traceback.format_exc())
    st.session_state.setdefault("errores_sistema", [])
    st.session_state.errores_sistema.append({
        "fecha": dt_datetime.now().isoformat(timespec="seconds"),
        "contexto": contexto,
        "error": str(error),
    })
    st.session_state.errores_sistema = st.session_state.errores_sistema[-50:]

def select_paginado(tabla, columnas="*", filtros=None, orden=None, desc=False, pagina=1000):
    """Lee todos los registros de una tabla evitando el límite por defecto de Supabase."""
    filas = []
    desde = 0
    while True:
        consulta = supabase.table(tabla).select(columnas)
        for metodo, campo, valor in (filtros or []):
            consulta = getattr(consulta, metodo)(campo, valor)
        if orden:
            consulta = consulta.order(orden, desc=desc)
        bloque = consulta.range(desde, desde + pagina - 1).execute().data or []
        filas.extend(bloque)
        if len(bloque) < pagina:
            break
        desde += pagina
    return filas

def registrar_auditoria(accion, modulo, registro_id=None, detalle=None):
    """Registra acciones importantes. Si aún no existe la tabla, no interrumpe la app."""
    try:
        supabase.table("auditoria").insert({
            "usuario": st.session_state.get("profesor_name") or st.session_state.get("role") or "sistema",
            "accion": accion,
            "modulo": modulo,
            "registro_id": str(registro_id) if registro_id is not None else None,
            "detalle": detalle or {},
            "fecha": dt_datetime.now().isoformat(),
        }).execute()
    except Exception as e:
        registrar_error("auditoria", e)


# ==============================================================================
# DIPLOMAS DIGITALES V15: REGISTRO, STORAGE Y VISTA PÚBLICA
# ==============================================================================
DIPLOMAS_BUCKET = "diplomas-digitales"
EQUIVALENCIA_HOJAS_NORMALES = 2.7
TONER_MIN_GRAMOS_POR_DIPLOMA = 0.15
TONER_MAX_GRAMOS_POR_DIPLOMA = 0.25

def obtener_config_diplomas():
    try:
        return dict(st.secrets["diplomas"])
    except Exception:
        return {}

def crear_codigo_diploma(area):
    area_limpia = (
        str(area or "GENE")
        .upper()
        .replace("Á", "A")
        .replace("É", "E")
        .replace("Í", "I")
        .replace("Ó", "O")
        .replace("Ú", "U")
        .replace("Ñ", "N")
    )
    abreviatura = re.sub(r"[^A-Z]", "", area_limpia)[:4] or "GENE"
    return f"CAV-{dt.date.today().year}-{abreviatura}-{uuid.uuid4().hex[:8].upper()}"

def crear_token_publico_diploma():
    return uuid.uuid4().hex + uuid.uuid4().hex

def obtener_url_base_aplicacion():
    """
    Obtiene la URL real desde la cual se está ejecutando Streamlit.

    Se prioriza la URL de la petición actual para evitar enlaces rotos por una
    configuración manual incorrecta, por ejemplo un subdominio que no existe.
    """
    candidatos = []

    # Streamlit moderno expone la URL completa de la sesión.
    try:
        url_contexto = str(st.context.url or "").strip()
        if url_contexto:
            candidatos.append(url_contexto)
    except Exception:
        pass

    # Respaldo mediante headers del proxy de Streamlit Cloud.
    try:
        headers = st.context.headers
        host = (
            headers.get("x-forwarded-host")
            or headers.get("host")
            or ""
        )
        proto = (
            headers.get("x-forwarded-proto")
            or "https"
        )
        if "," in proto:
            proto = proto.split(",", 1)[0].strip()
        if "," in host:
            host = host.split(",", 1)[0].strip()
        if host:
            candidatos.append(f"{proto}://{host}")
    except Exception:
        pass

    # Último respaldo: valor configurado manualmente en Secrets.
    config = obtener_config_diplomas()
    configurada = str(config.get("public_app_url", "")).strip()
    if configurada:
        candidatos.append(configurada)

    for candidata in candidatos:
        try:
            partes = urllib.parse.urlsplit(candidata)
            host = (partes.hostname or "").lower()
            if not partes.scheme or not partes.netloc:
                continue
            if host in {"localhost", "127.0.0.1", "0.0.0.0"}:
                continue
            return f"{partes.scheme}://{partes.netloc}".rstrip("/")
        except Exception:
            continue

    return ""

def construir_url_publica_diploma(token):
    base_url = obtener_url_base_aplicacion()
    if not base_url:
        return ""
    token_seguro = urllib.parse.quote(str(token), safe="")
    return f"{base_url}/?diploma={token_seguro}"

def reparar_urls_publicas_diplomas(registros):
    """
    Corrige registros antiguos que quedaron asociados a una URL equivocada.
    Solo actualiza la columna url_publica; no cambia códigos ni tokens.
    """
    base_url = obtener_url_base_aplicacion()
    if not base_url or not registros:
        return 0

    corregidos = 0
    for registro in registros:
        diploma_id = registro.get("id")
        token = registro.get("public_token")
        if not diploma_id or not token:
            continue

        url_correcta = f"{base_url}/?diploma={urllib.parse.quote(str(token), safe='')}"
        url_guardada = str(registro.get("url_publica") or "").strip()

        if url_guardada != url_correcta:
            try:
                supabase.table("diplomas_digitales").update({
                    "url_publica": url_correcta,
                    "actualizado_en": dt_datetime.now().isoformat(),
                }).eq("id", diploma_id).execute()
                registro["url_publica"] = url_correcta
                corregidos += 1
            except Exception as error:
                registrar_error("reparar_url_publica_diploma", error)

    return corregidos

def extraer_signed_url(respuesta):
    if isinstance(respuesta, str):
        return respuesta
    if isinstance(respuesta, dict):
        return (
            respuesta.get("signedURL")
            or respuesta.get("signedUrl")
            or respuesta.get("signed_url")
            or respuesta.get("url")
            or ""
        )
    data = getattr(respuesta, "data", None)
    if isinstance(data, dict):
        return (
            data.get("signedURL")
            or data.get("signedUrl")
            or data.get("signed_url")
            or data.get("url")
            or ""
        )
    return ""

def subir_archivo_diploma_storage(ruta, contenido, mime_type):
    """
    Sube PDF o PNG a Supabase Storage usando bytes reales.

    Algunas versiones de storage-py no reconocen io.BytesIO como archivo
    válido y lo intentan abrir como si fuera una ruta, provocando:
    "expected str, bytes or os.PathLike object, not BytesIO".
    """
    if isinstance(contenido, BytesIO):
        contenido = contenido.getvalue()
    elif hasattr(contenido, "read"):
        contenido = contenido.read()

    if isinstance(contenido, bytearray):
        contenido = bytes(contenido)

    if not isinstance(contenido, bytes):
        raise TypeError(
            "El contenido que se intenta subir debe ser bytes. "
            f"Tipo recibido: {type(contenido).__name__}"
        )

    if not contenido:
        raise ValueError(
            f"El archivo '{ruta}' está vacío y no puede subirse a Storage."
        )

    return (
        supabase.storage
        .from_(DIPLOMAS_BUCKET)
        .upload(
            path=str(ruta),
            file=contenido,
            file_options={
                "content-type": str(mime_type),
                "cache-control": "3600",
                "upsert": "true",
            },
        )
    )

def crear_url_firmada_diploma(ruta, expiracion=3600, descargar=False):
    if not ruta:
        return ""
    opciones = {"download": True} if descargar else None
    bucket = supabase.storage.from_(DIPLOMAS_BUCKET)
    if opciones:
        respuesta = bucket.create_signed_url(ruta, expiracion, opciones)
    else:
        respuesta = bucket.create_signed_url(ruta, expiracion)
    return extraer_signed_url(respuesta)

def registrar_evento_diploma(diploma_id, evento, detalle=None):
    try:
        supabase.table("diploma_eventos").insert({
            "diploma_id": diploma_id,
            "evento": evento,
            "usuario": (
                st.session_state.get("profesor_name")
                or st.session_state.get("role")
                or "sistema"
            ),
            "detalle": detalle or {},
            "fecha": dt_datetime.now().isoformat(),
        }).execute()
    except Exception as error:
        registrar_error("registrar_evento_diploma", error)

def crear_registro_diploma(datos, correo, origen="individual"):
    codigo = crear_codigo_diploma(datos.get("area"))
    token = crear_token_publico_diploma()
    url_publica = construir_url_publica_diploma(token)

    payload = {
        "codigo": codigo,
        "public_token": token,
        "nombre": str(datos.get("estudiante", "")).strip(),
        "correo": str(correo or "").strip().lower() or None,
        "curso": str(datos.get("curso", "")).strip(),
        "area": str(datos.get("area", "")).strip(),
        "titulo": str(datos.get("titulo", "")).strip(),
        "motivo": str(datos.get("motivo", "")).strip(),
        "profesor": str(datos.get("profesor", "")).strip(),
        "director": str(datos.get("director", "")).strip(),
        "fecha_diploma": str(datos.get("fecha")),
        "estado": "GENERADO",
        "origen": origen,
        "url_publica": url_publica or None,
        "creado_por": (
            st.session_state.get("profesor_name")
            or st.session_state.get("role")
            or "sistema"
        ),
        "creado_en": dt_datetime.now().isoformat(),
        "actualizado_en": dt_datetime.now().isoformat(),
    }

    respuesta = supabase.table("diplomas_digitales").insert(payload).execute()
    registro = (respuesta.data or [None])[0]
    if not registro:
        raise RuntimeError("Supabase no devolvió el registro del diploma.")

    registrar_evento_diploma(
        registro["id"],
        "GENERADO",
        {"codigo": codigo, "origen": origen},
    )
    return registro

def actualizar_archivos_diploma(diploma_id, pdf_path, preview_path):
    supabase.table("diplomas_digitales").update({
        "pdf_path": pdf_path,
        "preview_path": preview_path,
        "actualizado_en": dt_datetime.now().isoformat(),
    }).eq("id", diploma_id).execute()
    registrar_evento_diploma(
        diploma_id,
        "ARCHIVOS_GUARDADOS",
        {"pdf_path": pdf_path, "preview_path": preview_path},
    )

def marcar_diploma_enviado(diploma_id, correo):
    supabase.table("diplomas_digitales").update({
        "estado": "ENVIADO",
        "correo": str(correo).strip().lower(),
        "enviado_en": dt_datetime.now().isoformat(),
        "actualizado_en": dt_datetime.now().isoformat(),
        "error_envio": None,
    }).eq("id", diploma_id).execute()
    registrar_evento_diploma(
        diploma_id,
        "ENVIADO",
        {"correo": correo_enmascarado_global(correo)},
    )

def marcar_error_diploma(diploma_id, error):
    try:
        supabase.table("diplomas_digitales").update({
            "estado": "ERROR",
            "error_envio": str(error)[:1500],
            "actualizado_en": dt_datetime.now().isoformat(),
        }).eq("id", diploma_id).execute()
        registrar_evento_diploma(
            diploma_id,
            "ERROR",
            {"mensaje": str(error)[:500]},
        )
    except Exception as error_actualizacion:
        registrar_error("marcar_error_diploma", error_actualizacion)

def correo_enmascarado_global(correo):
    correo = str(correo or "")
    if "@" not in correo:
        return "***"
    usuario, dominio = correo.split("@", 1)
    visible = usuario[:2] if len(usuario) >= 2 else usuario[:1]
    return f"{visible}***@{dominio}"

def cargar_diplomas_registrados(creado_por=None):
    filtros = []
    if creado_por:
        filtros.append(("eq", "creado_por", creado_por))
    return select_paginado(
        "diplomas_digitales",
        "*",
        filtros=filtros,
        orden="creado_en",
        desc=True,
    )


def normalizar_texto_basico(valor):
    texto = str(valor or "").strip().lower()
    reemplazos = {
        "á": "a", "é": "e", "í": "i", "ó": "o", "ú": "u", "ñ": "n"
    }
    for origen, destino in reemplazos.items():
        texto = texto.replace(origen, destino)
    return texto

def obtener_tema_visual_diploma(area_texto):
    area = normalizar_texto_basico(area_texto)

    temas = [
        {
            "claves": ["lenguaje", "literatura", "comunicacion", "lectura", "escritura", "debate"],
            "nombre": "Lenguaje",
            "primario": "#7A1734",
            "secundario": "#D7A642",
            "fondo": "#FFF8F5",
            "panel": "#FFF3EA",
            "acento": "#4F0E22",
            "deco": ["📖", "✍️", "📝"],
        },
        {
            "claves": ["matematica", "calculo", "estadistica", "geometria"],
            "nombre": "Matemática",
            "primario": "#173B73",
            "secundario": "#D3A62E",
            "fondo": "#F7FAFF",
            "panel": "#ECF3FF",
            "acento": "#0B2851",
            "deco": ["∑", "π", "△"],
        },
        {
            "claves": ["ciencia", "biologia", "quimica", "fisica", "laboratorio", "cientifica"],
            "nombre": "Ciencias",
            "primario": "#126B70",
            "secundario": "#A8B83D",
            "fondo": "#F5FFFC",
            "panel": "#E7F8F0",
            "acento": "#084C50",
            "deco": ["🧪", "🔬", "⚛️"],
        },
        {
            "claves": ["arte", "artes", "pintura", "dibujo", "visual"],
            "nombre": "Artes",
            "primario": "#A23B72",
            "secundario": "#E39A37",
            "fondo": "#FFF7FB",
            "panel": "#FCEAF4",
            "acento": "#6F1E4A",
            "deco": ["🎨", "🖌️", "✨"],
        },
        {
            "claves": ["musica", "coro", "banda", "instrumental"],
            "nombre": "Música",
            "primario": "#5D3A8E",
            "secundario": "#D6A83B",
            "fondo": "#FBF8FF",
            "panel": "#EFE7FA",
            "acento": "#3A1E65",
            "deco": ["🎵", "🎼", "🎶"],
        },
        {
            "claves": ["tecnologia", "robotica", "computacion", "informatica", "programacion", "arduino", "ia"],
            "nombre": "Tecnología",
            "primario": "#1769AA",
            "secundario": "#31A8A0",
            "fondo": "#F7FBFF",
            "panel": "#E8F5FF",
            "acento": "#0E4778",
            "deco": ["💻", "🤖", "⚙️"],
        },
        {
            "claves": ["educacion fisica", "deporte", "voleibol", "futbol", "basquetbol", "atletismo"],
            "nombre": "Educación Física",
            "primario": "#237A57",
            "secundario": "#E1A833",
            "fondo": "#F7FFF8",
            "panel": "#E9F7EC",
            "acento": "#14533B",
            "deco": ["🏅", "⚽", "🏃"],
        },
        {
            "claves": ["teatro", "escenica", "danza"],
            "nombre": "Teatro",
            "primario": "#7B3049",
            "secundario": "#D7A642",
            "fondo": "#FFF9F7",
            "panel": "#FBEDE7",
            "acento": "#562034",
            "deco": ["🎭", "🎬", "🌟"],
        },
        {
            "claves": ["medioambiente", "ecologia", "ambiental", "verde", "reciclaje"],
            "nombre": "Medioambiente",
            "primario": "#2C6B2F",
            "secundario": "#97B93C",
            "fondo": "#F8FFF7",
            "panel": "#EAF7E7",
            "acento": "#1D4A20",
            "deco": ["🌿", "♻️", "🌎"],
        },
    ]

    for tema in temas:
        if any(clave in area for clave in tema["claves"]):
            return tema

    return {
        "nombre": "Institucional",
        "primario": "#800020",
        "secundario": "#C7A44A",
        "fondo": "#FFFDF8",
        "panel": "#FBF5E8",
        "acento": "#5E0017",
        "deco": ["🎓", "🏛️", "⭐"],
    }


def formatear_fecha_diploma_publica(valor):
    """
    Formatea fechas de Supabase sin depender de format_date_es(), que se
    declara más adelante en el archivo principal.
    """
    meses = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
    ]
    dias = [
        "lunes", "martes", "miércoles", "jueves",
        "viernes", "sábado", "domingo",
    ]

    try:
        if isinstance(valor, dt_datetime):
            fecha = valor.date()
        elif isinstance(valor, dt.date):
            fecha = valor
        else:
            fecha = pd.to_datetime(valor).date()

        return (
            f"{dias[fecha.weekday()]}, {fecha.day} de "
            f"{meses[fecha.month - 1]} de {fecha.year}"
        )
    except Exception:
        return str(valor or "")

def renderizar_diploma_publico(token):
    """
    Vista digital oficial V16.3.

    Se usa st.html() en lugar de st.markdown() para impedir que Markdown
    interprete las etiquetas HTML internas como un bloque de código.
    """
    try:
        respuesta = (
            supabase.table("diplomas_digitales")
            .select("*")
            .eq("public_token", str(token))
            .limit(1)
            .execute()
        )
        registros = respuesta.data or []
    except Exception as error:
        st.error("No fue posible consultar el diploma digital.")
        registrar_error("vista_publica_diploma", error)
        st.stop()

    if not registros:
        st.error("⚠️ El diploma consultado no existe o el enlace no es válido.")
        st.stop()

    diploma = registros[0]
    if diploma.get("estado") == "ANULADO":
        st.error("⚠️ Este diploma fue anulado por la institución.")
        st.stop()

    tema = obtener_tema_visual_diploma(diploma.get("area"))

    # Colores institucionales fijos.
    BURDEO_CAV = "#800020"
    BURDEO_OSCURO_CAV = "#590016"
    DORADO_CAV = "#C7A44A"

    fondo_tema = tema.get("fondo", "#FFFDF8")
    panel_tema = tema.get("panel", "#FBF5E8")
    acento_tema = tema.get("acento", BURDEO_CAV)

    try:
        pdf_url = crear_url_firmada_diploma(
            diploma.get("pdf_path"),
            expiracion=1800,
            descargar=True,
        )
    except Exception as error:
        pdf_url = ""
        registrar_error("firmar_urls_diploma", error)

    nombre = html_sanitizer.escape(str(diploma.get("nombre") or ""))
    titulo = html_sanitizer.escape(
        str(diploma.get("titulo") or "Diploma Digital")
    )
    curso = html_sanitizer.escape(str(diploma.get("curso") or ""))
    area = html_sanitizer.escape(str(diploma.get("area") or ""))
    motivo = html_sanitizer.escape(str(diploma.get("motivo") or ""))
    profesor = html_sanitizer.escape(str(diploma.get("profesor") or ""))
    director = html_sanitizer.escape(
        str(diploma.get("director") or "Director")
    )
    fecha = html_sanitizer.escape(
        formatear_fecha_diploma_publica(
            diploma.get("fecha_diploma") or dt.date.today()
        )
    )
    codigo = html_sanitizer.escape(str(diploma.get("codigo") or ""))
    area_label = html_sanitizer.escape(
        str(tema.get("nombre") or diploma.get("area") or "Institucional")
    )

    # Logo institucional embebido en la vista pública.
    logo_html = ""
    try:
        ruta_logo = Path(__file__).resolve().parent / "logocav.png"
        if ruta_logo.exists():
            logo_b64 = base64.b64encode(ruta_logo.read_bytes()).decode("ascii")
            logo_html = (
                '<img class="cav-logo" '
                f'src="data:image/png;base64,{logo_b64}" '
                'alt="Logo Colegio Antonio Varas">'
            )
    except Exception as error:
        registrar_error("logo_vista_publica_diploma", error)

    chips_html = "".join(
        f'<span class="cav-chip">{html_sanitizer.escape(str(icono))}</span>'
        for icono in (tema.get("deco") or [])[:3]
    )

    documento = f"""
<!doctype html>
<html lang="es">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<style>
:root {{
    --burdeo: {BURDEO_CAV};
    --burdeo-oscuro: {BURDEO_OSCURO_CAV};
    --dorado: {DORADO_CAV};
    --fondo-tema: {fondo_tema};
    --panel-tema: {panel_tema};
    --acento-tema: {acento_tema};
}}

* {{ box-sizing: border-box; }}

html, body {{
    margin: 0;
    padding: 0;
    color: #232833;
    background:
        radial-gradient(circle at 8% 4%, var(--panel-tema), transparent 25%),
        radial-gradient(circle at 96% 95%, rgba(199,164,74,.15), transparent 24%),
        linear-gradient(180deg, #fbfbfd, #f5f5f8);
    font-family: Inter, "Segoe UI", Arial, sans-serif;
}}

.cav-page {{
    width: min(1120px, calc(100% - 28px));
    margin: 18px auto 28px;
}}

.cav-topbar {{
    color: white;
    background: var(--burdeo);
    border-radius: 23px;
    padding: 24px 29px;
    border-bottom: 5px solid var(--dorado);
    box-shadow: 0 18px 40px rgba(60,0,16,.18);
    animation: cavFade .6s ease both;
}}

.cav-topbar h1 {{
    margin: 0;
    font-size: clamp(28px, 4.6vw, 43px);
    letter-spacing: -.025em;
}}

.cav-topbar p {{
    margin: 8px 0 0;
    opacity: .93;
    font-size: 1rem;
}}

.cav-card {{
    position: relative;
    isolation: isolate;
    overflow: hidden;
    margin-top: 20px;
    padding: 32px 42px 34px;
    background: var(--fondo-tema);
    border-radius: 28px;
    border: 1px solid rgba(0,0,0,.06);
    box-shadow: 0 20px 48px rgba(24,26,34,.09);
    animation: cavRise .8s ease both;
}}

.cav-card::before {{
    content: "";
    position: absolute;
    inset: 13px;
    border: 2px solid var(--dorado);
    border-radius: 22px;
    z-index: -1;
}}

.cav-card::after {{
    content: "";
    position: absolute;
    inset: 27px;
    border: 1px solid rgba(128,0,32,.15);
    border-radius: 17px;
    z-index: -1;
}}

.cav-version {{
    position: absolute;
    top: 24px;
    right: 31px;
    color: #7c6c45;
    font-size: .72rem;
    font-weight: 700;
    letter-spacing: .05em;
    text-transform: uppercase;
}}

.cav-badge {{
    display: inline-flex;
    align-items: center;
    min-height: 38px;
    padding: 8px 14px;
    color: var(--burdeo);
    background: var(--panel-tema);
    border: 1px solid rgba(128,0,32,.09);
    border-radius: 999px;
    font-size: .9rem;
    font-weight: 800;
}}

.cav-logo {{
    display: block;
    width: min(138px, 24vw);
    max-height: 138px;
    object-fit: contain;
    margin: 8px auto 8px;
    filter: drop-shadow(0 8px 12px rgba(0,0,0,.08));
}}

.cav-brand {{
    text-align: center;
    color: var(--burdeo);
}}

.cav-brand-kicker {{
    font-size: clamp(17px, 2.4vw, 25px);
    font-weight: 750;
}}

.cav-brand-name {{
    margin-top: 2px;
    font-size: clamp(23px, 3.2vw, 35px);
    font-weight: 900;
}}

.cav-gold-line {{
    width: min(690px, 72%);
    height: 2px;
    margin: 14px auto 25px;
    background: var(--dorado);
}}

.cav-title {{
    max-width: 920px;
    margin: 0 auto;
    color: var(--burdeo);
    text-align: center;
    font-size: clamp(34px, 5.2vw, 58px);
    font-weight: 900;
    line-height: 1.04;
    letter-spacing: -.035em;
    text-wrap: balance;
}}

.cav-intro {{
    margin-top: 17px;
    color: #6f7580;
    text-align: center;
    font-size: 1.08rem;
    font-style: italic;
}}

.cav-name {{
    max-width: 900px;
    margin: 15px auto 0;
    color: var(--burdeo);
    text-align: center;
    font-size: clamp(38px, 6.3vw, 68px);
    font-weight: 950;
    line-height: 1.03;
    letter-spacing: -.035em;
    text-wrap: balance;
}}

.cav-name-line {{
    width: min(820px, 82%);
    height: 2px;
    margin: 13px auto 14px;
    background: var(--dorado);
}}

.cav-meta {{
    color: var(--burdeo);
    text-align: center;
    font-size: 1.08rem;
    font-weight: 800;
}}

.cav-motivo {{
    max-width: 850px;
    margin: 22px auto 0;
    padding: 19px 24px;
    color: #3e4653;
    background: rgba(255,255,255,.9);
    border: 1px solid rgba(0,0,0,.075);
    border-radius: 18px;
    box-shadow: inset 0 1px 0 rgba(255,255,255,.9);
    text-align: center;
    font-size: clamp(17px, 2vw, 21px);
    line-height: 1.55;
}}

.cav-fecha {{
    margin-top: 17px;
    color: #606773;
    text-align: center;
    font-size: 1rem;
}}

.cav-icons {{
    display: flex;
    justify-content: center;
    gap: 15px;
    margin: 20px 0 2px;
}}

.cav-chip {{
    display: grid;
    place-items: center;
    width: 55px;
    height: 55px;
    background: white;
    border: 1px solid rgba(0,0,0,.075);
    border-radius: 17px;
    box-shadow: 0 9px 20px rgba(0,0,0,.065);
    font-size: 1.45rem;
    animation: cavFloat 3.4s ease-in-out infinite;
}}

.cav-chip:nth-child(2) {{ animation-delay: .35s; }}
.cav-chip:nth-child(3) {{ animation-delay: .7s; }}

.cav-footer {{
    display: grid;
    grid-template-columns: 1fr minmax(170px, auto) 1fr;
    align-items: end;
    gap: 24px;
    margin-top: 26px;
}}

.cav-sign {{
    padding-top: 16px;
    text-align: center;
}}

.cav-sign-line {{
    border-top: 2px solid #69717e;
    margin-bottom: 8px;
}}

.cav-sign strong {{
    color: var(--burdeo);
    font-size: 1rem;
}}

.cav-sign small {{
    display: block;
    margin-top: 4px;
    color: #737985;
}}

.cav-codebox {{
    min-width: 175px;
    padding: 13px 14px;
    background: white;
    border: 1px solid rgba(0,0,0,.075);
    border-radius: 17px;
    text-align: center;
}}

.cav-code-title {{
    color: var(--burdeo);
    font-weight: 850;
}}

.cav-code {{
    display: inline-block;
    margin-top: 8px;
    padding: 7px 11px;
    color: var(--burdeo);
    background: var(--panel-tema);
    border-radius: 999px;
    font-size: .85rem;
    font-weight: 850;
}}

.cav-code-help {{
    margin-top: 7px;
    color: #747b86;
    font-size: .8rem;
}}

.cav-eco {{
    margin-top: 20px;
    padding: 17px 19px;
    color: #174d29;
    background: #eff9f1;
    border: 1px solid #b8dfc1;
    border-radius: 17px;
}}

@keyframes cavRise {{
    from {{ opacity: 0; transform: translateY(14px); }}
    to {{ opacity: 1; transform: translateY(0); }}
}}

@keyframes cavFade {{
    from {{ opacity: 0; }}
    to {{ opacity: 1; }}
}}

@keyframes cavFloat {{
    0%, 100% {{ transform: translateY(0); }}
    50% {{ transform: translateY(-7px); }}
}}

@media (max-width: 760px) {{
    .cav-page {{
        width: min(100% - 16px, 1120px);
        margin-top: 8px;
    }}

    .cav-topbar {{
        padding: 21px 20px;
        border-radius: 19px;
    }}

    .cav-card {{
        padding: 27px 20px 28px;
        border-radius: 22px;
    }}

    .cav-version {{
        display: none;
    }}

    .cav-badge {{
        max-width: calc(100% - 20px);
        font-size: .8rem;
    }}

    .cav-footer {{
        grid-template-columns: 1fr;
    }}

    .cav-codebox {{
        order: -1;
    }}

    .cav-motivo {{
        padding: 17px 18px;
    }}
}}
</style>
</head>
<body>
<main class="cav-page">
    <header class="cav-topbar">
        <h1>🎓 Diploma Digital CAV</h1>
        <p>Reconocimiento oficial 100 % digital y amigable con el medio ambiente.</p>
    </header>

    <section class="cav-card">
        <span class="cav-version">Vista digital V16.3</span>
        <div class="cav-badge">{area_label} · Reconocimiento digital</div>

        {logo_html}

        <div class="cav-brand">
            <div class="cav-brand-kicker">Liceo Bicentenario de Excelencia</div>
            <div class="cav-brand-name">Colegio Antonio Varas</div>
        </div>

        <div class="cav-gold-line"></div>

        <h2 class="cav-title">{titulo}</h2>
        <div class="cav-intro">Se otorga el presente diploma a</div>
        <div class="cav-name">{nombre}</div>
        <div class="cav-name-line"></div>

        <div class="cav-meta">{curso} · {area}</div>

        <div class="cav-motivo">{motivo}</div>
        <div class="cav-fecha">Vicuña, {fecha}</div>

        <div class="cav-icons">{chips_html}</div>

        <div class="cav-footer">
            <div class="cav-sign">
                <div class="cav-sign-line"></div>
                <strong>{profesor}</strong>
                <small>Profesor responsable</small>
            </div>

            <div class="cav-codebox">
                <div class="cav-code-title">Código oficial</div>
                <div class="cav-code">{codigo}</div>
                <div class="cav-code-help">Diploma digital verificable</div>
            </div>

            <div class="cav-sign">
                <div class="cav-sign-line"></div>
                <strong>{director}</strong>
                <small>Director</small>
            </div>
        </div>
    </section>

    <div class="cav-eco">
        <strong>🌱 Acción ambiental institucional</strong><br>
        Este reconocimiento digital evita el uso de papel opalina y ayuda a
        reducir papel, tóner y residuos de impresión.
    </div>
</main>
</body>
</html>
"""

    # st.html renderiza HTML directamente y no lo procesa como Markdown.
    if hasattr(st, "html"):
        st.html(documento)
    else:
        # Respaldo para versiones antiguas de Streamlit.
        components.html(
            documento,
            height=1050,
            scrolling=True,
        )

    c1, c2, c3 = st.columns(3)
    c1.metric("Persona reconocida", diploma.get("nombre") or "—")
    c2.metric("Área", diploma.get("area") or "—")
    c3.metric("Estado", diploma.get("estado") or "—")

    if pdf_url:
        st.link_button(
            "📥 Descargar respaldo PDF",
            pdf_url,
            type="primary",
            use_container_width=True,
        )

    registrar_evento_diploma(
        diploma.get("id"),
        "VISTO",
        {"origen": "enlace_publico_v16_3"},
    )
    st.stop()

# La vista pública debe ejecutarse antes del login.
_token_publico_diploma = st.query_params.get("diploma")
if _token_publico_diploma:
    renderizar_diploma_publico(_token_publico_diploma)


def renderizar_monitor_semanal_publico():
    """
    Monitor semanal público de Enlaces.

    Se abre mediante:
        ?monitor=semana

    No requiere iniciar sesión y solo muestra información de agenda:
    fecha, horario, recurso, profesor y curso.
    """
    tz_monitor = pytz.timezone("America/Santiago")
    ahora = dt_datetime.now(tz_monitor)
    hoy = ahora.date()

    # Permite compartir semanas específicas:
    # ?monitor=semana&semana=2026-08-03
    semana_param = str(st.query_params.get("semana") or "").strip()
    if semana_param:
        try:
            fecha_referencia = pd.to_datetime(semana_param).date()
        except Exception:
            fecha_referencia = hoy
    else:
        fecha_referencia = hoy

    lunes = fecha_referencia - dt.timedelta(days=fecha_referencia.weekday())
    viernes = lunes + dt.timedelta(days=4)

    # Monitor público: refresco cada 2 minutos para reducir carga.
    st_autorefresh(
        interval=120_000,
        key=f"monitor_semanal_publico_{lunes.isoformat()}",
    )

    try:
        respuesta = (
            supabase.table("reservas")
            .select(
                "id, fecha, hora_inicio, hora_fin, "
                "profesores(nombre), cursos(nombre), recursos(nombre)"
            )
            .gte("fecha", lunes.isoformat())
            .lte("fecha", viernes.isoformat())
            .order("fecha")
            .order("hora_inicio")
            .execute()
        )
        reservas = respuesta.data or []
    except Exception as error:
        registrar_error("monitor_semanal_publico", error)
        reservas = []
        error_carga = str(error)
    else:
        error_carga = ""

    dias_nombres = {
        0: "Lunes",
        1: "Martes",
        2: "Miércoles",
        3: "Jueves",
        4: "Viernes",
    }

    agenda_por_dia = {
        lunes + dt.timedelta(days=i): []
        for i in range(5)
    }

    reservas_activas = []
    total_reservas = 0
    recursos_unicos = set()

    paleta = [
        "#800020",
        "#0F6B78",
        "#3659A8",
        "#2E7D4F",
        "#8952A8",
        "#B56818",
        "#A33B5B",
    ]

    def color_para_recurso(nombre):
        nombre = str(nombre or "Recurso")
        indice = sum(ord(caracter) for caracter in nombre) % len(paleta)
        return paleta[indice]

    for registro in reservas:
        try:
            fecha_reserva = pd.to_datetime(registro.get("fecha")).date()
        except Exception:
            continue

        if fecha_reserva not in agenda_por_dia:
            continue

        hora_inicio = str(registro.get("hora_inicio") or "")[:5]
        hora_fin = str(registro.get("hora_fin") or "")[:5]
        profesor = (
            (registro.get("profesores") or {}).get("nombre")
            or "Sin profesor"
        )
        curso = (
            (registro.get("cursos") or {}).get("nombre")
            or "Sin curso"
        )
        recurso = (
            (registro.get("recursos") or {}).get("nombre")
            or "Recurso de Enlaces"
        )

        item = {
            "id": registro.get("id"),
            "fecha": fecha_reserva,
            "hora_inicio": hora_inicio,
            "hora_fin": hora_fin,
            "profesor": str(profesor),
            "curso": str(curso),
            "recurso": str(recurso),
            "color": color_para_recurso(recurso),
        }
        agenda_por_dia[fecha_reserva].append(item)
        total_reservas += 1
        recursos_unicos.add(str(recurso))

        if fecha_reserva == hoy and hora_inicio and hora_fin:
            hora_actual = ahora.strftime("%H:%M")
            if hora_inicio <= hora_actual <= hora_fin:
                reservas_activas.append(item)

    for fecha_dia in agenda_por_dia:
        agenda_por_dia[fecha_dia].sort(
            key=lambda item: (
                item["hora_inicio"],
                item["recurso"].lower(),
            )
        )

    def escapar(valor):
        return html_sanitizer.escape(str(valor or ""))

    def construir_url_semana(fecha_lunes):
        base_url = obtener_url_base_aplicacion()
        ruta = (
            f"?monitor=semana&semana={fecha_lunes.isoformat()}"
        )
        return f"{base_url}/{ruta}" if base_url else ruta

    url_anterior = construir_url_semana(lunes - dt.timedelta(days=7))
    url_actual = construir_url_semana(hoy - dt.timedelta(days=hoy.weekday()))
    url_siguiente = construir_url_semana(lunes + dt.timedelta(days=7))
    url_login = obtener_url_base_aplicacion() or "/"

    logo_html = ""
    try:
        ruta_logo = Path(__file__).resolve().parent / "logocav.png"
        if ruta_logo.exists():
            logo_b64 = base64.b64encode(
                ruta_logo.read_bytes()
            ).decode("ascii")
            logo_html = (
                '<img class="monitor-logo" '
                f'src="data:image/png;base64,{logo_b64}" '
                'alt="Logo Colegio Antonio Varas">'
            )
    except Exception as error:
        registrar_error("logo_monitor_semanal", error)

    if reservas_activas:
        ahora_html = "".join(
            f"""
            <div class="monitor-now-item">
                <span class="monitor-now-dot"></span>
                <div>
                    <strong>{escapar(item["recurso"])}</strong>
                    <span>
                        {escapar(item["hora_inicio"])}–{escapar(item["hora_fin"])}
                        · {escapar(item["curso"])}
                        · {escapar(item["profesor"])}
                    </span>
                </div>
            </div>
            """
            for item in reservas_activas
        )
        ahora_estado = (
            f'<div class="monitor-now-list">{ahora_html}</div>'
        )
    else:
        ahora_estado = """
        <div class="monitor-now-empty">
            No hay una reserva en curso en este momento.
        </div>
        """

    columnas_dias = []
    for fecha_dia, items in agenda_por_dia.items():
        es_hoy = fecha_dia == hoy
        clase_hoy = " monitor-day-today" if es_hoy else ""

        if items:
            tarjetas = []
            for item in items:
                tarjetas.append(
                    f"""
                    <article
                        class="monitor-reserva"
                        style="--recurso-color:{item['color']};"
                    >
                        <div class="monitor-time">
                            {escapar(item["hora_inicio"])}
                            <span>–</span>
                            {escapar(item["hora_fin"])}
                        </div>
                        <div class="monitor-resource">
                            {escapar(item["recurso"])}
                        </div>
                        <div class="monitor-detail">
                            <span>👨‍🏫</span>
                            {escapar(item["profesor"])}
                        </div>
                        <div class="monitor-detail">
                            <span>📚</span>
                            {escapar(item["curso"])}
                        </div>
                    </article>
                    """
                )
            contenido_dia = "".join(tarjetas)
        else:
            contenido_dia = """
            <div class="monitor-empty">
                <span>✓</span>
                Sin reservas
            </div>
            """

        columnas_dias.append(
            f"""
            <section class="monitor-day{clase_hoy}">
                <header class="monitor-day-header">
                    <div>
                        <div class="monitor-day-name">
                            {dias_nombres[fecha_dia.weekday()]}
                        </div>
                        <div class="monitor-day-date">
                            {fecha_dia.strftime("%d/%m/%Y")}
                        </div>
                    </div>
                    {
                        '<span class="monitor-today-pill">HOY</span>'
                        if es_hoy else ""
                    }
                </header>
                <div class="monitor-day-list">
                    {contenido_dia}
                </div>
            </section>
            """
        )

    mensaje_error = ""
    if error_carga:
        mensaje_error = """
        <div class="monitor-error">
            No fue posible actualizar la agenda en este momento.
            La pantalla volverá a intentarlo automáticamente.
        </div>
        """

    periodo = (
        f"{lunes.strftime('%d/%m')} al "
        f"{viernes.strftime('%d/%m/%Y')}"
    )

    documento = f"""
    <!doctype html>
    <html lang="es">
    <head>
    <meta charset="utf-8">
    <meta
        name="viewport"
        content="width=device-width, initial-scale=1"
    >
    <style>
    :root {{
        --burdeo:#800020;
        --burdeo-oscuro:#590016;
        --dorado:#C7A44A;
        --fondo:#F5F7FA;
        --texto:#1F2937;
        --suave:#667085;
    }}

    * {{ box-sizing:border-box; }}

    html, body {{
        margin:0;
        padding:0;
        font-family:Inter, "Segoe UI", Arial, sans-serif;
        color:var(--texto);
        background:
            radial-gradient(
                circle at 7% 0%,
                rgba(199,164,74,.14),
                transparent 24%
            ),
            linear-gradient(180deg, #FCFCFD, var(--fondo));
    }}

    .monitor-page {{
        width:min(1500px, calc(100% - 24px));
        margin:12px auto 28px;
    }}

    .monitor-header {{
        display:grid;
        grid-template-columns:auto 1fr auto;
        align-items:center;
        gap:18px;
        padding:20px 24px;
        color:white;
        background:
            linear-gradient(
                135deg,
                var(--burdeo),
                var(--burdeo-oscuro)
            );
        border-radius:24px;
        border-bottom:5px solid var(--dorado);
        box-shadow:0 16px 36px rgba(70,0,18,.18);
    }}

    .monitor-logo {{
        width:70px;
        height:70px;
        object-fit:contain;
        padding:5px;
        background:white;
        border-radius:16px;
    }}

    .monitor-title {{
        margin:0;
        font-size:clamp(25px, 4vw, 42px);
        line-height:1.02;
        letter-spacing:-.025em;
    }}

    .monitor-subtitle {{
        margin-top:6px;
        color:rgba(255,255,255,.88);
        font-size:clamp(13px, 1.6vw, 17px);
    }}

    .monitor-clock {{
        min-width:174px;
        padding:12px 15px;
        text-align:center;
        color:var(--burdeo);
        background:white;
        border-radius:17px;
        font-weight:800;
    }}

    .monitor-clock-time {{
        display:block;
        margin-top:2px;
        font-size:1.35rem;
        font-variant-numeric:tabular-nums;
    }}

    .monitor-toolbar {{
        display:flex;
        justify-content:space-between;
        align-items:center;
        gap:12px;
        margin:16px 0;
        padding:12px 14px;
        background:white;
        border:1px solid rgba(0,0,0,.06);
        border-radius:18px;
        box-shadow:0 7px 18px rgba(0,0,0,.04);
    }}

    .monitor-week {{
        color:var(--burdeo);
        font-size:1.05rem;
        font-weight:850;
    }}

    .monitor-nav {{
        display:flex;
        gap:8px;
        flex-wrap:wrap;
    }}

    .monitor-nav a {{
        display:inline-flex;
        align-items:center;
        justify-content:center;
        min-height:39px;
        padding:8px 13px;
        color:var(--burdeo);
        background:#FFF8E8;
        border:1px solid rgba(128,0,32,.10);
        border-radius:12px;
        text-decoration:none;
        font-size:.9rem;
        font-weight:800;
    }}

    .monitor-nav a:hover {{
        background:#F7E8BE;
    }}

    .monitor-stats {{
        display:grid;
        grid-template-columns:repeat(3, minmax(0, 1fr));
        gap:12px;
        margin-bottom:14px;
    }}

    .monitor-stat {{
        padding:14px 16px;
        background:white;
        border:1px solid rgba(0,0,0,.06);
        border-radius:17px;
    }}

    .monitor-stat-label {{
        color:var(--suave);
        font-size:.78rem;
        font-weight:750;
        text-transform:uppercase;
        letter-spacing:.05em;
    }}

    .monitor-stat-value {{
        margin-top:4px;
        color:var(--burdeo);
        font-size:1.5rem;
        font-weight:900;
    }}

    .monitor-now {{
        margin-bottom:14px;
        padding:15px 17px;
        color:#174A2A;
        background:#EDF9F0;
        border:1px solid #B9DEC2;
        border-radius:18px;
    }}

    .monitor-now-title {{
        margin-bottom:8px;
        font-weight:900;
    }}

    .monitor-now-list {{
        display:grid;
        gap:7px;
    }}

    .monitor-now-item {{
        display:flex;
        align-items:flex-start;
        gap:9px;
    }}

    .monitor-now-dot {{
        width:10px;
        height:10px;
        margin-top:5px;
        flex:none;
        background:#20A45A;
        border-radius:999px;
        box-shadow:0 0 0 5px rgba(32,164,90,.13);
    }}

    .monitor-now-item strong {{
        display:block;
    }}

    .monitor-now-item span {{
        color:#41624B;
        font-size:.9rem;
    }}

    .monitor-now-empty {{
        color:#41624B;
    }}

    .monitor-grid {{
        display:grid;
        grid-template-columns:repeat(5, minmax(0, 1fr));
        gap:12px;
        align-items:start;
    }}

    .monitor-day {{
        min-width:0;
        overflow:hidden;
        background:white;
        border:1px solid rgba(0,0,0,.07);
        border-radius:20px;
        box-shadow:0 9px 22px rgba(0,0,0,.045);
    }}

    .monitor-day-today {{
        border:2px solid var(--dorado);
        box-shadow:0 11px 28px rgba(199,164,74,.16);
    }}

    .monitor-day-header {{
        display:flex;
        justify-content:space-between;
        align-items:center;
        gap:8px;
        padding:14px 14px 12px;
        color:white;
        background:
            linear-gradient(
                135deg,
                var(--burdeo),
                #A41A42
            );
    }}

    .monitor-day-name {{
        font-size:1.05rem;
        font-weight:900;
    }}

    .monitor-day-date {{
        margin-top:2px;
        color:rgba(255,255,255,.82);
        font-size:.8rem;
    }}

    .monitor-today-pill {{
        padding:5px 8px;
        color:var(--burdeo);
        background:var(--dorado);
        border-radius:999px;
        font-size:.68rem;
        font-weight:950;
    }}

    .monitor-day-list {{
        display:grid;
        gap:9px;
        padding:10px;
    }}

    .monitor-reserva {{
        position:relative;
        overflow:hidden;
        padding:11px 11px 12px 15px;
        background:#FCFCFD;
        border:1px solid rgba(0,0,0,.06);
        border-radius:14px;
    }}

    .monitor-reserva::before {{
        content:"";
        position:absolute;
        inset:0 auto 0 0;
        width:5px;
        background:var(--recurso-color);
    }}

    .monitor-time {{
        color:var(--recurso-color);
        font-size:.84rem;
        font-weight:900;
        font-variant-numeric:tabular-nums;
    }}

    .monitor-time span {{
        color:#9AA1AC;
        margin:0 2px;
    }}

    .monitor-resource {{
        margin-top:4px;
        color:#202733;
        font-size:.98rem;
        font-weight:900;
        line-height:1.16;
        overflow-wrap:anywhere;
    }}

    .monitor-detail {{
        display:flex;
        gap:6px;
        margin-top:6px;
        color:#5E6672;
        font-size:.79rem;
        line-height:1.25;
    }}

    .monitor-empty {{
        display:grid;
        place-items:center;
        min-height:94px;
        color:#8A929E;
        font-size:.84rem;
        text-align:center;
    }}

    .monitor-empty span {{
        display:grid;
        place-items:center;
        width:30px;
        height:30px;
        margin-bottom:4px;
        color:#2E8B57;
        background:#ECF8F0;
        border-radius:999px;
        font-weight:900;
    }}

    .monitor-error {{
        margin-bottom:14px;
        padding:13px 15px;
        color:#7B241C;
        background:#FDEDEC;
        border:1px solid #F5B7B1;
        border-radius:15px;
    }}

    .monitor-footer {{
        display:flex;
        justify-content:space-between;
        gap:12px;
        margin-top:15px;
        color:#777F8B;
        font-size:.78rem;
    }}

    .monitor-footer a {{
        color:var(--burdeo);
        font-weight:800;
        text-decoration:none;
    }}

    @media (max-width:1100px) {{
        .monitor-grid {{
            grid-template-columns:repeat(2, minmax(0, 1fr));
        }}

        .monitor-day:last-child {{
            grid-column:1 / -1;
        }}
    }}

    @media (max-width:720px) {{
        .monitor-page {{
            width:min(100% - 12px, 1500px);
            margin-top:6px;
        }}

        .monitor-header {{
            grid-template-columns:auto 1fr;
            padding:16px;
            border-radius:18px;
        }}

        .monitor-logo {{
            width:56px;
            height:56px;
            border-radius:13px;
        }}

        .monitor-clock {{
            grid-column:1 / -1;
            width:100%;
            min-width:0;
        }}

        .monitor-toolbar {{
            align-items:flex-start;
            flex-direction:column;
        }}

        .monitor-nav {{
            width:100%;
        }}

        .monitor-nav a {{
            flex:1;
            padding:8px 7px;
            font-size:.8rem;
        }}

        .monitor-stats {{
            grid-template-columns:repeat(3, 1fr);
            gap:7px;
        }}

        .monitor-stat {{
            padding:10px;
        }}

        .monitor-stat-value {{
            font-size:1.2rem;
        }}

        .monitor-grid {{
            grid-template-columns:1fr;
        }}

        .monitor-day:last-child {{
            grid-column:auto;
        }}

        .monitor-day-header {{
            position:sticky;
            top:0;
            z-index:2;
        }}

        .monitor-footer {{
            flex-direction:column;
            text-align:center;
        }}
    }}
    </style>
    </head>
    <body>
    <main class="monitor-page">
        <header class="monitor-header">
            {logo_html}
            <div>
                <h1 class="monitor-title">
                    Monitor Semanal de Enlaces
                </h1>
                <div class="monitor-subtitle">
                    Reservas de laboratorios, salas y recursos tecnológicos
                </div>
            </div>
            <div class="monitor-clock">
                Actualizado
                <span class="monitor-clock-time">
                    {ahora.strftime("%H:%M")}
                </span>
            </div>
        </header>

        <nav class="monitor-toolbar">
            <div class="monitor-week">
                🗓️ Semana del {periodo}
            </div>
            <div class="monitor-nav">
                <a href="{escapar(url_anterior)}">← Anterior</a>
                <a href="{escapar(url_actual)}">Semana actual</a>
                <a href="{escapar(url_siguiente)}">Siguiente →</a>
            </div>
        </nav>

        {mensaje_error}

        <section class="monitor-stats">
            <div class="monitor-stat">
                <div class="monitor-stat-label">
                    Reservas
                </div>
                <div class="monitor-stat-value">
                    {total_reservas}
                </div>
            </div>
            <div class="monitor-stat">
                <div class="monitor-stat-label">
                    Recursos
                </div>
                <div class="monitor-stat-value">
                    {len(recursos_unicos)}
                </div>
            </div>
            <div class="monitor-stat">
                <div class="monitor-stat-label">
                    En curso
                </div>
                <div class="monitor-stat-value">
                    {len(reservas_activas)}
                </div>
            </div>
        </section>

        <section class="monitor-now">
            <div class="monitor-now-title">
                🟢 En curso ahora
            </div>
            {ahora_estado}
        </section>

        <section class="monitor-grid">
            {''.join(columnas_dias)}
        </section>

        <footer class="monitor-footer">
            <span>
                Actualización automática cada 60 segundos ·
                Horario de Chile
            </span>
            <a href="{escapar(url_login)}">
                Volver al acceso principal
            </a>
        </footer>
    </main>
    </body>
    </html>
    """

    st.markdown(
        """
        <style>
        [data-testid="stHeader"],
        [data-testid="stSidebar"],
        [data-testid="stToolbar"],
        #MainMenu,
        footer {
            display:none !important;
        }

        .block-container {
            max-width:none !important;
            width:100% !important;
            padding:0 !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    if hasattr(st, "html"):
        st.html(documento)
    else:
        components.html(
            documento,
            height=1350,
            scrolling=True,
        )

    st.stop()


_modo_monitor_publico = str(
    st.query_params.get("monitor") or ""
).strip().lower()

if _modo_monitor_publico in {
    "semana",
    "semanal",
    "enlaces",
}:
    renderizar_monitor_semanal_publico()


TZ_CHILE = pytz.timezone("America/Santiago")

def convertir_datetime_chile(valor, por_defecto=None):
    """Convierte fechas de Supabase a datetime consciente en America/Santiago."""
    if valor in (None, ""):
        return por_defecto
    try:
        marca = pd.to_datetime(valor)
        fecha_python = marca.to_pydatetime() if hasattr(marca, "to_pydatetime") else marca
        if fecha_python.tzinfo is None:
            return TZ_CHILE.localize(fecha_python)
        return fecha_python.astimezone(TZ_CHILE)
    except Exception:
        return por_defecto

def combinar_fecha_hora_chile(fecha_valor, hora_valor):
    """Une date + time y aplica la zona horaria de Chile."""
    combinado = dt_datetime.combine(fecha_valor, hora_valor)
    return TZ_CHILE.localize(combinado)

def obtener_ventana_alerta(registro, ahora=None):
    """
    Devuelve (inicio, fin, estado).
    estado: EN_CURSO, PROGRAMADA, FINALIZADA o INVALIDA.
    Mantiene compatibilidad con alertas antiguas que solo tienen expiracion.
    """
    ahora = ahora or dt_datetime.now(TZ_CHILE)
    inicio = convertir_datetime_chile(registro.get("inicio_programado"))
    fin = convertir_datetime_chile(
        registro.get("fin_programado") or registro.get("expiracion")
    )

    if fin is None:
        return inicio, fin, "INVALIDA"

    if inicio is None:
        inicio = ahora - dt.timedelta(days=3650)

    if ahora < inicio:
        estado = "PROGRAMADA"
    elif inicio <= ahora <= fin:
        estado = "EN_CURSO"
    else:
        estado = "FINALIZADA"

    return inicio, fin, estado

def formato_fecha_hora_chile(valor):
    fecha = convertir_datetime_chile(valor)
    return fecha.strftime("%d/%m/%Y %H:%M") if fecha else "Sin fecha"

# ==============================================================================
# 📺 PANTALLA INFORMATIVA PÚBLICA (MODO KIOSCO) - MOTOR DE SONIDO INTEGRADO
# ==============================================================================
if st.session_state.get("ver_pantalla_tv", False):
    # Configuración de refresco y escala.
    # La escala real se conserva en una clave que no pertenece al slider.
    # Así Streamlit no la elimina cuando la alerta ejecuta st.stop().
    refresh_count = st_autorefresh(interval=30000, key="tv_refresh_global")

    if "tv_scale_saved" not in st.session_state:
        st.session_state.tv_scale_saved = 100

    escala_pct = int(st.session_state.get("tv_scale_saved", 100))

    # La pantalla se adapta automáticamente con CSS. El control manual queda
    # solo como ajuste fino y no puede volver a dejar la pantalla diminuta.
    if escala_pct < 80 or escala_pct > 130:
        escala_pct = 100

    escala_pct = max(80, min(130, escala_pct))
    st.session_state.tv_scale_saved = escala_pct
    escala = escala_pct / 100.0

    def guardar_escala_tv():
        valor = int(st.session_state.get("tv_scale_widget", 100))
        st.session_state.tv_scale_saved = max(80, min(130, valor))

    # 🕒 Sincronización absoluta con la hora local de Chile
    import pytz
    tz_chile = pytz.timezone("America/Santiago")
    now_dt = dt_datetime.now(tz_chile)
    hoy_str = now_dt.strftime("%Y-%m-%d")
    hora_actual_str = now_dt.strftime("%H:%M")
    
    # Preparar Logo
    ruta_logo = "logotv.png"
    logo_src = "<i class='ph-fill ph-shield-check header-logo-fallback'></i>"
    if os.path.exists(ruta_logo):
        with open(ruta_logo, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()
            logo_src = f"<img src='data:image/png;base64,{b64}' class='header-logo-img'/>"

    # 🚨 1. PRIORIDAD ABSOLUTA: ALERTA ROJA (999)
    # Admite alertas inmediatas y programadas por fecha/hora.
    # Si existen varias superpuestas, se muestra la más reciente.
    try:
        respuesta_alertas = (
            supabase.table("anuncios_urgentes")
            .select("*")
            .eq("is_active", True)
            .eq("prioridad", 999)
            .order("id", desc=True)
            .limit(100)
            .execute()
        )

        alerta = None
        ahora_chile = dt_datetime.now(TZ_CHILE)

        for candidata in (respuesta_alertas.data or []):
            inicio_alerta, fin_alerta, estado_alerta = obtener_ventana_alerta(
                candidata,
                ahora_chile,
            )

            if estado_alerta == "EN_CURSO" and alerta is None:
                alerta = candidata
            elif estado_alerta in ("FINALIZADA", "INVALIDA"):
                supabase.table("anuncios_urgentes").update(
                    {"is_active": False}
                ).eq("id", candidata["id"]).execute()

        if alerta:
            descripcion_alerta = html_sanitizer.escape(
                str(alerta.get("descripcion", ""))
            )

            st.markdown(f"""
                <style>
                html, body, .stApp,
                [data-testid="stAppViewContainer"],
                [data-testid="stMain"] {{
                    width: 100% !important;
                    min-width: 100% !important;
                    min-height: 100% !important;
                    zoom: 1 !important;
                    transform: none !important;
                }}

                .stApp {{
                    background: #ff0000 !important;
                    overflow: hidden !important;
                }}

                [data-testid="stHeader"],
                [data-testid="stSidebar"],
                [data-testid="stToolbar"],
                #MainMenu,
                footer {{
                    display: none !important;
                }}

                .block-container {{
                    max-width: none !important;
                    width: 100% !important;
                    padding: 0 !important;
                    margin: 0 !important;
                }}

                .alerta-total {{
                    position: fixed;
                    inset: 0;
                    z-index: 999999;
                    width: 100vw;
                    height: 100vh;
                    box-sizing: border-box;
                    display: flex;
                    flex-direction: column;
                    align-items: center;
                    justify-content: center;
                    padding: 4vh 5vw;
                    overflow: hidden;
                    background: #ff0000;
                    color: white;
                    text-align: center;
                    font-family: Inter, Arial, sans-serif;
                }}

                .at-titulo {{
                    font-size: clamp(46px, 7vw, 105px);
                    line-height: 1;
                    font-weight: 950;
                    text-shadow: 4px 4px 10px rgba(0,0,0,0.48);
                    animation: alertaPulso 1s infinite;
                }}

                .at-msg {{
                    width: min(92vw, 1500px);
                    margin-top: 4vh;
                    font-size: clamp(28px, 4.2vw, 64px);
                    line-height: 1.12;
                    font-weight: 750;
                    overflow-wrap: anywhere;
                    text-wrap: balance;
                    text-shadow: 2px 2px 7px rgba(0,0,0,0.35);
                }}

                @keyframes alertaPulso {{
                    0%, 100% {{ opacity: 1; transform: scale(1); }}
                    50% {{ opacity: .72; transform: scale(.985); }}
                }}
                </style>

                <div class="alerta-total">
                    <div class="at-titulo">⚠️ AVISO URGENTE ⚠️</div>
                    <div class="at-msg">{descripcion_alerta}</div>
                </div>
            """, unsafe_allow_html=True)

            alerta_id = alerta.get("id")
            if (
                os.path.exists("alarma.mp3")
                and st.session_state.get("ultima_alerta_roja_sonada") != alerta_id
            ):
                st.audio("alarma.mp3", format="audio/mp3", autoplay=True)
                st.session_state.ultima_alerta_roja_sonada = alerta_id

            st.stop()

    except Exception as e:
        registrar_error("alerta_roja_tv", e)

    # 📺 2. PROCESAMIENTO Y FILTRADO DE DATOS (Antes de renderizar)
    eventos = []
    avisos_vivos = []
    perfil = st.session_state.get("tv_profile", "General")

    # A. Cargar Cronograma (Eventos Generales)
    if st.session_state.get("url_calendario_tv"):
        for indice_cal, evento_cal in enumerate(
            obtener_eventos_google_calendar(st.session_state.url_calendario_tv)
        ):
            hora_cal = str(
                evento_cal.get("display_hora")
                or evento_cal.get("hora_sort")
                or "TODO EL DÍA"
            )
            eventos.append({
                "id_unico": f"cal_{indice_cal}_{evento_cal.get('titulo', 'evento')}",
                "hora_sort": str(evento_cal.get("hora_sort", "00:00")),
                "rango": hora_cal,
                "titulo": f"🗓️ {evento_cal.get('titulo', 'Evento especial')}",
                "desc": str(evento_cal.get("descripcion", "") or ""),
                "tipo": "calendario",
            })
    try:
        try:
            eventos_db = (
                supabase.table("eventos_tv")
                .select("*")
                .eq("is_active", True)
                .lte("fecha_inicio", hoy_str)
                .gte("fecha_fin", hoy_str)
                .order("hora_inicio")
                .execute()
                .data or []
            )
        except Exception:
            # Compatibilidad temporal antes de ejecutar la migración.
            eventos_db = (
                supabase.table("eventos_tv")
                .select("*")
                .eq("fecha_evento", hoy_str)
                .eq("is_active", True)
                .order("hora_inicio")
                .execute()
                .data or []
            )

        for e in eventos_db:
            hora_fin_ev = str(e.get("hora_fin", "23:59"))[:5]
            hora_ini_ev = str(e.get("hora_inicio", "00:00"))[:5]
            fecha_ini_ev = str(e.get("fecha_inicio") or e.get("fecha_evento") or hoy_str)
            fecha_fin_ev = str(e.get("fecha_fin") or e.get("fecha_evento") or hoy_str)

            if hora_actual_str <= hora_fin_ev:
                descripcion_evento = str(e.get("descripcion", "") or "").strip()
                if fecha_ini_ev != fecha_fin_ev:
                    try:
                        periodo_evento = (
                            f"Vigente del "
                            f"{dt_datetime.strptime(fecha_ini_ev, '%Y-%m-%d').strftime('%d/%m')} "
                            f"al {dt_datetime.strptime(fecha_fin_ev, '%Y-%m-%d').strftime('%d/%m')}"
                        )
                    except Exception:
                        periodo_evento = f"Vigencia: {fecha_ini_ev} al {fecha_fin_ev}"
                    descripcion_evento = (
                        f"{descripcion_evento} · {periodo_evento}"
                        if descripcion_evento
                        else periodo_evento
                    )

                eventos.append({
                    "id_unico": f"ev_{e['id']}",
                    "hora_sort": hora_ini_ev,
                    "rango": f"{hora_ini_ev} - {hora_fin_ev}",
                    "titulo": f"📢 {e['titulo']}",
                    "desc": descripcion_evento,
                    "tipo": "evento"
                })
        
        # B. Cargar Cronograma (Reservas según Perfil Técnico Real)
        if perfil in ["Profesores / PIE", "Inspectoría / UTP"]:
            res_res = supabase.table("reservas").select("*, profesores(nombre), recursos(nombre), cursos(nombre)").eq("fecha", hoy_str).execute()
            for r in (res_res.data or []):
                hora_fin_res = str(r.get("hora_fin", "23:59"))[:5]
                hora_ini_res = str(r.get("hora_inicio", "00:00"))[:5]
                if hora_actual_str <= hora_fin_res:
                    eventos.append({
                        "id_unico": f"res_{r['id']}",
                        "hora_sort": hora_ini_res, 
                        "rango": f"{hora_ini_res} - {hora_fin_res}", 
                        "titulo": f"🔒 {r['recursos']['nombre']} ➔ {r['cursos']['nombre']}", 
                        "desc": f"Docente: {r['profesores']['nombre']}", 
                        "tipo": "reserva"
                    })
    except Exception as e:
        st.error(f"Error cargando cronograma: {e}")

    eventos = sorted(eventos, key=lambda x: x['hora_sort'])

    # C. Cargar avisos vigentes. Se ordenan por prioridad y por registro reciente.
    try:
        avisos = (
            supabase.table("anuncios_urgentes")
            .select("*")
            .eq("is_active", True)
            .neq("prioridad", 999)
            .execute()
            .data or []
        )

        ahora_avisos = dt_datetime.now(TZ_CHILE)

        for aviso in avisos:
            inicio_aviso, fin_aviso, estado_aviso = obtener_ventana_alerta(
                aviso,
                ahora_avisos,
            )

            if estado_aviso == "EN_CURSO":
                aviso["_inicio_tv"] = inicio_aviso
                aviso["_expiracion_tv"] = fin_aviso
                aviso["_estado_tv"] = estado_aviso
                avisos_vivos.append(aviso)

            elif estado_aviso in ("FINALIZADA", "INVALIDA"):
                # Limpieza automática de avisos vencidos o sin fecha válida.
                try:
                    supabase.table("anuncios_urgentes").update(
                        {"is_active": False}
                    ).eq("id", aviso["id"]).execute()
                except Exception as error_limpieza:
                    registrar_error("limpiar_aviso_tv", error_limpieza)

        avisos_vivos = sorted(
            avisos_vivos,
            key=lambda aviso: (
                int(aviso.get("prioridad", 2) or 2),
                -int(aviso.get("id", 0) or 0),
            ),
        )
    except Exception as e:
        registrar_error("cargar_avisos_tv", e)

    # 🎵 3. DETECTOR INTELIGENTE DE NUEVOS ELEMENTOS (SISTEMA DE AUDIO)
    if "tv_elementos_vistos" not in st.session_state:
        st.session_state.tv_elementos_vistos = None

    # Recopilar firmas de todo lo actual
    firmas_actuales = set()
    for ev in eventos:
        firmas_actuales.add(ev.get("id_unico", f"ev_fallback_{ev['titulo']}"))
    for av in avisos_vivos:
        firmas_actuales.add(f"aviso_{av['id']}")

    reproducir_notificacion = False
    if st.session_state.tv_elementos_vistos is None:
        # Estado base inicial para no sonar al encender la pantalla
        st.session_state.tv_elementos_vistos = firmas_actuales
    else:
        # Si hay algo actual que no estaba guardado en memoria, ¡es nuevo!
        nuevos = firmas_actuales - st.session_state.tv_elementos_vistos
        if nuevos:
            reproducir_notificacion = True
        # Mantener memoria actualizada de lo que se queda en pantalla
        st.session_state.tv_elementos_vistos = firmas_actuales

    # 📺 4. INTERFAZ GRÁFICA (ESTILOS Y MAQUETACIÓN)
    st.markdown(f"""
    <style>
        @import url('https://unpkg.com/@phosphor-icons/web@2.1.1/src/fill/style.css');

        :root {{
            --tv-scale: {escala};
            --tv-page-gap: clamp(8px, min(1.2vw, 1.8vh), 24px);
            --tv-card-gap: clamp(9px, min(1vw, 1.4vh), 18px);
            --tv-header-pad-y: clamp(10px, min(1.2vw, 1.7vh), 18px);
            --tv-header-pad-x: clamp(14px, min(2vw, 2.2vh), 30px);
            --tv-heading-size:
                clamp(1.25rem, min(2.15vw, 3.15vh), 2.55rem);
            --tv-subtitle-size:
                clamp(.78rem, min(1vw, 1.55vh), 1.05rem);
            --tv-notice-title-size:
                clamp(1.05rem, min(1.75vw, 2.55vh), 2.15rem);
            --tv-notice-body-size:
                clamp(.9rem, min(1.22vw, 1.9vh), 1.5rem);
            --tv-agenda-title-size:
                clamp(.9rem, min(1.05vw, 1.65vh), 1.2rem);
        }}

        html,
        body,
        .stApp,
        [data-testid="stAppViewContainer"],
        [data-testid="stMain"] {{
            width: 100% !important;
            min-width: 100% !important;
            min-height: 100dvh !important;
            margin: 0 !important;
            padding: 0 !important;
            overflow-x: hidden !important;
        }}

        .stApp {{
            background-color: #0f172a !important;
            color: #f8fafc;
            font-family: 'Inter', sans-serif;
        }}

        [data-testid="stHeader"],
        [data-testid="stSidebar"],
        [data-testid="stToolbar"],
        #MainMenu,
        footer {{
            display: none !important;
        }}

        [data-testid="stMainBlockContainer"],
        .block-container {{
            width: 100% !important;
            max-width: none !important;
            min-height: 100dvh !important;
            margin: 0 !important;
            padding: var(--tv-page-gap) !important;
        }}

        /* El autorefresh no debe reservar altura visible. */
        iframe[title*="autorefresh"],
        iframe[src*="streamlit_autorefresh"] {{
            position: absolute !important;
            width: 0 !important;
            height: 0 !important;
            min-height: 0 !important;
            opacity: 0 !important;
            border: 0 !important;
            pointer-events: none !important;
        }}

        div[data-testid="stElementContainer"]:has(
            iframe[title*="autorefresh"]
        ),
        div[data-testid="stElementContainer"]:has(
            iframe[src*="streamlit_autorefresh"]
        ) {{
            display: none !important;
            width: 0 !important;
            height: 0 !important;
            min-height: 0 !important;
            margin: 0 !important;
            padding: 0 !important;
            overflow: hidden !important;
        }}

        [data-testid="stHorizontalBlock"] {{
            gap: clamp(12px, 1.45vw, 30px) !important;
            align-items: flex-start !important;
        }}

        [data-testid="column"] {{
            min-width: 0 !important;
        }}

        .tv-header {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            gap: clamp(12px, 1.5vw, 28px);
            min-height: clamp(72px, 11vh, 126px);
            margin: 0 0 clamp(9px, 1.7vh, 22px) 0;
            padding:
                var(--tv-header-pad-y)
                var(--tv-header-pad-x);
            color: #0f172a;
            background:
                linear-gradient(
                    135deg,
                    #ffffff 0%,
                    #f8fafc 100%
                );
            border: 1px solid #e2e8f0;
            border-radius: clamp(14px, 1.4vw, 22px);
            box-shadow: 0 10px 30px rgba(0,0,0,.34);
        }}

        .header-logo-img {{
            width: auto;
            height:
                clamp(
                    54px,
                    min(7.5vw, 10vh),
                    112px
                );
            object-fit: contain;
        }}

        .header-info {{
            display: flex;
            align-items: center;
            justify-content: flex-end;
            gap: clamp(10px, 1.45vw, 26px);
            min-width: 0;
            color: #334155;
            font-size:
                calc(
                    clamp(
                        .82rem,
                        min(1.32vw, 2.05vh),
                        1.55rem
                    )
                    * var(--tv-scale)
                );
            font-weight: 800;
            flex-wrap: wrap;
        }}

        .time-highlight {{
            padding:
                clamp(5px, .55vh, 8px)
                clamp(10px, .95vw, 17px);
            color: #2563eb !important;
            background: #eff6ff;
            border: 1px solid #bfdbfe;
            border-radius: 10px;
            font-weight: 900;
            white-space: nowrap;
        }}

        .progress-bar {{
            width: 0%;
            height: clamp(3px, .55vh, 6px);
            margin-top: calc(-1 * clamp(4px, .7vh, 9px));
            margin-bottom: clamp(8px, 1.6vh, 20px);
            background:
                linear-gradient(
                    90deg,
                    #3b82f6,
                    #60a5fa
                );
            border-radius: 10px;
            animation: load 20s linear infinite;
        }}

        @keyframes load {{
            0% {{ width: 0%; }}
            100% {{ width: 100%; }}
        }}

        @keyframes slideIn {{
            from {{
                opacity: 0;
                transform: translateY(14px);
            }}
            to {{
                opacity: 1;
                transform: translateY(0);
            }}
        }}

        .stExpander {{
            margin-top: clamp(8px, 1.4vh, 18px) !important;
            background-color: #1e293b !important;
            border: 1px solid #3b82f640 !important;
            border-radius: 14px !important;
        }}

        .stExpander * {{
            color: #f8fafc !important;
        }}

        .stExpander summary {{
            font-size:
                calc(
                    clamp(.88rem, min(1vw, 1.55vh), 1.08rem)
                    * var(--tv-scale)
                ) !important;
            font-weight: 800 !important;
        }}

        .tv-section-head {{
            display: flex;
            align-items: flex-start;
            justify-content: space-between;
            gap: clamp(10px, 1vw, 18px);
            margin: 0 0 clamp(8px, 1.2vh, 14px) 0;
        }}

        .tv-section-title {{
            margin: 0 !important;
            color: #f8fafc !important;
            font-size:
                calc(
                    var(--tv-heading-size)
                    * var(--tv-scale)
                ) !important;
            line-height: 1.08 !important;
            font-weight: 900 !important;
            text-shadow: 2px 2px 5px rgba(0,0,0,.34);
            text-wrap: balance;
        }}

        .tv-section-subtitle {{
            margin-top: clamp(3px, .55vh, 6px);
            color: #94a3b8;
            font-size:
                calc(
                    var(--tv-subtitle-size)
                    * var(--tv-scale)
                );
            font-weight: 650;
        }}

        .tv-count-badge {{
            flex-shrink: 0;
            padding:
                clamp(5px, .65vh, 8px)
                clamp(9px, .8vw, 13px);
            color: #e2e8f0;
            background: rgba(255,255,255,.08);
            border: 1px solid rgba(255,255,255,.16);
            border-radius: 999px;
            font-size:
                calc(
                    clamp(.68rem, min(.82vw, 1.25vh), .92rem)
                    * var(--tv-scale)
                );
            font-weight: 850;
            white-space: nowrap;
        }}

        .tv-notice-card {{
            position: relative;
            overflow: hidden;
            min-height: 0;
            margin-bottom: var(--tv-card-gap);
            padding:
                clamp(14px, min(1.45vw, 2.1vh), 27px)
                clamp(16px, min(1.55vw, 2.1vh), 29px)
                clamp(13px, min(1.35vw, 1.9vh), 24px)
                clamp(20px, min(1.75vw, 2.25vh), 33px);
            border: 1px solid rgba(255,255,255,.78);
            border-radius: clamp(14px, 1.3vw, 21px);
            box-shadow: 0 10px 25px rgba(0,0,0,.25);
            animation: slideIn .42s ease-out;
        }}

        .tv-notice-card::before {{
            content: "";
            position: absolute;
            inset: 0 auto 0 0;
            width: clamp(7px, .7vw, 11px);
        }}

        .tv-notice-high {{
            color: #1e293b;
            background:
                linear-gradient(
                    135deg,
                    #fff1f2 0%,
                    #ffffff 72%
                );
        }}

        .tv-notice-high::before {{
            background: #e11d48;
        }}

        .tv-notice-medium {{
            color: #1e293b;
            background:
                linear-gradient(
                    135deg,
                    #fffbeb 0%,
                    #ffffff 72%
                );
        }}

        .tv-notice-medium::before {{
            background: #d97706;
        }}

        .tv-notice-top {{
            display: flex;
            align-items: flex-start;
            justify-content: space-between;
            gap: clamp(10px, 1vw, 18px);
        }}

        .tv-notice-heading {{
            display: flex;
            align-items: center;
            gap: clamp(8px, .9vw, 14px);
            min-width: 0;
        }}

        .tv-notice-symbol {{
            display: flex;
            align-items: center;
            justify-content: center;
            width:
                clamp(
                    38px,
                    min(4vw, 6vh),
                    54px
                );
            height:
                clamp(
                    38px,
                    min(4vw, 6vh),
                    54px
                );
            flex: 0 0 auto;
            background: rgba(255,255,255,.82);
            border: 1px solid rgba(15,23,42,.08);
            border-radius: clamp(10px, 1vw, 15px);
            font-size:
                calc(
                    clamp(1.05rem, min(1.5vw, 2.3vh), 1.55rem)
                    * var(--tv-scale)
                );
        }}

        .tv-notice-title {{
            color: #0f172a;
            font-size:
                calc(
                    var(--tv-notice-title-size)
                    * var(--tv-scale)
                );
            line-height: 1.11;
            font-weight: 950;
            overflow-wrap: anywhere;
            text-wrap: balance;
        }}

        .tv-priority-pill {{
            flex-shrink: 0;
            padding:
                clamp(5px, .6vh, 8px)
                clamp(8px, .75vw, 13px);
            border-radius: 999px;
            font-size:
                calc(
                    clamp(.62rem, min(.72vw, 1.1vh), .83rem)
                    * var(--tv-scale)
                );
            font-weight: 950;
            letter-spacing: .04em;
            white-space: nowrap;
        }}

        .tv-priority-high {{
            color: #9f1239;
            background: #ffe4e6;
            border: 1px solid #fecdd3;
        }}

        .tv-priority-medium {{
            color: #92400e;
            background: #fef3c7;
            border: 1px solid #fde68a;
        }}

        .tv-notice-text {{
            display: flex;
            flex-direction: column;
            gap: clamp(5px, .8vh, 10px);
            margin:
                clamp(10px, 1.3vh, 19px)
                0 0
                clamp(0px, 4.3vw, 61px);
            color: #334155;
            font-size:
                calc(
                    var(--tv-notice-body-size)
                    * var(--tv-scale)
                );
            line-height: 1.36;
            font-weight: 700;
            text-align: left;
            overflow-wrap: anywhere;
        }}

        .tv-notice-line {{
            display: block;
            width: 100%;
        }}

        .tv-notice-footer {{
            display: flex;
            justify-content: flex-end;
            margin:
                clamp(9px, 1.2vh, 18px)
                0 0
                clamp(0px, 4.3vw, 61px);
            color: #64748b;
            font-size:
                calc(
                    clamp(.67rem, min(.78vw, 1.18vh), .88rem)
                    * var(--tv-scale)
                );
            font-weight: 750;
        }}

        .tv-agenda-card {{
            margin-bottom: clamp(7px, .95vh, 12px);
            padding:
                clamp(10px, min(1vw, 1.45vh), 14px);
            background: #ffffff;
            border: 1px solid #e2e8f0;
            border-left: clamp(5px, .5vw, 7px) solid #3b82f6;
            border-radius: clamp(11px, 1vw, 15px);
            box-shadow: 0 5px 15px rgba(0,0,0,.18);
            animation: slideIn .38s ease-out;
        }}

        .tv-agenda-card.reserva {{
            border-left-color: #10b981;
        }}

        .tv-agenda-card.calendario {{
            border-left-color: #8b5cf6;
        }}

        .tv-agenda-time {{
            display: inline-flex;
            align-items: center;
            margin-bottom: clamp(4px, .55vh, 7px);
            padding:
                clamp(3px, .4vh, 5px)
                clamp(6px, .55vw, 9px);
            color: #1e40af;
            background: #eff6ff;
            border: 1px solid #bfdbfe;
            border-radius: 8px;
            font-size:
                calc(
                    clamp(.64rem, min(.75vw, 1.15vh), .86rem)
                    * var(--tv-scale)
                );
            font-weight: 900;
        }}

        .tv-agenda-title {{
            color: #0f172a;
            font-size:
                calc(
                    var(--tv-agenda-title-size)
                    * var(--tv-scale)
                );
            line-height: 1.18;
            font-weight: 900;
            overflow-wrap: anywhere;
        }}

        .tv-agenda-desc {{
            margin-top: clamp(3px, .45vh, 6px);
            color: #64748b;
            font-size:
                calc(
                    clamp(.68rem, min(.8vw, 1.2vh), .9rem)
                    * var(--tv-scale)
                );
            line-height: 1.23;
            font-weight: 650;
            overflow-wrap: anywhere;
        }}

        .tv-empty-state {{
            display: flex;
            min-height:
                clamp(
                    120px,
                    24vh,
                    220px
                );
            align-items: center;
            justify-content: center;
            flex-direction: column;
            gap: clamp(5px, .7vh, 9px);
            padding: clamp(18px, 2vw, 30px);
            color: #cbd5e1;
            background: rgba(30,41,59,.72);
            border: 1px dashed rgba(148,163,184,.45);
            border-radius: clamp(14px, 1.2vw, 19px);
            text-align: center;
        }}

        .tv-empty-icon {{
            font-size:
                calc(
                    clamp(1.55rem, min(2vw, 3vh), 2.25rem)
                    * var(--tv-scale)
                );
        }}

        .tv-empty-title {{
            color: #f8fafc;
            font-size:
                calc(
                    clamp(.9rem, min(1.05vw, 1.55vh), 1.2rem)
                    * var(--tv-scale)
                );
            font-weight: 900;
        }}

        .tv-empty-text {{
            color: #94a3b8;
            font-size:
                calc(
                    clamp(.7rem, min(.82vw, 1.2vh), .92rem)
                    * var(--tv-scale)
                );
            font-weight: 650;
        }}

        /* Pantallas con poca altura: compactar automáticamente. */
        @media (max-height: 850px) and (min-width: 1000px) {{
            :root {{
                --tv-page-gap: 8px;
                --tv-card-gap: 8px;
                --tv-heading-size:
                    clamp(1.12rem, min(1.9vw, 2.75vh), 2rem);
                --tv-notice-title-size:
                    clamp(.98rem, min(1.5vw, 2.15vh), 1.75rem);
                --tv-notice-body-size:
                    clamp(.82rem, min(1.05vw, 1.55vh), 1.2rem);
            }}

            .tv-header {{
                min-height: 66px;
                margin-bottom: 8px;
                padding: 8px 16px;
            }}

            .header-logo-img {{
                height: clamp(48px, 7vh, 72px);
            }}

            .progress-bar {{
                margin-bottom: 8px;
            }}

            .tv-notice-card {{
                padding: 12px 18px 11px 22px;
            }}

            .tv-notice-text {{
                margin-top: 8px;
                gap: 4px;
            }}

            .tv-notice-footer {{
                margin-top: 8px;
            }}

            .stExpander {{
                margin-top: 7px !important;
            }}
        }}

        /* Tablets y pantallas estrechas: columnas verticales. */
        @media (max-width: 1050px) {{
            [data-testid="stHorizontalBlock"] {{
                display: flex !important;
                flex-direction: column !important;
                gap: 14px !important;
            }}

            [data-testid="column"] {{
                width: 100% !important;
                min-width: 100% !important;
                flex: 1 1 100% !important;
            }}

            .tv-header {{
                align-items: flex-start;
                flex-direction: column;
            }}

            .header-info {{
                width: 100%;
                justify-content: flex-start;
            }}

            .tv-notice-text,
            .tv-notice-footer {{
                margin-left: 0;
            }}
        }}

        /* Celulares. */
        @media (max-width: 620px) {{
            [data-testid="stMainBlockContainer"],
            .block-container {{
                padding: 7px !important;
            }}

            .tv-header {{
                gap: 9px;
                border-radius: 14px;
            }}

            .header-logo-img {{
                max-width: 100%;
                height: auto;
                max-height: 68px;
            }}

            .header-info {{
                gap: 8px;
                font-size: .84rem;
            }}

            .tv-section-head {{
                align-items: flex-start;
            }}

            .tv-priority-pill {{
                display: none;
            }}

            .tv-count-badge {{
                max-width: 42%;
                white-space: normal;
                text-align: center;
            }}

            .tv-notice-card {{
                padding: 14px 14px 13px 18px;
            }}
        }}

        /* Pantallas 4K o monitores muy grandes. */
        @media (min-width: 2200px) and (min-height: 1200px) {{
            :root {{
                --tv-heading-size: 3rem;
                --tv-notice-title-size: 2.45rem;
                --tv-notice-body-size: 1.72rem;
                --tv-agenda-title-size: 1.35rem;
            }}

            [data-testid="stMainBlockContainer"],
            .block-container {{
                padding: 28px !important;
            }}
        }}

        @media (prefers-reduced-motion: reduce) {{
            .progress-bar,
            .tv-notice-card,
            .tv-agenda-card {{
                animation: none !important;
            }}
        }}
    </style>
    """, unsafe_allow_html=True)

    clima = obtener_clima_vicuna()
    st.markdown(f"""
    <div class="tv-header">
        <div>{logo_src}</div>
        <div class="header-info">
            <div><i class="ph-fill ph-calendar"></i> {now_dt.day}/{now_dt.month}</div>
            <div style="color:#cbd5e1">|</div>
            <div>{clima}</div>
            <div style="color:#cbd5e1">|</div>
            <div class="time-highlight"><i class="ph-fill ph-clock"></i> {hora_actual_str}</div>
        </div>
    </div><div class="progress-bar"></div>
    """, unsafe_allow_html=True)

    # Los avisos ocupan el espacio principal. La agenda queda como resumen lateral.
    col_avisos, col_agenda = st.columns([2.35, 1], gap="large")

    def formatear_vigencia_aviso_tv(aviso):
        try:
            expira = aviso.get("_expiracion_tv")
            if expira is None:
                expira = pd.to_datetime(aviso.get("expiracion"))
                if expira.tzinfo is None:
                    expira = tz_chile.localize(expira.to_pydatetime())
                else:
                    expira = expira.tz_convert(tz_chile)
            return expira.strftime("Vigente hasta %d/%m · %H:%M")
        except Exception:
            return "Aviso vigente"

    def formatear_descripcion_aviso_tv(texto):
        """Conserva los saltos de línea del formulario y genera HTML seguro."""
        contenido = str(texto or "")
        contenido = contenido.replace("\r\n", "\n").replace("\r", "\n")

        # Cada renglón escrito en el formulario se muestra como una línea separada.
        lineas = [linea.strip() for linea in contenido.split("\n") if linea.strip()]
        if not lineas and contenido.strip():
            lineas = [contenido.strip()]

        return "".join(
            f'<div class="tv-notice-line">{html_sanitizer.escape(linea)}</div>'
            for linea in lineas
        )

    # ==============================================================
    # ÁREA PRINCIPAL: AVISOS DESTACADOS
    # ==============================================================
    with col_avisos:
        total_avisos = len(avisos_vivos)

        if total_avisos:
            longitud_max_aviso = max(
                (
                    len(str(aviso.get("titulo", "") or ""))
                    + len(str(aviso.get("descripcion", "") or ""))
                )
                for aviso in avisos_vivos
            )
            AVISOS_POR_PAGINA = (
                1 if longitud_max_aviso > 430 else 2
            )

            paginas_avisos = max(
                1,
                (total_avisos + AVISOS_POR_PAGINA - 1) // AVISOS_POR_PAGINA,
            )
            pagina_avisos = refresh_count % paginas_avisos
            avisos_pagina = avisos_vivos[
                pagina_avisos * AVISOS_POR_PAGINA:
                (pagina_avisos + 1) * AVISOS_POR_PAGINA
            ]
            indicador_avisos = (
                f"{pagina_avisos + 1}/{paginas_avisos}"
                if paginas_avisos > 1
                else f"{total_avisos} vigente(s)"
            )
        else:
            avisos_pagina = []
            indicador_avisos = "Sin avisos"

        st.markdown(
            f"""
            <div class="tv-section-head">
                <div>
                    <h2 class="tv-section-title">🚨 Avisos importantes</h2>
                    <div class="tv-section-subtitle">
                        Información prioritaria para la comunidad educativa
                    </div>
                </div>
                <div class="tv-count-badge">{indicador_avisos}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        if not avisos_pagina:
            st.markdown(
                """
                <div class="tv-empty-state">
                    <div class="tv-empty-icon">✅</div>
                    <div class="tv-empty-title">No hay avisos pendientes</div>
                    <div class="tv-empty-text">
                        La comunidad se encuentra sin comunicaciones especiales vigentes.
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )
        else:
            for aviso in avisos_pagina:
                es_alta = str(aviso.get("prioridad", 2)) == "1"
                clase_tarjeta = "tv-notice-high" if es_alta else "tv-notice-medium"
                clase_prioridad = (
                    "tv-priority-high" if es_alta else "tv-priority-medium"
                )
                texto_prioridad = "PRIORIDAD ALTA" if es_alta else "INFORMACIÓN"
                simbolo = "⚠️" if es_alta else "📣"

                titulo_aviso = html_sanitizer.escape(
                    str(aviso.get("titulo", "Aviso institucional") or "Aviso institucional")
                )
                descripcion_aviso = formatear_descripcion_aviso_tv(
                    aviso.get("descripcion", "")
                )
                vigencia_aviso = html_sanitizer.escape(
                    formatear_vigencia_aviso_tv(aviso)
                )

                st.markdown(
                    f"""
                    <article class="tv-notice-card {clase_tarjeta}">
                        <div class="tv-notice-top">
                            <div class="tv-notice-heading">
                                <div class="tv-notice-symbol">{simbolo}</div>
                                <div class="tv-notice-title">{titulo_aviso}</div>
                            </div>
                            <div class="tv-priority-pill {clase_prioridad}">
                                {texto_prioridad}
                            </div>
                        </div>
                        <div class="tv-notice-text">{descripcion_aviso}</div>
                        <div class="tv-notice-footer">🕒 {vigencia_aviso}</div>
                    </article>
                    """,
                    unsafe_allow_html=True,
                )

    # ==============================================================
    # COLUMNA LATERAL: CRONOGRAMA COMPACTO + CONTROLES
    # ==============================================================
    with col_agenda:
        total_eventos = len(eventos)

        if total_eventos:
            longitud_max_agenda = max(
                (
                    len(str(item.get("titulo", "") or ""))
                    + len(str(item.get("desc", "") or ""))
                )
                for item in eventos
            )
            AGENDA_POR_PAGINA = (
                3 if longitud_max_agenda > 260 else 4
            )

            paginas_agenda = max(
                1,
                (total_eventos + AGENDA_POR_PAGINA - 1) // AGENDA_POR_PAGINA,
            )
            pagina_agenda = refresh_count % paginas_agenda
            items_agenda = eventos[
                pagina_agenda * AGENDA_POR_PAGINA:
                (pagina_agenda + 1) * AGENDA_POR_PAGINA
            ]
            indicador_agenda = (
                f"{pagina_agenda + 1}/{paginas_agenda}"
                if paginas_agenda > 1
                else f"{total_eventos} actividad(es)"
            )
        else:
            items_agenda = []
            indicador_agenda = "Sin actividades"

        st.markdown(
            f"""
            <div class="tv-section-head">
                <div>
                    <h2 class="tv-section-title">📅 Cronograma</h2>
                    <div class="tv-section-subtitle">Próximas actividades de hoy</div>
                </div>
                <div class="tv-count-badge">{indicador_agenda}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        if not items_agenda:
            st.markdown(
                f"""
                <div class="tv-empty-state" style="min-height:150px;">
                    <div class="tv-empty-icon">🗓️</div>
                    <div class="tv-empty-title">Sin actividades próximas</div>
                    <div class="tv-empty-text">
                        Perfil actual: {html_sanitizer.escape(str(perfil))}
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )
        else:
            for item in items_agenda:
                tipo_item = str(item.get("tipo", "evento"))
                clase_item = (
                    "reserva"
                    if tipo_item == "reserva"
                    else "calendario"
                    if tipo_item == "calendario"
                    else ""
                )
                titulo_item = html_sanitizer.escape(
                    str(item.get("titulo", "Actividad"))
                )
                rango_item = html_sanitizer.escape(
                    str(item.get("rango", item.get("hora_sort", "")))
                )
                descripcion_item = html_sanitizer.escape(
                    str(item.get("desc", "") or "")
                )

                st.markdown(
                    f"""
                    <div class="tv-agenda-card {clase_item}">
                        <div class="tv-agenda-time">🕒 {rango_item}</div>
                        <div class="tv-agenda-title">{titulo_item}</div>
                        <div class="tv-agenda-desc">{descripcion_item}</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

        with st.expander("⚙️ Ajustes de pantalla"):
            st.selectbox(
                "👁️ Perfil visual",
                ["General", "Profesores / PIE", "Inspectoría / UTP"],
                key="tv_profile",
            )

            if "tv_scale_widget" not in st.session_state:
                st.session_state.tv_scale_widget = int(
                    st.session_state.get("tv_scale_saved", 100)
                )

            st.slider(
                "🔍 Ajuste manual fino (%)",
                80,
                130,
                key="tv_scale_widget",
                step=5,
                on_change=guardar_escala_tv,
                help=(
                    "La adaptación principal es automática según el ancho y "
                    "alto de la pantalla. Este control solo permite un ajuste fino."
                ),
            )

        st.write("")
        if st.button(
            "🔙 VOLVER AL MENÚ PRINCIPAL",
            use_container_width=True,
            type="primary",
        ):
            st.session_state.ver_pantalla_tv = False
            st.rerun()

    # 🎛️ EJECUCIÓN FÍSICA DEL SONIDO DE NOTIFICACIÓN NUEVA
    if reproducir_notificacion and os.path.exists("notificacion.mp3"):
        st.audio("notificacion.mp3", format="audio/mp3", autoplay=True)

    st.stop()
# ──────────────────────────────────────────────────────────────────────────────
# 0) CONFIGURACIÓN GLOBAL Y ESTILO
# ──────────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
    :root {
        --primary-color: #0072C6;
        --background-color: #F8F9FA;
        --sidebar-background: #FFFFFF;
        --card-background: #FFFFFF;
        --text-color: #343A40;
        --subtle-text-color: #6C757D;
        --border-color: #DEE2E6;
        --hover-color: #E9ECEF;
        --card-shadow: 0 4px 6px rgba(0,0,0,0.05);
    }
    .main .block-container { padding: 2rem; }
    h1 {
        color: var(--primary-color);
        font-weight: 600;
        border-bottom: 2px solid var(--border-color);
        padding-bottom: 0.5rem;
        margin-bottom: 1.5rem;
    }
    [data-testid="stSidebar"] {
        background-color: var(--sidebar-background);
        border-right: 1px solid var(--border-color);
        padding: 1rem;
    }
    .st-emotion-cache-1r4qj8v, [data-testid="stForm"], [data-testid="stExpander"] {
        border: 1px solid var(--border-color);
        border-radius: 0.75rem;
        padding: 1.5rem;
        box-shadow: var(--card-shadow);
        background-color: var(--card-background);
    }
    [data-testid="stMetric"] {
        background-color: var(--card-background);
        border: 1px solid var(--border-color);
        border-radius: 0.75rem;
        padding: 1.5rem;
        box-shadow: var(--card-shadow);
    }
    [data-testid="stSidebarNav"] a:hover {
        background-color: var(--hover-color);
        color: var(--primary-color);
    }
    [data-testid="stSidebarNav"] a[aria-current="page"] {
        background-color: var(--primary-color);
        color: white;
    }
    @media (prefers-color-scheme: dark) {
        :root {
            --primary-color: #58A6FF;
            --background-color: #0D1117;
            --sidebar-background: #161B22;
            --card-background: #161B22;
            --text-color: #C9D1D9;
            --subtle-text-color: #8B949E;
            --border-color: #30363D;
            --hover-color: #252b33;
        }
        body, .stApp { background-color: var(--background-color); color: var(--text-color); }
        .st-emotion-cache-1r4qj8v, [data-testid="stForm"], [data-testid="stExpander"], [data-testid="stMetric"] { border-color: var(--border-color); }
        .tooltip-text { background-color: #f0f2f6 !important; color: #111 !important; }
    }
    .reservation-card { 
        border-radius: 5px; 
        padding: 6px; 
        margin-bottom: 4px;
        font-size: 0.8em; 
        line-height: 1.3; 
        word-wrap: break-word; 
        border: 1px solid rgba(0,0,0,0.1);
        position: relative;
        cursor: default;
    }
    .tooltip-text {
        visibility: hidden; width: 220px; background-color: #333; color: #fff; text-align: left;
        border-radius: 6px; padding: 10px; position: absolute; z-index: 10; bottom: 105%;
        left: 50%; margin-left: -110px; opacity: 0; transition: opacity 0.3s; box-shadow: 0 4px 8px rgba(0,0,0,0.2);
        pointer-events: none;
    }
    .reservation-card:hover .tooltip-text { visibility: visible; opacity: 1; }
    .dataframe td { vertical-align: top; }
</style>
""", unsafe_allow_html=True)

# ---- Funciones de utilidad ----
def parse_date(val):
    if isinstance(val, date) and not isinstance(val, dt_datetime): return val
    if isinstance(val, dt_datetime): return val.date()
    if isinstance(val, pd.Timestamp): return val.to_pydatetime().date()
    s = str(val).strip()
    for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%Y-%m-%d %H:%M:%S"):
        try: return dt_datetime.strptime(s, fmt).date()
        except (ValueError, TypeError): continue
    raise ValueError(f"Formato de fecha inválido: {val!r}")

def as_time(val):
    if isinstance(val, dt.time): return val
    if isinstance(val, dt.datetime): return val.time()
    if isinstance(val, str):
        s = val.strip()
        for fmt in ('%H:%M:%S', '%H:%M'):
            try: return dt_datetime.strptime(s, fmt).time()
            except ValueError: continue
    try: return pd.to_datetime(val).time()
    except (ValueError, TypeError): raise ValueError(f"Formato de hora inválido: {val!r}")

def overlap(s1, e1, s2, e2): return max(s1, s2) < min(e1, e2)

def get_color_from_string(input_string: str) -> str:
    if not input_string: return "#CCCCCC"
    hash_obj = hashlib.md5(input_string.encode())
    hash_int = int(hash_obj.hexdigest(), 16)
    hue = hash_int % 360; saturation = 75; lightness = 90
    return f"hsl({hue}, {saturation}%, {lightness}%)"

def sort_time_key(time_string):
    try:
        start_time_str = time_string.split(' a ')[0].strip()
        return dt.datetime.strptime(start_time_str, '%H:%M').time()
    except (ValueError, IndexError): return dt.time(23, 59)

def format_date_es(date_obj):
    dias = ["Lunes", "Martes", "Miércoles", "Jueves", "Viernes", "Sábado", "Domingo"]
    meses = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
    return f"{dias[date_obj.weekday()]}, {date_obj.day} de {meses[date_obj.month - 1]} de {date_obj.year}"

def send_email(subject, body, recipient_email):
    try:
        creds = st.secrets["email_credentials"]
        sender_email, password = creds["smtp_username"], creds["smtp_password"]
        msg = MIMEMultipart(); msg['From'] = sender_email; msg['To'] = recipient_email; msg['Subject'] = subject
        msg.attach(MIMEText(body, 'html'))
        server = smtplib.SMTP(creds["smtp_server"], creds["smtp_port"]); server.starttls()
        server.login(sender_email, password); server.send_message(msg); server.quit()
        st.toast(f"📧 Notificación enviada a {recipient_email}")
    except Exception as e:
        pass 

# ==============================================================================
# --- MODO PÚBLICO: ENRUTAMIENTO VÍA CÓDIGO QR ---
# ==============================================================================
# 1. Captura limpia de parámetros (SIN imprimir nada en pantalla)
query_params = st.query_params
page_param = query_params.get("page", "")

if page_param == "reporte":
    # 2. ESTILOS AESTHETIC (Ocultar menús y decorar)
    st.markdown("""
        <style>
            [data-testid="stHeader"] {display: none;}
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            .stApp { background-color: #F8FAFC; }
            [data-testid="stForm"] {
                background-color: #FFFFFF;
                border-radius: 16px;
                padding: 25px;
                box-shadow: 0 4px 15px rgba(0,0,0,0.05);
                border: 1px solid #E2E8F0;
            }
        </style>
    """, unsafe_allow_html=True)

    # 3. ENCABEZADO
    st.markdown("<h1 style='text-align: center; font-size: 50px;'>🏫</h1>", unsafe_allow_html=True)
    st.markdown("<h2 style='text-align: center; color: #1E3A8A; font-weight: 800; margin-bottom: 0px;'>🚨 Reportar Falla</h2>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: #64748b; font-size: 14px; margin-top: 5px;'>Colegio Antonio Varas - Depto. Enlaces</p>", unsafe_allow_html=True)
    st.write("") 
    
    # 4. LÓGICA DEL REPORTE MAESTRO
    try:
        # Traer todos los recursos para el selectbox
        recursos_activos = supabase.table("recursos").select("id, nombre").execute().data or []
        opciones_recursos = {r['nombre']: r['id'] for r in recursos_activos}
    except Exception as e:
        st.error(f"Error al cargar recursos: {e}")
        opciones_recursos = {}

    if not opciones_recursos:
        st.error("❌ No hay recursos en la base de datos para reportar.")
        st.stop()

    with st.form("form_reporte_publico", clear_on_submit=True):
        st.markdown("#### 📝 Detalles del Problema")
        st.markdown("<span style='color: #64748b; font-size: 14px;'>Ayúdanos con estos datos para que Soporte Técnico lo resuelva rápido.</span>", unsafe_allow_html=True)
        st.write("")
        
        # EL NUEVO SELECTBOX (Menú desplegable)
        recurso_seleccionado = st.selectbox("🖥️ Selecciona el equipo o sala que presenta la falla:", ["-- Selecciona un Equipo --"] + sorted(list(opciones_recursos.keys())))
        
        nombre_reporta = st.text_input("👤 Tu Nombre Completo:", placeholder="Ej. Juan Pérez")
        descripcion = st.text_area("🔧 Describe el problema detalladamente:", height=120, placeholder="Ej. El proyector no enciende y parpadea una luz roja...")
        
        st.write("")
        submit = st.form_submit_button("🚀 Enviar Reporte Técnico", type="primary", use_container_width=True)
        
        if submit:
            if recurso_seleccionado == "-- Selecciona un Equipo --":
                st.error("⚠️ Debes seleccionar un equipo de la lista.")
            elif not nombre_reporta.strip() or not descripcion.strip():
                st.error("⚠️ Faltan datos: Por favor ingresa tu nombre y la descripción.")
            else:
                with st.spinner("Enviando reporte a Enlaces..."):
                    try:
                        recurso_id = opciones_recursos[recurso_seleccionado]
                        supabase.table("mantenimientos").insert({
                            "recurso_id": recurso_id,
                            "descripcion": descripcion.strip(),
                            "estado": "Reportado (Vía QR)",
                            "reportado_por": nombre_reporta.strip(),
                            "fecha": str(dt.date.today())  
                        }).execute()
                        
                        st.success("✅ ¡Reporte enviado con éxito! Gracias por avisarnos.")
                        st.balloons() # 🎈 ¡AQUÍ ESTÁN LOS GLOBOS! 🎈
                        time.sleep(1.2)
                        st.rerun() # Limpia la pantalla para el siguiente reporte
                    except Exception as e:
                        st.error(f"Error técnico al guardar: {e}")

    st.stop() # Detiene la app aquí para que no cargue el Login del sistema interno
# ==============================================================================
# ------------------------------------------------------------------
# 1) INICIALIZACIÓN DE DATOS
# ------------------------------------------------------------------
if 'PROFESORES' not in globals(): PROFESORES = []
if 'RECURSOS' not in globals(): RECURSOS = []
if 'CURSOS' not in globals(): CURSOS = []

def custom_course_sort_key(course_name):
    course_name = str(course_name).strip()
    if 'Dif' in course_name: return (3, 0, course_name)
    match = re.match(r"(\d+)°\s*(BÁSICO|MEDIO)\s*([A-Z])?", course_name, re.IGNORECASE)
    if match:
        num, level, letter = match.groups()
        level_priority = 0 if 'BÁSICO' in level.upper() else 1
        return (level_priority, int(num), letter or '')
    return (4, 0, course_name)

@st.cache_data(ttl=600, show_spinner=False)
def cargar_datos_login():
    try:
        p_res = supabase.table("profesores").select("nombre").execute().data
        profs = sorted([p["nombre"] for p in p_res]) if p_res else []
        r_res = supabase.table("recursos").select("nombre").execute().data
        recs = sorted([r["nombre"] for r in r_res]) if r_res else []
        c_res = supabase.table("cursos").select("nombre").execute().data
        curs = sorted([c["nombre"] for c in c_res], key=custom_course_sort_key) if c_res else []
        return profs, recs, curs
    except:
        return [], [], []

PROFESORES, RECURSOS, CURSOS = cargar_datos_login()

# ------------------------------------------------------------------
# 2) SISTEMA DE LOGIN HORIZONTAL
# ------------------------------------------------------------------
if "logged" not in st.session_state:
    st.session_state.logged = False
    st.session_state.role = None
    st.session_state.profesor_name = None

if not st.session_state.logged:
    st.markdown("""
        <style>
            .block-container { padding-top: 3rem !important; }
            [data-testid="stVerticalBlock"] { gap: 0.5rem !important; }
            .login-card { padding: 20px; border-radius: 15px; border: 1px solid #eeeeee; background-color: white; box-shadow: 0 4px 6px rgba(0,0,0,0.05); }
            .stTextInput, .stSelectbox { margin-bottom: -10px; }
            label { font-size: 0.85rem !important; font-weight: 600 !important; }
        </style>
    """, unsafe_allow_html=True)

    main_container = st.container()
    
    with main_container:
        col_logo, col_form = st.columns([1, 1.8], gap="large")

        with col_logo:
            st.markdown("<br><br>", unsafe_allow_html=True)
            BASE_DIR = Path(__file__).parent
            logo_path = BASE_DIR / "logocav.png"
            if logo_path.exists(): 
                st.image(str(logo_path), use_container_width=True)
            else: 
                st.info("Logo CAV")

        with col_form:
            st.markdown("<h2 style='text-align: center; color: #1E3A8A; margin-bottom: 0px;'>SISTEMA CAV</h2>", unsafe_allow_html=True)
            st.markdown("<p style='text-align: center; color: gray; font-size: 0.9rem;'>Reservas, recursos, diplomas y monitor semanal · V23 · navegación estable + modo rendimiento</p>", unsafe_allow_html=True)
            
            with st.container(border=True):
                # Quitamos "Profesor" de las opciones. Ahora solo quedan Admin y Mensajería
                tipo_user = st.radio("Acceder como:", ["Administrador", "Profesores / Diplomas", "Mensajería Interna"], horizontal=True)
                st.markdown("---")

                if tipo_user == "Administrador":
                    with st.form("admin_form", clear_on_submit=True):
                        u_adm = st.text_input(
                            "Usuario (Soporte Técnico)",
                            placeholder="Ej: Edgar",
                            autocomplete="username",
                        )
                        p_adm = st.text_input(
                            "Contraseña",
                            type="password",
                            placeholder="••••••••",
                            autocomplete="current-password",
                        )
                        st.caption(
                            "🔐 Tu navegador puede ofrecer guardar esta clave "
                            "después de iniciar sesión."
                        )
                        
                        if st.form_submit_button("INICIAR SESIÓN", use_container_width=True, type="primary"):
                            try:
                                # 🔒 Llamada segura al panel de Secrets de Streamlit
                                clave_admin_segura = st.secrets["passwords"]["admin_pass"]
                            except KeyError:
                                st.error("❌ Error: Falta configurar [passwords] admin_pass en los Secrets de Streamlit.")
                                st.stop()

                            if u_adm.strip().upper() in ["EDGAR", "GLORIA", "CARLOS", "ALEXIS"] and p_adm == clave_admin_segura:
                                st.session_state.logged = True
                                st.session_state.role = "admin"
                                st.session_state.profesor_name = u_adm.strip().capitalize()
                                st.rerun()
                            else:
                                st.error("Acceso denegado. Credenciales incorrectas.")
                                
                elif tipo_user == "Mensajería Interna":
                    st.info("💡 Acceso protegido para la gestión de Pantallas Informativas.")
                    with st.form("mensajeria_form", clear_on_submit=True):
                        u_msg = st.text_input(
                            "Nombre del operador",
                            placeholder="Ej. Inspectoría",
                            autocomplete="username",
                        )
                        p_msg = st.text_input(
                            "PIN de Mensajería",
                            type="password",
                            placeholder="••••••",
                            autocomplete="current-password",
                        )
                        st.caption(
                            "🔐 Tu navegador puede ofrecer guardar este acceso."
                        )
                        if st.form_submit_button("ENTRAR AL PANEL DE MENSAJERÍA", use_container_width=True, type="primary"):
                            try:
                                clave_msg = st.secrets["passwords"]["mensajeria_pass"]
                            except KeyError:
                                st.error("❌ Falta configurar [passwords] mensajeria_pass en Streamlit Secrets.")
                                st.stop()
                            if u_msg.strip() and p_msg == clave_msg:
                                st.session_state.logged = True
                                st.session_state.role = "mensajeria"
                                st.session_state.profesor_name = u_msg.strip()
                                registrar_auditoria("inicio_sesion", "autenticacion", detalle={"rol": "mensajeria"})
                                st.rerun()
                            else:
                                st.error("Acceso denegado. Revisa el nombre y el PIN.")
                        
                else:
                    with st.form("profe_form", clear_on_submit=True):
                        u_profe = st.selectbox("Busca tu nombre", PROFESORES, index=None, placeholder="Selecciona...")
                        p_profe = st.text_input(
                            "Clave de Acceso",
                            type="password",
                            placeholder="••••",
                            autocomplete="current-password",
                        )
                        st.caption(
                            "🔐 El navegador puede ofrecer recordar la clave "
                            "en este dispositivo."
                        )
                        
                        if st.form_submit_button("ENTRAR AL PANEL", use_container_width=True, type="primary"):
                            try:
                                # 🔒 Llamada segura al panel de Secrets de Streamlit
                                clave_profe_segura = st.secrets["passwords"]["profe_pass"]
                            except KeyError:
                                st.error("❌ Error: Falta configurar [passwords] profe_pass en los Secrets de Streamlit.")
                                st.stop()

                            if u_profe and p_profe == clave_profe_segura:
                                st.session_state.logged = True
                                st.session_state.role = "profesor"
                                st.session_state.profesor_name = u_profe
                                st.rerun()
                            elif not u_profe:
                                st.warning("Por favor selecciona tu nombre")
                            else:
                                st.error("Contraseña incorrecta")
                
            # --- AQUÍ VA EL ACCESO PÚBLICO (Fuera de los formularios) ---
            st.markdown("---")
            st.markdown("<h4 style='text-align:center;'>Acceso Público</h4>", unsafe_allow_html=True)

            # Pantalla informativa pública dentro de la sesión actual.
            if st.button(
                "📺 Abrir Pantalla Informativa",
                use_container_width=True,
            ):
                st.session_state.ver_pantalla_tv = True
                st.rerun()

            # Monitor semanal mediante enlace público compartible.
            base_publica = obtener_url_base_aplicacion()
            enlace_monitor = (
                f"{base_publica}/?monitor=semana"
                if base_publica
                else "?monitor=semana"
            )

            st.link_button(
                "🗓️ Abrir Monitor Semanal de Enlaces",
                enlace_monitor,
                type="primary",
                use_container_width=True,
            )

            with st.expander(
                "🔗 Copiar enlace público del monitor",
                expanded=False,
            ):
                st.code(enlace_monitor, language=None)
                st.caption(
                    "Este enlace puede abrirse desde celulares, tablets, "
                    "televisores o computadores sin iniciar sesión."
                )
            # -------------------------------------------------------------

    # ESTO SIEMPRE DEBE IR AL FINAL DEL BLOQUE DE LOGIN
    st.stop()
    
# ------------------------------------------------------------------
# 3) CARGA DE LA BASE DE DATOS PRINCIPAL 
# ------------------------------------------------------------------
@st.cache_data(ttl=120, show_spinner=False)
def cargar_reservas_y_datos():
    horas_corregidas = [
        '8:00 a 8:45', '8:45 a 9:30', '8:00 a 9:30',
        '9:45 a 10:30', '10:30 a 11:15', '9:45 a 11:15',
        '11:30 a 12:15', '12:15 a 13:00', '11:30 a 13:00',
        '14:00 a 14:45', '14:45 a 15:30', '14:00 a 15:30',
        '14:00 a 16:30', '14:45 a 16:30', '15:45 a 16:30',
        '17:00 a 18:30', '17:30 a 18:30'
    ]
    
    try:
        res_data = select_paginado(
            "reservas",
            "id, fecha, hora_inicio, hora_fin, observaciones, profesores(nombre), cursos(nombre), recursos(nombre)",
            orden="fecha",
            desc=False
        )
        reservas_limpias = []
        for r in res_data:
            reservas_limpias.append({
                "id": r["id"],
                "Fecha": parse_date(r["fecha"]),
                "Hora inicio": as_time(r["hora_inicio"]),
                "Hora fin": as_time(r["hora_fin"]),
                "Profesor": r["profesores"]["nombre"] if r.get("profesores") else "",
                "Curso": r["cursos"]["nombre"] if r.get("cursos") else "",
                "Recurso": r["recursos"]["nombre"] if r.get("recursos") else "",
                "Observaciones": r["observaciones"]
            })
        df_res = pd.DataFrame(reservas_limpias) if reservas_limpias else pd.DataFrame(columns=['id', 'Fecha', 'Hora inicio', 'Hora fin', 'Profesor', 'Curso', 'Recurso', 'Observaciones'])
            
        try:
            # CORRECCIÓN DEFINITIVA DE MANTENIMIENTO: Carga simple sin JOIN para evitar errores
            mant_data = supabase.table("mantenimientos").select("*").execute().data
            rec_data = supabase.table("recursos").select("id, nombre").execute().data
            mant_map_rec = {r['id']: r['nombre'] for r in rec_data} if rec_data else {}
            
            df_mant = pd.DataFrame(mant_data) if mant_data else pd.DataFrame()
            
            if not df_mant.empty:
                df_mant = df_mant[df_mant['estado'].isin(['Reportado (Vía QR)', 'En Revisión'])]
                df_mant['FechaInicio_dt'] = df_mant['fecha'].apply(parse_date) if 'fecha' in df_mant.columns else dt.date.today()
                df_mant['FechaFin_dt'] = df_mant['FechaInicio_dt']
                df_mant['HoraInicio'] = dt.time(0, 0)
                df_mant['HoraFin'] = dt.time(23, 59)
                if 'recurso_id' in df_mant.columns:
                    df_mant['Recurso'] = df_mant['recurso_id'].apply(lambda x: mant_map_rec.get(x, 'Desconocido'))
                else: df_mant['Recurso'] = 'Desconocido'
            else:
                df_mant = pd.DataFrame(columns=['Recurso', 'FechaInicio_dt', 'HoraInicio', 'FechaFin_dt', 'HoraFin'])
        except Exception as e:
            df_mant = pd.DataFrame(columns=['Recurso', 'FechaInicio_dt', 'HoraInicio', 'FechaFin_dt', 'HoraFin'])

        return df_res, horas_corregidas, df_mant
    except Exception as e:
        return pd.DataFrame(columns=['id', 'Fecha', 'Hora inicio', 'Hora fin', 'Profesor', 'Curso', 'Recurso', 'Observaciones']), horas_corregidas, pd.DataFrame()

df, HORAS, df_mantenimiento = cargar_reservas_y_datos()

@st.cache_data(ttl=600, show_spinner=False)
def cargar_catalogos_runtime():
    """Catálogos de apoyo para formularios y correos."""
    map_prof_local = {}
    map_cur_local = {}
    map_rec_local = {}
    profesor_data_local = {}

    try:
        prof_data_db = (
            supabase.table("profesores")
            .select("id, nombre, email")
            .execute()
            .data
            or []
        )
        map_prof_local = {
            p["nombre"]: p["id"]
            for p in prof_data_db
        }
        profesor_data_local = {
            p["nombre"]: p.get("email", "")
            for p in prof_data_db
        }
    except Exception as error:
        registrar_error("catalogo_profesores", error)

    try:
        cursos_db = (
            supabase.table("cursos")
            .select("id, nombre")
            .execute()
            .data
            or []
        )
        map_cur_local = {
            c["nombre"]: c["id"]
            for c in cursos_db
        }
    except Exception as error:
        registrar_error("catalogo_cursos", error)

    try:
        recursos_db = (
            supabase.table("recursos")
            .select("id, nombre")
            .execute()
            .data
            or []
        )
        map_rec_local = {
            r["nombre"]: r["id"]
            for r in recursos_db
        }
    except Exception as error:
        registrar_error("catalogo_recursos", error)

    return (
        map_prof_local,
        map_cur_local,
        map_rec_local,
        profesor_data_local,
    )


map_prof, map_cur, map_rec, PROFESOR_DATA = cargar_catalogos_runtime()


@st.cache_data(ttl=90, show_spinner=False)
def cargar_resumen_inicio():
    """Resumen cacheado del Centro de Operaciones."""
    try:
        tickets = select_paginado(
            "mantenimientos",
            "*",
            orden="fecha",
            desc=True,
        )
    except Exception as error:
        registrar_error("inicio_tickets", error)
        tickets = []

    try:
        equipos = select_paginado("equipos", "*")
    except Exception as error:
        registrar_error("inicio_bajas", error)
        equipos = []

    return tickets, equipos


# ------------------------------------------------------------------
# 4) NAVEGACIÓN Y VISTAS
# ------------------------------------------------------------------
# V23: sin refresco automático global durante el trabajo privado.

sidebar_title = f"Panel de {st.session_state.role.capitalize()}"
if st.session_state.role == 'profesor': sidebar_title = f"Hola, {st.session_state.profesor_name.split(' ')[0]}"

BASE_DIR = Path(__file__).parent
logo_path = BASE_DIR / "logocav.png"

if logo_path.exists():
    st.markdown("""<style>[data-testid="stSidebarNav"]::before { content: ""; display: none; margin-top: 0px; } section[data-testid="stSidebar"] div.st-emotion-cache-16t70r2 { padding-top: 0.5rem !important; }</style>""", unsafe_allow_html=True)
    col1, col2, col3 = st.sidebar.columns([0.1, 2.8, 0.1])
    with col2: st.image(str(logo_path), use_container_width=True)

st.sidebar.markdown(f"<div style='text-align: center; color: var(--primary-color); font-weight: bold; margin-bottom: 0px; font-size: 1.1em; letter-spacing: 1px;'>{sidebar_title.upper()}</div>", unsafe_allow_html=True)
st.sidebar.markdown("<hr style='margin: 8px 0px 5px 0px; padding: 0;'>", unsafe_allow_html=True)

html_reloj = """
<!DOCTYPE html>
<html>
<head>
<style>
    body { margin: 0; padding: 0; font-family: "Source Sans Pro", sans-serif; text-align: center; display: flex; justify-content: center; align-items: center; overflow: hidden; }
    .container { padding: 0px; width: 100%; }
    @media (prefers-color-scheme: dark) { .container { background-color: transparent; color: #FAFAFA; } .reloj { color: #ff4b4b; } .mensaje { color: #c6c6d1; } }
    @media (prefers-color-scheme: light) { .container { background-color: transparent; color: #31333F; } .reloj { color: #ff4b4b; } .mensaje { color: #555555; } }
    .fecha { font-weight: 600; font-size: 0.9em; text-transform: capitalize; margin-bottom: 2px;}
    .reloj { font-size: 1.6em; font-weight: 700; margin: 2px 0; font-variant-numeric: tabular-nums;}
    .mensaje { font-size: 0.8em; font-style: italic; margin-top: 2px;}
</style>
</head>
<body>
    <div class="container">
        <div id="fecha" class="fecha"></div>
        <div id="reloj" class="reloj"></div>
        <div id="mensaje" class="mensaje"></div>
    </div>
    <script>
        const mensajes = ["¡Que tengas un excelente día! ☀️", "Cada día es una nueva oportunidad para brillar. ✨", "Tu esfuerzo de hoy es el éxito de mañana. 💪", "Haz que las cosas pasen. ¡Tú puedes! 🚀", "Pequeños pasos todos los días llevan a grandes resultados. 🏔️", "Sonríe, respira y sigue adelante. 🌻", "La actitud lo es todo. ¡A dar el 100%! 💯", "Hoy es un buen día para hacer la diferencia. 🌟", "Tu trabajo y dedicación son muy valiosos. 🤝", "¡Mucho éxito en todas tus tareas de hoy! 🎯"];
        document.getElementById("mensaje").innerText = mensajes[Math.floor(Math.random() * mensajes.length)];
        function actualizarReloj() {
            const ahora = new Date();
            document.getElementById("reloj").innerText = ahora.toLocaleTimeString('es-CL', { hour: '2-digit', minute: '2-digit', second: '2-digit' });
            document.getElementById("fecha").innerText = ahora.toLocaleDateString('es-CL', { weekday: 'long', year: 'numeric', month: 'long', day: 'numeric' });
        }
        setInterval(actualizarReloj, 1000); actualizarReloj();
    </script>
</body>
</html>
"""

# ==============================================================================
with st.sidebar:
    components.html(html_reloj, height=85)
    st.markdown("<hr style='margin: 0px 0px 10px 0px; padding: 0;'>", unsafe_allow_html=True)

    PAGES_CONFIG = {
        "Inicio": {"icon": "🏠", "roles": ["admin", "profesor", "mensajeria"]},
        "Mis Reservas": {"icon": "👤", "roles": ["profesor"]},
        "Registrar": {"icon": "📝", "roles": ["admin"]},
        "Base de datos": {"icon": "🗃️", "roles": ["admin"]},
        "Semana": {"icon": "🗓️", "roles": ["admin", "profesor"]},
        "Dashboard": {"icon": "📈", "roles": ["admin"]},
        "Técnicos": {"icon": "🔧", "roles": ["admin"]},
        "Diplomas": {"icon": "🎓", "roles": ["admin", "profesor"]},
        "Inventario": {"icon": "💻", "roles": ["admin"]},
        "Mantención preventiva": {"icon": "🧰", "roles": ["admin"]},
        "Auditoría": {"icon": "🧾", "roles": ["admin"]},
        "Configuración": {"icon": "⚙️", "roles": ["admin"]},
        "Modo TV": {"icon": "📺", "roles": ["admin", "mensajeria"]},
    }

    available_pages = [
        p
        for p, conf in PAGES_CONFIG.items()
        if st.session_state.role in conf["roles"]
    ]
    default_page = "Inicio"

    if default_page not in available_pages:
        default_page = available_pages[0]

    # Navegación diferida: aplicar ANTES de crear el widget del menú.
    pagina_pendiente = st.session_state.pop(
        "_pending_nav_page",
        None,
    )

    if pagina_pendiente in available_pages:
        st.session_state.nav_page = pagina_pendiente

    if (
        "nav_page" not in st.session_state
        or st.session_state.nav_page not in available_pages
    ):
        st.session_state.nav_page = default_page

    page = st.sidebar.radio(
        "Navegación",
        available_pages,
        format_func=lambda p: f"{PAGES_CONFIG[p]['icon']} {p}",
        label_visibility="collapsed",
        key="nav_page",
    )

    st.sidebar.markdown("---")

    st.sidebar.caption(
        "⚡ Modo rendimiento activo · sin refrescos automáticos mientras trabajas"
    )

    if st.sidebar.button("🔄 Refrescar datos", use_container_width=True):
        st.cache_data.clear()
        st.rerun()

    if st.sidebar.button("🚪 Cerrar Sesión", use_container_width=True):
        for key in st.session_state.keys(): 
            del st.session_state[key]
        st.rerun()
# ==============================================================================
# 🤖 ASISTENTE IA FLOTANTE Y CONTEXTUAL V21
# ==============================================================================
def obtener_contexto_pagina_asistente(pagina_actual):
    """
    Entrega a Gemini contexto de la sección que el usuario está visitando.
    No expone secretos ni credenciales.
    """
    rol_actual = st.session_state.get("role", "usuario")
    usuario_actual = (
        st.session_state.get("profesor_name")
        or st.session_state.get("mensajeria_user")
        or "Usuario"
    )

    contexto_base = [
        "SISTEMA: Sistema Institucional CAV / Departamento de Informática y Enlaces.",
        f"PÁGINA ACTUAL: {pagina_actual}.",
        f"ROL ACTUAL: {rol_actual}.",
        f"USUARIO: {usuario_actual}.",
        (
            "Tu prioridad es ayudar sobre la sección visible. "
            "Si la pregunta no tiene relación con la página actual, puedes responder, "
            "pero primero aclara brevemente que estás cambiando de contexto."
        ),
        (
            "Nunca muestres claves API, contraseñas, tokens, service_role, secretos "
            "de Streamlit ni credenciales SMTP."
        ),
    ]

    contextos = {
        "Inicio": """
La página Inicio es el Centro de Operaciones. Muestra métricas rápidas del día,
tickets pendientes/en revisión, bajas registradas, próximas actividades y
Acciones rápidas mediante botones que llevan a Registrar, Tickets, Baja de Equipos
y Modo TV.
""",
        "Registrar": """
La página Registrar permite crear reservas de recursos. El usuario selecciona
profesor, curso, recurso, fecha, hora de inicio, hora de término y observaciones.
El sistema debe prevenir reservas superpuestas del mismo recurso.
""",
        "Mis Reservas": """
La página Mis Reservas permite revisar reservas asociadas al profesor autenticado,
consultar fechas y horarios y gestionar sus registros según los permisos disponibles.
""",
        "Semana": """
La página Semana es la vista semanal de reservas. Sirve para revisar rápidamente
qué recurso está ocupado, por quién, para qué curso y en qué horario.
""",
        "Dashboard": """
Dashboard contiene estadísticas de uso, gráficos, métricas de reservas y resúmenes
para apoyar la gestión del Departamento de Informática/Enlaces.
""",
        "Base de datos": """
Base de datos es el Centro avanzado de registros. Tiene cuatro pestañas:
Consultar, Editar, Eliminar y Exportar.
Consultar permite búsqueda global, filtros por profesor/curso/recurso/fecha y paginación.
Editar funciona seleccionando primero una fecha en calendario; después muestra solo
los registros de ese día y permite elegir exactamente una reserva para modificarla.
Eliminar permite seleccionar registros con confirmación explícita.
Exportar descarga los resultados filtrados en CSV.
""",
        "Técnicos": f"""
Técnicos es el centro de soporte. El submódulo actualmente seleccionado es:
{st.session_state.get("tecnico_modulo", "Tickets")}.
Incluye Tickets, Baja de Equipos y Generador QR. Los tickets administran fallas
reportadas; Baja de Equipos genera documentación técnica; Generador QR crea
identificadores para reportes/equipos.
""",
        "Inventario": """
Inventario registra equipos y recursos tecnológicos: nombre, categoría, código
patrimonial, serie, marca, modelo, ubicación, responsable, garantía, estado y
observaciones. También permite búsqueda y consulta del pasaporte del equipo.
""",
        "Mantención preventiva": """
Mantención preventiva permite programar tareas de mantenimiento para equipos del
inventario, asignar prioridad, responsable, frecuencia, estado y resultado.
""",
        "Auditoría": """
Auditoría permite revisar acciones realizadas dentro del sistema, usuario,
módulo, registro afectado, detalle y fecha.
""",
        "Diplomas": """
Diplomas Digitales CAV genera reconocimientos digitales oficiales. Mantiene registro
en Supabase, código único, enlace digital, PDF de respaldo, envío por correo
institucional y métricas de impacto ambiental. El administrador puede revisar el
dashboard histórico de diplomas.
""",
        "Configuración": """
Configuración reúne opciones administrativas, estado de servicios y controles
generales del sistema. Debe utilizarse con cuidado porque afecta el funcionamiento
institucional.
""",
        "Modo TV": """
Modo TV administra la pantalla informativa pública. Incluye avisos, alertas rojas,
eventos y cronograma. Las alertas y avisos pueden tener fecha/hora de inicio y fin.
La pantalla se adapta automáticamente a la resolución del dispositivo.
""",
    }

    contexto_especifico = contextos.get(
        pagina_actual,
        "Esta es una sección del Sistema Institucional CAV."
    )

    # Añadir un pequeño resumen de datos sin exponer información sensible.
    resumen_datos = ""
    try:
        if pagina_actual in {"Inicio", "Semana", "Base de datos", "Mis Reservas"}:
            total_reservas = len(df) if isinstance(df, pd.DataFrame) else 0
            resumen_datos += f"\nReservas cargadas actualmente en memoria: {total_reservas}."

        if pagina_actual == "Base de datos":
            filtros_activos = []
            if st.session_state.get("bd_busqueda_global"):
                filtros_activos.append(
                    f"texto='{st.session_state.get('bd_busqueda_global')}'"
                )
            if st.session_state.get("bd_filtro_profesor"):
                filtros_activos.append("filtro profesor activo")
            if st.session_state.get("bd_filtro_curso"):
                filtros_activos.append("filtro curso activo")
            if st.session_state.get("bd_filtro_recurso"):
                filtros_activos.append("filtro recurso activo")
            if st.session_state.get("bd_calendario_edicion"):
                filtros_activos.append(
                    "fecha de edición="
                    + str(st.session_state.get("bd_calendario_edicion"))
                )
            if filtros_activos:
                resumen_datos += (
                    "\nEstado visible de filtros: "
                    + ", ".join(filtros_activos)
                    + "."
                )
    except Exception:
        pass

    return "\n".join(contexto_base) + "\n" + contexto_especifico + resumen_datos


def responder_asistente_contextual(pregunta, pagina_actual):
    if "asistente_historial" not in st.session_state:
        st.session_state.asistente_historial = []

    historial_reciente = st.session_state.asistente_historial[-8:]
    historial_texto = "\n".join(
        f"{'Usuario' if item['role'] == 'user' else 'Asistente'}: {item['content']}"
        for item in historial_reciente
    )

    contexto = obtener_contexto_pagina_asistente(pagina_actual)

    prompt = f"""
Eres el Asistente IA del Departamento de Informática/Enlaces del
Liceo Bicentenario de Excelencia Colegio Antonio Varas.

Tu tono debe ser profesional, cercano, claro y resolutivo.
Debes orientar al usuario sobre la página que está viendo y explicar
acciones concretas dentro del sistema.

{contexto}

HISTORIAL RECIENTE:
{historial_texto if historial_texto else "Sin conversación previa."}

PREGUNTA DEL USUARIO:
{pregunta}

INSTRUCCIONES DE RESPUESTA:
- Responde en español.
- Prioriza instrucciones relacionadas con la página actual.
- Cuando corresponda, indica la ruta exacta dentro del sistema.
- Sé breve cuando la consulta sea simple.
- Para procedimientos, usa pasos claros.
- No inventes funciones que no estén descritas en el contexto.
- No reveles secretos ni credenciales.
"""

    respuesta = consultar_gemini(prompt)

    st.session_state.asistente_historial.append(
        {"role": "user", "content": pregunta}
    )
    st.session_state.asistente_historial.append(
        {"role": "assistant", "content": respuesta}
    )

    # Limitar memoria visible para no hacer crecer indefinidamente la sesión.
    if len(st.session_state.asistente_historial) > 30:
        st.session_state.asistente_historial = (
            st.session_state.asistente_historial[-30:]
        )

    return respuesta


@st.fragment
def renderizar_asistente_flotante(pagina_actual):
    """
    Asistente flotante V21.1.

    La versión anterior dependía del selector interno data-testid="stPopover".
    En versiones nuevas de Streamlit ese selector puede cambiar.

    Esta versión asigna una KEY explícita al popover y usa la clase CSS
    oficial generada por Streamlit: st-key-asistente_flotante_cav.
    """

    st.markdown(
        """
        <style>
        /* ==========================================================
           GLOBITO FLOTANTE DEL ASISTENTE
           ========================================================== */

        .st-key-asistente_flotante_cav,
        div.st-key-asistente_flotante_cav,
        [class*="st-key-asistente_flotante_cav"] {
            position: fixed !important;
            left: max(16px, env(safe-area-inset-left)) !important;
            bottom: max(18px, env(safe-area-inset-bottom)) !important;
            right: auto !important;
            top: auto !important;
            width: 64px !important;
            min-width: 64px !important;
            max-width: 64px !important;
            height: 64px !important;
            min-height: 64px !important;
            max-height: 64px !important;
            z-index: 2147483000 !important;
            margin: 0 !important;
            padding: 0 !important;
            overflow: visible !important;
        }

        .st-key-asistente_flotante_cav button,
        [class*="st-key-asistente_flotante_cav"] button {
            display: flex !important;
            align-items: center !important;
            justify-content: center !important;
            gap: 0 !important;

            width: 64px !important;
            min-width: 64px !important;
            max-width: 64px !important;
            height: 64px !important;
            min-height: 64px !important;
            max-height: 64px !important;

            padding: 0 !important;
            margin: 0 !important;

            color: #ffffff !important;
            background:
                linear-gradient(
                    145deg,
                    #800020 0%,
                    #5b0017 100%
                ) !important;

            border: 2px solid #C7A44A !important;
            border-radius: 999px !important;

            box-shadow:
                0 14px 30px rgba(15, 23, 42, .34),
                0 0 0 6px rgba(128, 0, 32, .12) !important;

            cursor: pointer !important;
            transition:
                transform .18s ease,
                box-shadow .18s ease !important;
        }

        .st-key-asistente_flotante_cav button:hover,
        [class*="st-key-asistente_flotante_cav"] button:hover {
            transform: translateY(-3px) scale(1.05) !important;
            box-shadow:
                0 18px 38px rgba(15, 23, 42, .38),
                0 0 0 8px rgba(199, 164, 74, .16) !important;
        }

        /* Ocultar el texto y dejar solo el robot */
        .st-key-asistente_flotante_cav button p,
        [class*="st-key-asistente_flotante_cav"] button p {
            display: none !important;
        }

        .st-key-asistente_flotante_cav button svg,
        [class*="st-key-asistente_flotante_cav"] button svg {
            width: 29px !important;
            height: 29px !important;
        }

        /* Pulso visual discreto para hacerlo fácil de encontrar */
        .st-key-asistente_flotante_cav::after,
        [class*="st-key-asistente_flotante_cav"]::after {
            content: "";
            position: absolute;
            inset: -7px;
            border: 2px solid rgba(199, 164, 74, .33);
            border-radius: 999px;
            pointer-events: none;
            animation: cavAssistantPulse 2.6s ease-out infinite;
        }

        @keyframes cavAssistantPulse {
            0% {
                opacity: .65;
                transform: scale(.92);
            }
            72% {
                opacity: 0;
                transform: scale(1.25);
            }
            100% {
                opacity: 0;
                transform: scale(1.25);
            }
        }

        /* El panel emergente */
        [data-baseweb="popover"] {
            z-index: 2147483001 !important;
        }

        [data-baseweb="popover"] > div {
            max-width: min(440px, calc(100vw - 24px)) !important;
        }

        /* En teléfono: globito algo más pequeño */
        @media (max-width: 700px) {
            .st-key-asistente_flotante_cav,
            div.st-key-asistente_flotante_cav,
            [class*="st-key-asistente_flotante_cav"] {
                left: 10px !important;
                bottom: 12px !important;
                width: 58px !important;
                min-width: 58px !important;
                max-width: 58px !important;
                height: 58px !important;
                min-height: 58px !important;
                max-height: 58px !important;
            }

            .st-key-asistente_flotante_cav button,
            [class*="st-key-asistente_flotante_cav"] button {
                width: 58px !important;
                min-width: 58px !important;
                max-width: 58px !important;
                height: 58px !important;
                min-height: 58px !important;
                max-height: 58px !important;
            }
        }

        @media (prefers-reduced-motion: reduce) {
            .st-key-asistente_flotante_cav::after,
            [class*="st-key-asistente_flotante_cav"]::after {
                animation: none !important;
            }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # KEY explícita: Streamlit genera la clase CSS
    # "st-key-asistente_flotante_cav", que usamos arriba.
    asistente_popover = st.popover(
        "Asistente",
        icon="🤖",
        help=(
            f"Asistente IA contextual · Página actual: {pagina_actual}"
        ),
        type="primary",
        width="content",
        key="asistente_flotante_cav",
    )

    with asistente_popover:
        st.markdown("### 🤖 Asistente CAV")
        st.caption(
            f"Contexto actual: **{pagina_actual}**"
        )

        st.info(
            "Pregúntame cómo usar esta sección, dónde encontrar una opción "
            "o qué significa la información que estás viendo."
        )

        sugerencias = st.columns(3)
        pregunta_sugerida = None

        if sugerencias[0].button(
            "🧭 ¿Qué hago aquí?",
            key=f"asistente_que_hago_{pagina_actual}",
            use_container_width=True,
        ):
            pregunta_sugerida = (
                "Explícame qué puedo hacer en esta página y cuál es "
                "el flujo recomendado."
            )

        if sugerencias[1].button(
            "📋 Guíame",
            key=f"asistente_guiame_{pagina_actual}",
            use_container_width=True,
        ):
            pregunta_sugerida = (
                "Guíame paso a paso para usar correctamente esta sección."
            )

        if sugerencias[2].button(
            "⚠️ Ayuda",
            key=f"asistente_ayuda_{pagina_actual}",
            use_container_width=True,
        ):
            pregunta_sugerida = (
                "¿Qué errores comunes debo evitar en esta sección?"
            )

        historial = st.session_state.get(
            "asistente_historial",
            [],
        )

        contenedor_historial = st.container(
            height=320,
            border=True,
        )

        with contenedor_historial:
            if not historial:
                st.markdown(
                    "👋 **Hola.** Puedes preguntarme algo sobre "
                    f"**{pagina_actual}**."
                )
            else:
                for item in historial[-12:]:
                    with st.chat_message(item["role"]):
                        st.markdown(item["content"])

        with st.form(
            f"form_asistente_flotante_{pagina_actual}",
            clear_on_submit=True,
        ):
            pregunta_manual = st.text_input(
                "Tu pregunta",
                placeholder=(
                    f"Ej. ¿Cómo uso {pagina_actual}?"
                ),
                label_visibility="collapsed",
            )

            enviar = st.form_submit_button(
                "Enviar ➜",
                type="primary",
                use_container_width=True,
            )

        pregunta_final = pregunta_sugerida

        if enviar and pregunta_manual.strip():
            pregunta_final = pregunta_manual.strip()

        if pregunta_final:
            with st.spinner("Analizando esta sección..."):
                responder_asistente_contextual(
                    pregunta_final,
                    pagina_actual,
                )

            st.rerun(scope="fragment")

        if st.button(
            "🧹 Limpiar conversación",
            key=f"limpiar_asistente_{pagina_actual}",
            use_container_width=True,
        ):
            st.session_state.asistente_historial = []
            st.rerun(scope="fragment")

        st.caption(
            "🔒 El asistente usa el contexto de la sección actual y no "
            "tiene acceso a tus contraseñas ni claves secretas."
        )


renderizar_asistente_flotante(page)


# ------------------------------------------------------------------
# PÁGINAS
# ------------------------------------------------------------------

if page == "Inicio":
    st.title("🏠 Centro de Operaciones")
    st.caption("⚡ V23 · rendimiento optimizado · asistente contextual activo")
    nombre_corto = (st.session_state.get("profesor_name") or "Usuario").split(" ")[0]
    st.caption(f"Bienvenido, {nombre_corto}. Resumen actualizado del sistema.")

    hoy_inicio = dt.date.today()
    reservas_hoy_inicio = df[df["Fecha"] == hoy_inicio] if not df.empty else pd.DataFrame()
    tickets_inicio, bajas_inicio = cargar_resumen_inicio()

    tickets_pend = [
        t
        for t in tickets_inicio
        if t.get("estado") == "Reportado (Vía QR)"
    ]
    tickets_rev = [
        t
        for t in tickets_inicio
        if t.get("estado") == "En Revisión"
    ]

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("📅 Reservas hoy", len(reservas_hoy_inicio))
    c2.metric("🔴 Tickets pendientes", len(tickets_pend))
    c3.metric("🟡 En revisión", len(tickets_rev))
    c4.metric("🗑️ Bajas registradas", len(bajas_inicio))

    st.markdown("### ⚡ Acciones rápidas")

    def ir_accion_rapida(
        pagina,
        modulo_tecnico=None,
        abrir_config_tv=False,
    ):
        if pagina not in available_pages:
            st.warning(
                "Tu perfil no tiene permisos para abrir esta sección."
            )
            return

        # No modificamos directamente `nav_page` porque el radio ya existe.
        st.session_state["_pending_nav_page"] = pagina

        if modulo_tecnico:
            st.session_state["_pending_tecnico_modulo"] = modulo_tecnico

        if abrir_config_tv:
            st.session_state.ver_pantalla_tv = False

        st.rerun()

    if st.session_state.get("role") == "admin":
        acciones = st.columns(4)

        with acciones[0]:
            with st.container(border=True):
                st.markdown("#### 📝 Nueva reserva")
                st.caption(
                    "Registrar una nueva reserva de sala, laboratorio "
                    "o recurso."
                )
                if st.button(
                    "Crear reserva",
                    key="qa_nueva_reserva",
                    type="primary",
                    use_container_width=True,
                ):
                    ir_accion_rapida("Registrar")

        with acciones[1]:
            with st.container(border=True):
                st.markdown("#### 🎫 Gestionar tickets")
                st.caption(
                    "Revisar fallas reportadas y actualizar su estado."
                )
                if st.button(
                    "Abrir tickets",
                    key="qa_tickets",
                    use_container_width=True,
                ):
                    ir_accion_rapida(
                        "Técnicos",
                        modulo_tecnico="🎫 Tickets",
                    )

        with acciones[2]:
            with st.container(border=True):
                st.markdown("#### 🗑️ Dar de baja")
                st.caption(
                    "Abrir directamente el módulo de Baja de Equipos."
                )
                if st.button(
                    "Procesar baja",
                    key="qa_baja",
                    use_container_width=True,
                ):
                    ir_accion_rapida(
                        "Técnicos",
                        modulo_tecnico="🗑️ Baja de Equipos",
                    )

        with acciones[3]:
            with st.container(border=True):
                st.markdown("#### 📺 Publicar aviso")
                st.caption(
                    "Abrir la administración de Modo TV para avisos "
                    "y alertas."
                )
                if st.button(
                    "Abrir Modo TV",
                    key="qa_modo_tv",
                    use_container_width=True,
                ):
                    ir_accion_rapida(
                        "Modo TV",
                        abrir_config_tv=True,
                    )

    elif st.session_state.get("role") == "profesor":
        acciones = st.columns(3)

        with acciones[0]:
            with st.container(border=True):
                st.markdown("#### 👤 Mis reservas")
                if st.button(
                    "Ver mis reservas",
                    key="qa_prof_reservas",
                    type="primary",
                    use_container_width=True,
                ):
                    ir_accion_rapida("Mis Reservas")

        with acciones[1]:
            with st.container(border=True):
                st.markdown("#### 🗓️ Semana")
                if st.button(
                    "Ver horario semanal",
                    key="qa_prof_semana",
                    use_container_width=True,
                ):
                    ir_accion_rapida("Semana")

        with acciones[2]:
            with st.container(border=True):
                st.markdown("#### 🎓 Diplomas")
                if st.button(
                    "Abrir diplomas",
                    key="qa_prof_diplomas",
                    use_container_width=True,
                ):
                    ir_accion_rapida("Diplomas")

    else:
        acciones = st.columns(2)

        with acciones[0]:
            with st.container(border=True):
                st.markdown("#### 📺 Modo TV")
                if st.button(
                    "Administrar pantalla",
                    key="qa_msg_tv",
                    type="primary",
                    use_container_width=True,
                ):
                    ir_accion_rapida(
                        "Modo TV",
                        abrir_config_tv=True,
                    )

        with acciones[1]:
            with st.container(border=True):
                st.markdown("#### 🏠 Inicio")
                if st.button(
                    "Volver al inicio",
                    key="qa_msg_inicio",
                    use_container_width=True,
                ):
                    ir_accion_rapida("Inicio")

    st.markdown("### 🕒 Próximas actividades")
    if reservas_hoy_inicio.empty:
        st.info("No hay reservas para hoy.")
    else:
        for _, row in reservas_hoy_inicio.sort_values("Hora inicio").head(6).iterrows():
            with st.container(border=True):
                st.markdown(
                    f"**{row['Hora inicio'].strftime('%H:%M')}–{row['Hora fin'].strftime('%H:%M')} · {row['Recurso']}**  \n"
                    f"👨‍🏫 {row['Profesor']} · 📚 {row['Curso']}"
                )

elif page == "Mis Reservas":
    st.title("👤 Mis Próximas Reservas")
    if not df.empty:
        prof_df = df[df['Profesor'] == st.session_state.profesor_name]
        future_reservas = prof_df[prof_df['Fecha'] >= dt.date.today()].sort_values(by="Fecha")
    else: future_reservas = pd.DataFrame()

    if future_reservas.empty: st.info("No tienes reservas programadas para el futuro.")
    else:
        for _, row in future_reservas.iterrows():
            with st.container(border=True):
                st.markdown(f"#### {format_date_es(row['Fecha'])}")
                st.markdown(f"**Hora:** {row['Hora inicio'].strftime('%H:%M')} - {row['Hora fin'].strftime('%H:%M')}<br>**Curso:** {row['Curso']}<br>**Recurso:** {row['Recurso']}", unsafe_allow_html=True)
                if row['Observaciones']: st.markdown(f"> *{row['Observaciones']}*")

if page == "Registrar":
    st.title("📝 Registrar Nuevo Horario")
    def check_all_conflicts(fechas, recursos, hora_inicio, hora_fin, df_reservas, df_mantenimiento):
        conflictos_reserva, conflictos_mantenimiento = [], []
        if not df_reservas.empty:
            for fecha in fechas:
                for rec in recursos:
                    df_check = df_reservas[(df_reservas['Fecha'] == fecha) & (df_reservas['Recurso'] == rec)]
                    for _, row in df_check.iterrows():
                        if overlap(hora_inicio, hora_fin, row['Hora inicio'], row['Hora fin']):
                            conflictos_reserva.append(f"<li>{rec} el {fecha.strftime('%d/%m/%Y')} (con {row['Profesor']})</li>")
        if not df_mantenimiento.empty:
            for fecha in fechas:
                for rec in recursos:
                    mant_check = df_mantenimiento[(df_mantenimiento['Recurso'] == rec) & (df_mantenimiento['FechaInicio_dt'] <= fecha) & (df_mantenimiento['FechaFin_dt'] >= fecha)]
                    for _, m_row in mant_check.iterrows():
                        if overlap(hora_inicio, hora_fin, as_time(m_row['HoraInicio']), as_time(m_row['HoraFin'])):
                             conflictos_mantenimiento.append(f"<li>{rec} el {fecha.strftime('%d/%m/%Y')} (en mantenimiento)</li>")
        return conflictos_reserva, list(set(conflictos_mantenimiento))
        
    with st.container(border=True):
        tipo_reserva = st.radio("Tipo de Reserva", ["Única", "Múltiples Fechas", "Semanal Recurrente"], horizontal=True, key="tipo_reserva")
        with st.form("entry_form"):
            st.markdown("---")
            c1, c2 = st.columns(2)
            fechas_a_registrar = []
            
            # 1. Definir la hora exacta en Chile
            import pytz
            tz_chile = pytz.timezone('America/Santiago')
            hoy_chile = dt_datetime.now(tz_chile).date()

            if tipo_reserva == "Única":
                # Forzamos a que el valor por defecto sea hoy_chile
                fecha = c1.date_input('Fecha', value=hoy_chile, format="DD/MM/YYYY")
                fechas_a_registrar.append(fecha)
                
            elif tipo_reserva == "Múltiples Fechas":
                # Usamos hoy_chile como punto de partida
                date_range = [hoy_chile + dt.timedelta(days=i) for i in range(180)]
                df_dates = pd.DataFrame({"Seleccionar": [False] * len(date_range), "Fecha Disponible": [format_date_es(d) for d in date_range], "_date_obj": date_range})
                with c1:
                    st.write("Selecciona las fechas deseadas:")
                    edited_dates_df = st.data_editor(df_dates, column_config={"Seleccionar": st.column_config.CheckboxColumn(required=True), "_date_obj": None}, hide_index=True, height=200, use_container_width=True)
                selected_dates_df = edited_dates_df[edited_dates_df["Seleccionar"]]
                fechas_a_registrar = sorted(selected_dates_df["_date_obj"].tolist())
                
            else:
                # Forzamos el inicio de la recurrencia a hoy_chile
                fecha_inicio = c1.date_input('Fecha de Inicio', value=hoy_chile, format="DD/MM/YYYY")
                num_semanas = c1.number_input('Repetir durante (semanas)', min_value=1, max_value=52, value=4)
                fechas_a_registrar = [fecha_inicio + dt.timedelta(weeks=i) for i in range(num_semanas)]
                if fechas_a_registrar: c1.info(f"Se registrarán {len(fechas_a_registrar)} fechas.")
            
            hora  = c1.selectbox('Bloque Horario', HORAS)
            obs   = c1.text_area('Observaciones (Opcional)')
            prof   = c2.selectbox('Profesor', PROFESORES)
            curso  = c2.selectbox('Curso',    CURSOS)
            recs = c2.multiselect('Recursos', RECURSOS, placeholder="Selecciona uno o más recursos")
            
            if st.form_submit_button('💾 Guardar Registro', use_container_width=True, type="primary"):
                if recs and fechas_a_registrar:
                    h_inicio, h_fin = [dt.datetime.strptime(t.strip(), '%H:%M').time() for t in hora.split(' a ')]
                    
                    # Validación 1: Revisar contra mantenimientos y caché local
                    conflictos_r, conflictos_m = check_all_conflicts(fechas_a_registrar, recs, h_inicio, h_fin, df, df_mantenimiento)
                    
                    # Validación 2: CONSULTA ESTRICTA EN TIEMPO REAL A SUPABASE
                    # Esto evita choques si dos profes guardan al mismo tiempo
                    hay_choque_tiempo_real = False
                    
                    for fecha in fechas_a_registrar:
                        for rec in recs:
                            id_rec_buscado = map_rec.get(rec)
                            # Pedimos directamente a la base de datos las reservas de ese día y recurso
                            db_check = supabase.table("reservas").select("hora_inicio, hora_fin, profesores(nombre)").eq("fecha", str(fecha)).eq("recurso", id_rec_buscado).execute()
                            
                            for registro_db in (db_check.data or []):
                                h_ini_db = as_time(registro_db['hora_inicio'])
                                h_fin_db = as_time(registro_db['hora_fin'])
                                
                                # Si hay solapamiento exacto
                                if overlap(h_inicio, h_fin, h_ini_db, h_fin_db):
                                    profe_ocupante = registro_db['profesores']['nombre'] if registro_db.get('profesores') else 'Otro usuario'
                                    mensaje_choque = f"<li>{rec} el {fecha.strftime('%d/%m/%Y')} (Ya reservado por {profe_ocupante} justo ahora)</li>"
                                    if mensaje_choque not in conflictos_r:
                                        conflictos_r.append(mensaje_choque)
                                    hay_choque_tiempo_real = True

                    # Comprobamos si saltó alguna de las dos alarmas
                    if conflictos_r or conflictos_m or hay_choque_tiempo_real:
                        if conflictos_r: 
                            st.error(f"❌ **Error: Alguien acaba de ocupar este horario:**")
                            st.markdown(f"<ul>{''.join(set(conflictos_r))}</ul>", unsafe_allow_html=True)
                        if conflictos_m: 
                            st.error(f"❌ **Error: Equipo en Mantenimiento:**")
                            st.markdown(f"<ul>{''.join(set(conflictos_m))}</ul>", unsafe_allow_html=True)
                    else:
                        # Si todo está libre, procedemos a guardar
                        nuevas_reservas = []
                        for fecha in fechas_a_registrar:
                            for rec in recs:
                                nuevas_reservas.append({
                                    'fecha': str(fecha), 
                                    'hora_inicio': str(h_inicio), 
                                    'hora_fin': str(h_fin), 
                                    'profesor': map_prof.get(prof), 
                                    'curso': map_cur.get(curso), 
                                    'recurso': map_rec.get(rec), 
                                    'observaciones': obs
                                })
                        try:
                            supabase.table("reservas").insert(nuevas_reservas).execute()
                            st.success("✅ ¡Reservas guardadas exitosamente!")
                            st.cache_data.clear() # Limpiamos la caché global
                            
                            email_to = PROFESOR_DATA.get(prof)
                            if email_to:
                                subject = f"Confirmación de Reserva de Recursos - {curso}"
                                body = f"""<html><body><p>Hola {prof.split(' ')[0]},</p><p>Se ha(n) confirmado la(s) siguiente(s) reserva(s) a tu nombre:</p><ul><li><b>Curso:</b> {curso}</li><li><b>Recurso(s):</b> {', '.join(recs)}</li><li><b>Horario:</b> {hora}</li></ul><p><b>Fechas Registradas:</b></p><ul>{''.join([f'<li>{format_date_es(f)}</li>' for f in fechas_a_registrar])}</ul>{f"<p><b>Observaciones:</b> {obs}</p>" if obs else ""}<p>Saludos,<br>Sistema de Horarios CAV</p></body></html>"""
                                send_email(subject, body, email_to)
                                
                            time.sleep(0.25)
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error al guardar en la nube: {e}")
if page == "Base de datos":
    st.title("🗃️ Centro avanzado de registros")
    st.caption(
        "Busca, consulta, edita, elimina y exporta reservas de forma segura. "
        "Los filtros se aplican sin modificar la información original."
    )

    def bd_fecha(valor):
        try:
            return pd.to_datetime(valor).date()
        except Exception:
            return None

    def bd_hora(valor):
        if isinstance(valor, dt.time):
            return valor
        try:
            return pd.to_datetime(str(valor)).time()
        except Exception:
            try:
                return dt.datetime.strptime(
                    str(valor)[:5],
                    "%H:%M",
                ).time()
            except Exception:
                return dt.time(0, 0)

    def bd_hora_texto(valor):
        hora = bd_hora(valor)
        return hora.strftime("%H:%M")

    def bd_etiqueta_registro(fila):
        return (
            f"ID {int(fila['id'])} · "
            f"{bd_fecha(fila['Fecha']).strftime('%d/%m/%Y') if bd_fecha(fila['Fecha']) else 'Sin fecha'} · "
            f"{bd_hora_texto(fila['Hora inicio'])} · "
            f"{fila.get('Recurso', 'Sin recurso')} · "
            f"{fila.get('Profesor', 'Sin profesor')}"
        )

    if df.empty:
        st.info("No hay reservas registradas en la base de datos.")
    else:
        datos_bd = df.copy()

        # Columnas normalizadas solo para filtros y orden.
        datos_bd["_fecha_filtro"] = pd.to_datetime(
            datos_bd["Fecha"],
            errors="coerce",
        ).dt.date
        datos_bd["_hora_inicio_filtro"] = datos_bd["Hora inicio"].apply(
            bd_hora
        )
        datos_bd["_texto_busqueda"] = (
            datos_bd[
                [
                    "Profesor",
                    "Curso",
                    "Recurso",
                    "Observaciones",
                ]
            ]
            .fillna("")
            .astype(str)
            .agg(" ".join, axis=1)
            .str.lower()
        )

        st.markdown("### 🔎 Búsqueda avanzada")

        with st.container(border=True):
            fila_1 = st.columns([1.5, 1, 1, 1])

            texto_busqueda = fila_1[0].text_input(
                "Buscar en todos los campos",
                placeholder=(
                    "Profesor, curso, recurso, observación o ID"
                ),
                key="bd_busqueda_global",
            )

            profesores_filtro = fila_1[1].multiselect(
                "Profesor",
                sorted(
                    datos_bd["Profesor"]
                    .dropna()
                    .astype(str)
                    .unique()
                    .tolist()
                ),
                placeholder="Todos",
                key="bd_filtro_profesor",
            )

            cursos_filtro = fila_1[2].multiselect(
                "Curso",
                sorted(
                    datos_bd["Curso"]
                    .dropna()
                    .astype(str)
                    .unique()
                    .tolist()
                ),
                placeholder="Todos",
                key="bd_filtro_curso",
            )

            recursos_filtro = fila_1[3].multiselect(
                "Recurso",
                sorted(
                    datos_bd["Recurso"]
                    .dropna()
                    .astype(str)
                    .unique()
                    .tolist()
                ),
                placeholder="Todos",
                key="bd_filtro_recurso",
            )

            fila_2 = st.columns([1, 1, 1, 1])

            fecha_minima = datos_bd["_fecha_filtro"].dropna().min()
            fecha_maxima = datos_bd["_fecha_filtro"].dropna().max()

            usar_todas_fechas = fila_2[0].checkbox(
                "Todas las fechas",
                value=True,
                key="bd_todas_fechas",
            )

            fecha_desde = fila_2[1].date_input(
                "Desde",
                value=fecha_minima or dt.date.today(),
                format="DD/MM/YYYY",
                disabled=usar_todas_fechas,
                key="bd_fecha_desde",
            )

            fecha_hasta = fila_2[2].date_input(
                "Hasta",
                value=fecha_maxima or dt.date.today(),
                format="DD/MM/YYYY",
                disabled=usar_todas_fechas,
                key="bd_fecha_hasta",
            )

            solo_futuras = fila_2[3].checkbox(
                "Solo reservas vigentes o futuras",
                value=False,
                key="bd_solo_futuras",
            )

            fila_3 = st.columns([1.3, 1, 1, 1])

            campo_orden = fila_3[0].selectbox(
                "Ordenar por",
                [
                    "Fecha",
                    "Hora inicio",
                    "Profesor",
                    "Curso",
                    "Recurso",
                    "id",
                ],
                key="bd_orden_campo",
            )

            orden_descendente = fila_3[1].toggle(
                "Orden descendente",
                value=False,
                key="bd_orden_desc",
            )

            resultados_pagina = fila_3[2].selectbox(
                "Resultados por página",
                [10, 20, 50, 100],
                index=1,
                key="bd_resultados_pagina",
            )

            mostrar_observaciones = fila_3[3].toggle(
                "Mostrar observaciones",
                value=True,
                key="bd_mostrar_observaciones",
            )

        filtrados = datos_bd.copy()

        if texto_busqueda.strip():
            texto_normalizado = texto_busqueda.strip().lower()
            mascara_texto = filtrados["_texto_busqueda"].str.contains(
                texto_normalizado,
                regex=False,
                na=False,
            )
            if texto_normalizado.isdigit():
                mascara_texto = (
                    mascara_texto
                    | filtrados["id"]
                    .astype(str)
                    .str.contains(
                        texto_normalizado,
                        regex=False,
                        na=False,
                    )
                )
            filtrados = filtrados[mascara_texto]

        if profesores_filtro:
            filtrados = filtrados[
                filtrados["Profesor"].isin(profesores_filtro)
            ]

        if cursos_filtro:
            filtrados = filtrados[
                filtrados["Curso"].isin(cursos_filtro)
            ]

        if recursos_filtro:
            filtrados = filtrados[
                filtrados["Recurso"].isin(recursos_filtro)
            ]

        if not usar_todas_fechas:
            filtrados = filtrados[
                filtrados["_fecha_filtro"].between(
                    fecha_desde,
                    fecha_hasta,
                    inclusive="both",
                )
            ]

        if solo_futuras:
            filtrados = filtrados[
                filtrados["_fecha_filtro"] >= dt.date.today()
            ]

        columna_orden_real = {
            "Fecha": "_fecha_filtro",
            "Hora inicio": "_hora_inicio_filtro",
        }.get(campo_orden, campo_orden)

        filtrados = filtrados.sort_values(
            columna_orden_real,
            ascending=not orden_descendente,
            na_position="last",
        )

        total_encontrados = len(filtrados)
        hoy_bd = dt.date.today()
        reservas_hoy = int(
            (filtrados["_fecha_filtro"] == hoy_bd).sum()
        )
        reservas_futuras = int(
            (filtrados["_fecha_filtro"] >= hoy_bd).sum()
        )
        recursos_encontrados = int(
            filtrados["Recurso"].dropna().nunique()
        )

        metricas = st.columns(4)
        metricas[0].metric("Resultados", total_encontrados)
        metricas[1].metric("Reservas de hoy", reservas_hoy)
        metricas[2].metric("Vigentes o futuras", reservas_futuras)
        metricas[3].metric(
            "Recursos distintos",
            recursos_encontrados,
        )

        pestaña_consultar, pestaña_editar, pestaña_eliminar, pestaña_exportar = st.tabs(
            [
                "🔍 Consultar",
                "✏️ Editar",
                "🗑️ Eliminar",
                "⬇️ Exportar",
            ]
        )

        with pestaña_consultar:
            if filtrados.empty:
                st.warning(
                    "No se encontraron registros con los filtros actuales."
                )
            else:
                total_paginas = max(
                    1,
                    (
                        total_encontrados
                        + resultados_pagina
                        - 1
                    )
                    // resultados_pagina,
                )

                if (
                    st.session_state.get("bd_pagina_actual", 1)
                    > total_paginas
                ):
                    st.session_state.bd_pagina_actual = 1

                navegacion = st.columns([1, 1, 1.5, 1, 1])

                if navegacion[0].button(
                    "⏮️ Primera",
                    use_container_width=True,
                    disabled=(
                        st.session_state.get(
                            "bd_pagina_actual",
                            1,
                        )
                        <= 1
                    ),
                ):
                    st.session_state.bd_pagina_actual = 1
                    st.rerun()

                if navegacion[1].button(
                    "◀️ Anterior",
                    use_container_width=True,
                    disabled=(
                        st.session_state.get(
                            "bd_pagina_actual",
                            1,
                        )
                        <= 1
                    ),
                ):
                    st.session_state.bd_pagina_actual = max(
                        1,
                        st.session_state.get(
                            "bd_pagina_actual",
                            1,
                        )
                        - 1,
                    )
                    st.rerun()

                pagina_actual = navegacion[2].number_input(
                    "Página",
                    min_value=1,
                    max_value=total_paginas,
                    value=min(
                        st.session_state.get(
                            "bd_pagina_actual",
                            1,
                        ),
                        total_paginas,
                    ),
                    step=1,
                    key="bd_selector_pagina",
                )
                st.session_state.bd_pagina_actual = int(
                    pagina_actual
                )

                if navegacion[3].button(
                    "Siguiente ▶️",
                    use_container_width=True,
                    disabled=(
                        st.session_state.bd_pagina_actual
                        >= total_paginas
                    ),
                ):
                    st.session_state.bd_pagina_actual = min(
                        total_paginas,
                        st.session_state.bd_pagina_actual + 1,
                    )
                    st.rerun()

                if navegacion[4].button(
                    "Última ⏭️",
                    use_container_width=True,
                    disabled=(
                        st.session_state.bd_pagina_actual
                        >= total_paginas
                    ),
                ):
                    st.session_state.bd_pagina_actual = total_paginas
                    st.rerun()

                inicio = (
                    st.session_state.bd_pagina_actual - 1
                ) * resultados_pagina
                fin = inicio + resultados_pagina

                pagina_df = filtrados.iloc[inicio:fin].copy()
                pagina_df.insert(0, "Seleccionar", False)

                columnas_visibles = [
                    "Seleccionar",
                    "id",
                    "Fecha",
                    "Hora inicio",
                    "Hora fin",
                    "Profesor",
                    "Curso",
                    "Recurso",
                ]
                if mostrar_observaciones:
                    columnas_visibles.append("Observaciones")

                editor_consulta = st.data_editor(
                    pagina_df[columnas_visibles],
                    hide_index=True,
                    use_container_width=True,
                    disabled=[
                        columna
                        for columna in columnas_visibles
                        if columna != "Seleccionar"
                    ],
                    column_config={
                        "Seleccionar": st.column_config.CheckboxColumn(
                            "Seleccionar",
                            help=(
                                "Marca uno o más registros para "
                                "consultarlos."
                            ),
                        ),
                        "id": st.column_config.NumberColumn(
                            "ID",
                            format="%d",
                        ),
                        "Fecha": st.column_config.DateColumn(
                            "Fecha",
                            format="DD/MM/YYYY",
                        ),
                        "Hora inicio": st.column_config.TimeColumn(
                            "Inicio",
                            format="HH:mm",
                        ),
                        "Hora fin": st.column_config.TimeColumn(
                            "Fin",
                            format="HH:mm",
                        ),
                    },
                    key=(
                        f"bd_consulta_"
                        f"{st.session_state.bd_pagina_actual}_"
                        f"{total_encontrados}"
                    ),
                )

                ids_seleccionados = (
                    editor_consulta.loc[
                        editor_consulta["Seleccionar"],
                        "id",
                    ]
                    .astype(int)
                    .tolist()
                )

                st.caption(
                    f"Mostrando {inicio + 1}–"
                    f"{min(fin, total_encontrados)} de "
                    f"{total_encontrados} registros."
                )

                if ids_seleccionados:
                    st.markdown("#### 📄 Detalle seleccionado")
                    detalle_df = datos_bd[
                        datos_bd["id"].isin(ids_seleccionados)
                    ].drop(
                        columns=[
                            "_fecha_filtro",
                            "_hora_inicio_filtro",
                            "_texto_busqueda",
                        ],
                        errors="ignore",
                    )

                    st.dataframe(
                        detalle_df,
                        use_container_width=True,
                        hide_index=True,
                    )

                    if len(ids_seleccionados) == 1:
                        fila_detalle = detalle_df.iloc[0]
                        with st.container(border=True):
                            d1, d2, d3 = st.columns(3)
                            d1.write(
                                f"**Profesor:** "
                                f"{fila_detalle.get('Profesor', '—')}"
                            )
                            d2.write(
                                f"**Curso:** "
                                f"{fila_detalle.get('Curso', '—')}"
                            )
                            d3.write(
                                f"**Recurso:** "
                                f"{fila_detalle.get('Recurso', '—')}"
                            )

                            st.write(
                                f"**Fecha y horario:** "
                                f"{bd_fecha(fila_detalle.get('Fecha')).strftime('%d/%m/%Y') if bd_fecha(fila_detalle.get('Fecha')) else '—'} "
                                f"· {bd_hora_texto(fila_detalle.get('Hora inicio'))}"
                                f"–{bd_hora_texto(fila_detalle.get('Hora fin'))}"
                            )
                            st.write(
                                f"**Observaciones:** "
                                f"{fila_detalle.get('Observaciones') or 'Sin observaciones'}"
                            )

        with pestaña_editar:
            st.markdown("#### 📅 Selecciona el día que deseas editar")
            st.caption(
                "Selecciona primero una fecha. El sistema mostrará "
                "solamente las reservas registradas para ese día."
            )

            fechas_disponibles = sorted(
                fecha
                for fecha in datos_bd["_fecha_filtro"].dropna().unique()
            )

            fecha_inicial_edicion = dt.date.today()
            if (
                fecha_inicial_edicion not in fechas_disponibles
                and fechas_disponibles
            ):
                fechas_futuras = [
                    fecha
                    for fecha in fechas_disponibles
                    if fecha >= dt.date.today()
                ]
                fecha_inicial_edicion = (
                    fechas_futuras[0]
                    if fechas_futuras
                    else fechas_disponibles[-1]
                )

            with st.container(border=True):
                c_cal1, c_cal2 = st.columns([1, 2])

                fecha_registros_editar = c_cal1.date_input(
                    "Fecha de la reserva",
                    value=fecha_inicial_edicion,
                    format="DD/MM/YYYY",
                    key="bd_calendario_edicion",
                )

                registros_dia = datos_bd[
                    datos_bd["_fecha_filtro"] == fecha_registros_editar
                ].copy()

                registros_dia = registros_dia.sort_values(
                    ["_hora_inicio_filtro", "Recurso", "Profesor"],
                    ascending=[True, True, True],
                    na_position="last",
                )

                c_cal2.metric(
                    "Registros encontrados ese día",
                    len(registros_dia),
                )

            if registros_dia.empty:
                st.info(
                    "No hay reservas registradas para "
                    f"{fecha_registros_editar.strftime('%d/%m/%Y')}."
                )

                if fechas_disponibles:
                    with st.expander(
                        "📌 Ver fechas que sí tienen registros",
                        expanded=False,
                    ):
                        fechas_texto = ", ".join(
                            fecha.strftime("%d/%m/%Y")
                            for fecha in fechas_disponibles[-30:]
                        )
                        st.write(fechas_texto)

            else:
                st.markdown(
                    f"##### Reservas del "
                    f"{fecha_registros_editar.strftime('%d/%m/%Y')}"
                )

                resumen_dia = registros_dia[
                    [
                        "id",
                        "Hora inicio",
                        "Hora fin",
                        "Profesor",
                        "Curso",
                        "Recurso",
                    ]
                ].copy()

                st.dataframe(
                    resumen_dia,
                    use_container_width=True,
                    hide_index=True,
                    column_config={
                        "id": st.column_config.NumberColumn(
                            "ID",
                            format="%d",
                        ),
                        "Hora inicio": st.column_config.TimeColumn(
                            "Inicio",
                            format="HH:mm",
                        ),
                        "Hora fin": st.column_config.TimeColumn(
                            "Fin",
                            format="HH:mm",
                        ),
                    },
                )

                opciones_edicion = {
                    int(fila["id"]): (
                        f"{bd_hora_texto(fila['Hora inicio'])}"
                        f"–{bd_hora_texto(fila['Hora fin'])} · "
                        f"{fila.get('Recurso', 'Sin recurso')} · "
                        f"{fila.get('Curso', 'Sin curso')} · "
                        f"{fila.get('Profesor', 'Sin profesor')}"
                    )
                    for _, fila in registros_dia.iterrows()
                }

                st.markdown("##### ✏️ Elige el registro exacto")

                id_editar = st.selectbox(
                    "Reserva a modificar",
                    list(opciones_edicion.keys()),
                    format_func=lambda valor: opciones_edicion[valor],
                    key=(
                        "bd_id_editar_"
                        f"{fecha_registros_editar.isoformat()}"
                    ),
                )

                fila_editar = datos_bd[
                    datos_bd["id"].astype(int) == int(id_editar)
                ].iloc[0]

                with st.container(border=True):
                    vista1, vista2, vista3, vista4 = st.columns(4)
                    vista1.metric(
                        "Horario",
                        (
                            f"{bd_hora_texto(fila_editar['Hora inicio'])}"
                            f"–{bd_hora_texto(fila_editar['Hora fin'])}"
                        ),
                    )
                    vista2.metric(
                        "Profesor",
                        str(fila_editar.get("Profesor") or "—"),
                    )
                    vista3.metric(
                        "Curso",
                        str(fila_editar.get("Curso") or "—"),
                    )
                    vista4.metric(
                        "Recurso",
                        str(fila_editar.get("Recurso") or "—"),
                    )

                st.markdown("##### 🛠️ Modificar reserva")

                with st.form(
                    f"bd_form_editar_{id_editar}",
                    clear_on_submit=False,
                ):
                    e1, e2, e3 = st.columns(3)

                    fecha_editada = e1.date_input(
                        "Fecha",
                        value=bd_fecha(
                            fila_editar["Fecha"]
                        ) or dt.date.today(),
                        format="DD/MM/YYYY",
                    )
                    hora_inicio_editada = e2.time_input(
                        "Hora de inicio",
                        value=bd_hora(
                            fila_editar["Hora inicio"]
                        ),
                        step=300,
                    )
                    hora_fin_editada = e3.time_input(
                        "Hora de término",
                        value=bd_hora(
                            fila_editar["Hora fin"]
                        ),
                        step=300,
                    )

                    e4, e5, e6 = st.columns(3)

                    profesor_editado = e4.selectbox(
                        "Profesor",
                        PROFESORES,
                        index=(
                            PROFESORES.index(
                                fila_editar["Profesor"]
                            )
                            if fila_editar["Profesor"] in PROFESORES
                            else 0
                        ),
                    )
                    curso_editado = e5.selectbox(
                        "Curso",
                        CURSOS,
                        index=(
                            CURSOS.index(
                                fila_editar["Curso"]
                            )
                            if fila_editar["Curso"] in CURSOS
                            else 0
                        ),
                    )
                    recurso_editado = e6.selectbox(
                        "Recurso",
                        RECURSOS,
                        index=(
                            RECURSOS.index(
                                fila_editar["Recurso"]
                            )
                            if fila_editar["Recurso"] in RECURSOS
                            else 0
                        ),
                    )

                    observaciones_editadas = st.text_area(
                        "Observaciones",
                        value=str(
                            fila_editar.get(
                                "Observaciones",
                                "",
                            )
                            or ""
                        ),
                        height=100,
                    )

                    notificar_edicion = st.checkbox(
                        "Notificar el cambio al profesor por correo",
                        value=True,
                    )

                    guardar_edicion = st.form_submit_button(
                        "💾 Guardar modificación",
                        type="primary",
                        use_container_width=True,
                    )

                if guardar_edicion:
                    if hora_inicio_editada >= hora_fin_editada:
                        st.error(
                            "La hora de término debe ser posterior "
                            "a la hora de inicio."
                        )
                    else:
                        try:
                            recurso_id_editado = map_rec.get(
                                recurso_editado
                            )

                            consulta_conflictos = (
                                supabase.table("reservas")
                                .select(
                                    "id,hora_inicio,hora_fin,"
                                    "profesores(nombre)"
                                )
                                .eq(
                                    "fecha",
                                    fecha_editada.isoformat(),
                                )
                                .eq(
                                    "recurso",
                                    recurso_id_editado,
                                )
                                .neq("id", int(id_editar))
                                .execute()
                            )

                            conflictos = []
                            for conflicto in (
                                consulta_conflictos.data or []
                            ):
                                if overlap(
                                    hora_inicio_editada,
                                    hora_fin_editada,
                                    bd_hora(
                                        conflicto.get(
                                            "hora_inicio"
                                        )
                                    ),
                                    bd_hora(
                                        conflicto.get(
                                            "hora_fin"
                                        )
                                    ),
                                ):
                                    nombre_conflicto = (
                                        (
                                            conflicto.get(
                                                "profesores"
                                            )
                                            or {}
                                        ).get("nombre")
                                        or "Otro usuario"
                                    )
                                    conflictos.append(
                                        nombre_conflicto
                                    )

                            if conflictos:
                                st.error(
                                    "No se puede guardar porque el "
                                    "recurso ya está reservado en ese "
                                    "horario por: "
                                    + ", ".join(
                                        sorted(
                                            set(conflictos)
                                        )
                                    )
                                )
                            else:
                                datos_actualizados = {
                                    "fecha": fecha_editada.isoformat(),
                                    "hora_inicio": (
                                        hora_inicio_editada
                                        .strftime("%H:%M:%S")
                                    ),
                                    "hora_fin": (
                                        hora_fin_editada
                                        .strftime("%H:%M:%S")
                                    ),
                                    "profesor": map_prof.get(
                                        profesor_editado
                                    ),
                                    "curso": map_cur.get(
                                        curso_editado
                                    ),
                                    "recurso": recurso_id_editado,
                                    "observaciones": (
                                        observaciones_editadas.strip()
                                    ),
                                }

                                (
                                    supabase.table("reservas")
                                    .update(datos_actualizados)
                                    .eq("id", int(id_editar))
                                    .execute()
                                )

                                registrar_auditoria(
                                    "editó reserva",
                                    "Base de datos",
                                    registro_id=id_editar,
                                    detalle={
                                        "fecha": (
                                            fecha_editada.isoformat()
                                        ),
                                        "profesor": profesor_editado,
                                        "curso": curso_editado,
                                        "recurso": recurso_editado,
                                    },
                                )

                                if notificar_edicion:
                                    email_to = PROFESOR_DATA.get(
                                        profesor_editado
                                    )
                                    if email_to:
                                        subject = (
                                            "Actualización de reserva "
                                            f"- {curso_editado}"
                                        )
                                        body = f"""
                                        <html>
                                        <body>
                                        <p>Hola {profesor_editado.split(' ')[0]},</p>
                                        <p>Se actualizó una reserva registrada a tu nombre:</p>
                                        <ul>
                                            <li><b>Fecha:</b> {format_date_es(fecha_editada)}</li>
                                            <li><b>Horario:</b> {hora_inicio_editada.strftime('%H:%M')}–{hora_fin_editada.strftime('%H:%M')}</li>
                                            <li><b>Curso:</b> {curso_editado}</li>
                                            <li><b>Recurso:</b> {recurso_editado}</li>
                                        </ul>
                                        <p>Saludos,<br>Sistema de Horarios CAV</p>
                                        </body>
                                        </html>
                                        """
                                        send_email(
                                            subject,
                                            body,
                                            email_to,
                                        )

                                st.success(
                                    "✅ Registro actualizado correctamente."
                                )
                                st.cache_data.clear()
                                time.sleep(0.15)
                                st.rerun()

                        except Exception as error:
                            registrar_error(
                                "editar_reserva_bd",
                                error,
                            )
                            st.error(
                                "No fue posible actualizar el registro. "
                                f"Detalle técnico: {error}"
                            )

        with pestaña_eliminar:
            st.markdown("#### 📅 Selecciona el día de las reservas a eliminar")
            st.caption(
                "Primero elige una fecha en el calendario. Después se "
                "mostrarán solamente las reservas de ese día para que puedas "
                "seleccionar exactamente cuál o cuáles deseas eliminar."
            )

            fechas_disponibles_eliminar = sorted(
                fecha
                for fecha in datos_bd["_fecha_filtro"].dropna().unique()
            )

            fecha_inicial_eliminar = dt.date.today()

            if (
                fecha_inicial_eliminar not in fechas_disponibles_eliminar
                and fechas_disponibles_eliminar
            ):
                fechas_futuras_eliminar = [
                    fecha
                    for fecha in fechas_disponibles_eliminar
                    if fecha >= dt.date.today()
                ]

                fecha_inicial_eliminar = (
                    fechas_futuras_eliminar[0]
                    if fechas_futuras_eliminar
                    else fechas_disponibles_eliminar[-1]
                )

            with st.container(border=True):
                col_fecha_eliminar, col_total_eliminar = st.columns([1, 2])

                fecha_registros_eliminar = col_fecha_eliminar.date_input(
                    "Fecha de la reserva",
                    value=fecha_inicial_eliminar,
                    format="DD/MM/YYYY",
                    key="bd_calendario_eliminar",
                )

                registros_dia_eliminar = datos_bd[
                    datos_bd["_fecha_filtro"]
                    == fecha_registros_eliminar
                ].copy()

                registros_dia_eliminar = registros_dia_eliminar.sort_values(
                    [
                        "_hora_inicio_filtro",
                        "Recurso",
                        "Profesor",
                    ],
                    ascending=[True, True, True],
                    na_position="last",
                )

                col_total_eliminar.metric(
                    "Reservas encontradas ese día",
                    len(registros_dia_eliminar),
                )

            if registros_dia_eliminar.empty:
                st.info(
                    "No existen reservas registradas para "
                    f"{fecha_registros_eliminar.strftime('%d/%m/%Y')}."
                )

                if fechas_disponibles_eliminar:
                    with st.expander(
                        "📌 Ver fechas que sí tienen reservas",
                        expanded=False,
                    ):
                        fechas_texto_eliminar = ", ".join(
                            fecha.strftime("%d/%m/%Y")
                            for fecha in fechas_disponibles_eliminar[-30:]
                        )
                        st.write(fechas_texto_eliminar)

            else:
                st.markdown(
                    f"##### Reservas del "
                    f"{fecha_registros_eliminar.strftime('%d/%m/%Y')}"
                )

                tabla_eliminar = registros_dia_eliminar[
                    [
                        "id",
                        "Hora inicio",
                        "Hora fin",
                        "Profesor",
                        "Curso",
                        "Recurso",
                        "Observaciones",
                    ]
                ].copy()

                # Selector visual mediante checkboxes dentro de una tabla.
                tabla_eliminar.insert(0, "Eliminar", False)

                seleccion_eliminar = st.data_editor(
                    tabla_eliminar,
                    hide_index=True,
                    use_container_width=True,
                    disabled=[
                        "id",
                        "Hora inicio",
                        "Hora fin",
                        "Profesor",
                        "Curso",
                        "Recurso",
                        "Observaciones",
                    ],
                    column_config={
                        "Eliminar": st.column_config.CheckboxColumn(
                            "Eliminar",
                            help=(
                                "Marca únicamente las reservas que deseas "
                                "eliminar."
                            ),
                        ),
                        "id": st.column_config.NumberColumn(
                            "ID",
                            format="%d",
                        ),
                        "Hora inicio": st.column_config.TimeColumn(
                            "Inicio",
                            format="HH:mm",
                        ),
                        "Hora fin": st.column_config.TimeColumn(
                            "Fin",
                            format="HH:mm",
                        ),
                        "Observaciones": st.column_config.TextColumn(
                            "Observaciones",
                            width="large",
                        ),
                    },
                    key=(
                        "bd_tabla_eliminar_"
                        f"{fecha_registros_eliminar.isoformat()}"
                    ),
                )

                ids_eliminar = (
                    seleccion_eliminar.loc[
                        seleccion_eliminar["Eliminar"],
                        "id",
                    ]
                    .astype(int)
                    .tolist()
                )

                if not ids_eliminar:
                    st.info(
                        "☝️ Marca en la primera columna la reserva o las "
                        "reservas que quieres eliminar."
                    )

                else:
                    st.markdown("##### 🗑️ Registros seleccionados")

                    detalle_eliminar = datos_bd[
                        datos_bd["id"]
                        .astype(int)
                        .isin(ids_eliminar)
                    ][
                        [
                            "id",
                            "Fecha",
                            "Hora inicio",
                            "Hora fin",
                            "Profesor",
                            "Curso",
                            "Recurso",
                            "Observaciones",
                        ]
                    ].copy()

                    st.dataframe(
                        detalle_eliminar,
                        use_container_width=True,
                        hide_index=True,
                        column_config={
                            "id": st.column_config.NumberColumn(
                                "ID",
                                format="%d",
                            ),
                            "Fecha": st.column_config.DateColumn(
                                "Fecha",
                                format="DD/MM/YYYY",
                            ),
                            "Hora inicio": st.column_config.TimeColumn(
                                "Inicio",
                                format="HH:mm",
                            ),
                            "Hora fin": st.column_config.TimeColumn(
                                "Fin",
                                format="HH:mm",
                            ),
                        },
                    )

                    # Resumen visual antes de eliminar.
                    resumen_1, resumen_2, resumen_3 = st.columns(3)

                    resumen_1.metric(
                        "Seleccionadas",
                        len(ids_eliminar),
                    )
                    resumen_2.metric(
                        "Profesores afectados",
                        detalle_eliminar["Profesor"]
                        .dropna()
                        .nunique(),
                    )
                    resumen_3.metric(
                        "Recursos liberados",
                        detalle_eliminar["Recurso"]
                        .dropna()
                        .nunique(),
                    )

                    st.warning(
                        "⚠️ Esta acción es permanente. Una vez eliminadas, "
                        "las reservas dejarán de aparecer en el calendario, "
                        "Modo Monitor y demás vistas del sistema."
                    )

                    notificar_eliminacion = st.checkbox(
                        "📧 Notificar por correo a los profesores afectados",
                        value=True,
                        key=(
                            "bd_notificar_eliminacion_"
                            f"{fecha_registros_eliminar.isoformat()}"
                        ),
                    )

                    texto_confirmacion_esperado = (
                        f"ELIMINAR {len(ids_eliminar)}"
                    )

                    st.markdown(
                        "Para confirmar, escribe exactamente:"
                    )
                    st.code(
                        texto_confirmacion_esperado,
                        language=None,
                    )

                    confirmacion_eliminar = st.text_input(
                        "Confirmación de eliminación",
                        placeholder=texto_confirmacion_esperado,
                        key=(
                            "bd_confirmacion_eliminar_"
                            f"{fecha_registros_eliminar.isoformat()}_"
                            f"{len(ids_eliminar)}"
                        ),
                    )

                    confirmacion_correcta = (
                        confirmacion_eliminar.strip().upper()
                        == texto_confirmacion_esperado
                    )

                    if not confirmacion_correcta:
                        st.caption(
                            "🔒 El botón permanecerá bloqueado hasta que la "
                            "confirmación coincida exactamente."
                        )

                    if st.button(
                        (
                            "🗑️ Eliminar "
                            f"{len(ids_eliminar)} "
                            + (
                                "reserva"
                                if len(ids_eliminar) == 1
                                else "reservas"
                            )
                        ),
                        type="primary",
                        use_container_width=True,
                        disabled=not confirmacion_correcta,
                        key=(
                            "bd_boton_eliminar_"
                            f"{fecha_registros_eliminar.isoformat()}_"
                            f"{len(ids_eliminar)}"
                        ),
                    ):
                        eliminados = 0
                        errores_eliminar = []

                        # Guardamos los datos originales antes de borrar.
                        filas_a_eliminar = datos_bd[
                            datos_bd["id"]
                            .astype(int)
                            .isin(ids_eliminar)
                        ].copy()

                        with st.spinner(
                            "Eliminando las reservas seleccionadas..."
                        ):
                            for _, fila_original in filas_a_eliminar.iterrows():
                                id_borrar = int(fila_original["id"])

                                try:
                                    (
                                        supabase.table("reservas")
                                        .delete()
                                        .eq("id", id_borrar)
                                        .execute()
                                    )

                                    registrar_auditoria(
                                        "eliminó reserva",
                                        "Base de datos",
                                        registro_id=id_borrar,
                                        detalle={
                                            "fecha": str(
                                                fila_original["Fecha"]
                                            ),
                                            "hora_inicio": (
                                                bd_hora_texto(
                                                    fila_original[
                                                        "Hora inicio"
                                                    ]
                                                )
                                            ),
                                            "hora_fin": (
                                                bd_hora_texto(
                                                    fila_original[
                                                        "Hora fin"
                                                    ]
                                                )
                                            ),
                                            "profesor": (
                                                fila_original["Profesor"]
                                            ),
                                            "curso": (
                                                fila_original["Curso"]
                                            ),
                                            "recurso": (
                                                fila_original["Recurso"]
                                            ),
                                        },
                                    )

                                    if notificar_eliminacion:
                                        profesor_borrado = (
                                            fila_original["Profesor"]
                                        )

                                        email_to = PROFESOR_DATA.get(
                                            profesor_borrado
                                        )

                                        if email_to:
                                            subject = (
                                                "Cancelación de reserva "
                                                f"- {fila_original['Curso']}"
                                            )

                                            body = f"""
                                            <html>
                                            <body>
                                                <p>
                                                    Hola {
                                                        str(
                                                            profesor_borrado
                                                        ).split(' ')[0]
                                                    },
                                                </p>

                                                <p>
                                                    Se canceló la siguiente
                                                    reserva registrada a tu
                                                    nombre:
                                                </p>

                                                <ul>
                                                    <li>
                                                        <b>Fecha:</b>
                                                        {
                                                            format_date_es(
                                                                bd_fecha(
                                                                    fila_original[
                                                                        'Fecha'
                                                                    ]
                                                                )
                                                            )
                                                        }
                                                    </li>
                                                    <li>
                                                        <b>Horario:</b>
                                                        {
                                                            bd_hora_texto(
                                                                fila_original[
                                                                    'Hora inicio'
                                                                ]
                                                            )
                                                        }–{
                                                            bd_hora_texto(
                                                                fila_original[
                                                                    'Hora fin'
                                                                ]
                                                            )
                                                        }
                                                    </li>
                                                    <li>
                                                        <b>Curso:</b>
                                                        {
                                                            fila_original[
                                                                'Curso'
                                                            ]
                                                        }
                                                    </li>
                                                    <li>
                                                        <b>Recurso:</b>
                                                        {
                                                            fila_original[
                                                                'Recurso'
                                                            ]
                                                        }
                                                    </li>
                                                </ul>

                                                <p>
                                                    Saludos,<br>
                                                    Sistema de Horarios CAV
                                                </p>
                                            </body>
                                            </html>
                                            """

                                            send_email(
                                                subject,
                                                body,
                                                email_to,
                                            )

                                    eliminados += 1

                                except Exception as error:
                                    errores_eliminar.append(
                                        f"ID {id_borrar}: {error}"
                                    )

                                    registrar_error(
                                        "eliminar_reserva_bd",
                                        error,
                                    )

                        if eliminados:
                            st.success(
                                "✅ "
                                f"{eliminados} "
                                + (
                                    "reserva fue eliminada "
                                    if eliminados == 1
                                    else "reservas fueron eliminadas "
                                )
                                + "correctamente."
                            )

                        if errores_eliminar:
                            st.error(
                                "Algunos registros no pudieron "
                                "eliminarse:\n\n"
                                + "\n".join(
                                    f"- {mensaje}"
                                    for mensaje
                                    in errores_eliminar
                                )
                            )
                        else:
                            st.cache_data.clear()
                            time.sleep(0.2)
                            st.rerun()

        with pestaña_exportar:
            st.markdown("#### Descargar resultados filtrados")

            exportar_df = filtrados.drop(
                columns=[
                    "_fecha_filtro",
                    "_hora_inicio_filtro",
                    "_texto_busqueda",
                ],
                errors="ignore",
            ).copy()

            st.write(
                f"El archivo incluirá **{len(exportar_df)}** registros."
            )

            csv_exportacion = exportar_df.to_csv(
                index=False,
            ).encode("utf-8-sig")

            st.download_button(
                "⬇️ Descargar CSV",
                data=csv_exportacion,
                file_name=(
                    "reservas_filtradas_"
                    f"{dt.date.today().isoformat()}.csv"
                ),
                mime="text/csv",
                type="primary",
                use_container_width=True,
            )

            st.markdown("#### Vista previa")
            st.dataframe(
                exportar_df.head(100),
                use_container_width=True,
                hide_index=True,
            )

# --- VISTA SEMANAL ---
elif page == "Semana": 
    import json # Necesario para la animación del monitor
    st.header("🗓️ Vista Semanal")
    
    def get_color_from_string(s):
        import hashlib
        hash_val = int(hashlib.md5(s.encode('utf-8')).hexdigest(), 16)
        hue = hash_val % 360
        return f"hsl({hue}, 75%, 50%)"

    with st.container(border=True):
        st.write("🔍 **Filtros de Búsqueda Avanzados**")
        col_d, col_r, col_p, col_c = st.columns(4)
        
        recursos_list = df['Recurso'].dropna().unique().tolist() if not df.empty else []
        profesores_list = df['Profesor'].dropna().unique().tolist() if not df.empty else []
        cursos_list = df['Curso'].dropna().unique().tolist() if not df.empty else []
        
        if not df.empty:
            df['fecha_obj'] = pd.to_datetime(df['Fecha']).dt.date
            default_date_week = dt.date.today() # ✅ CORREGIDO: Siempre usa la fecha de hoy
        else:
            default_date_week = dt.date.today()
            
        selected_date = col_d.date_input("Semana del", value=default_date_week, format="DD/MM/YYYY")
        
        selected_recursos = col_r.multiselect("Filtrar Recursos", recursos_list, placeholder="Todos")
        selected_profesores = col_p.multiselect("Filtrar Profesores", profesores_list, placeholder="Todos")
        selected_cursos = col_c.multiselect("Filtrar Cursos", cursos_list, placeholder="Todos")

    start_of_week = selected_date - dt.timedelta(days=selected_date.weekday())
    week_days = [start_of_week + dt.timedelta(days=i) for i in range(5)]
    
    if not df.empty:
        mask = (df['fecha_obj'] >= week_days[0]) & (df['fecha_obj'] <= week_days[-1])
        if selected_recursos: mask &= df['Recurso'].isin(selected_recursos)
        if selected_profesores: mask &= df['Profesor'].isin(selected_profesores)
        if selected_cursos: mask &= df['Curso'].isin(selected_cursos)
        df_week = df[mask]
    else:
        df_week = pd.DataFrame()

    st.markdown("---")
    
    # --- INTERRUPTOR MODO TV ---
    modo_tv = st.toggle("📺 Activar Modo Monitor (Pantalla Automática)", value=False)
    
    dias_es = {0: 'Lunes', 1: 'Martes', 2: 'Miércoles', 3: 'Jueves', 4: 'Viernes', 5: 'Sábado', 6: 'Domingo'}

    if modo_tv:
        # ======= MODO PANTALLA ANIMADA Y FULLSCREEN =======
        if df_week.empty:
            st.info("No hay reservas esta semana para mostrar en la pantalla.")
        else:
            df_tv = df_week.sort_values(by=['fecha_obj', 'Hora inicio'])
            
            # Preparar los datos para Javascript
            tv_data = []
            for _, row in df_tv.iterrows():
                tv_data.append({
                    "dia": dias_es[row['fecha_obj'].weekday()],
                    "fecha": row['fecha_obj'].strftime('%d/%m'),
                    "horario": f"{str(row['Hora inicio'])[:5]} a {str(row['Hora fin'])[:5]}",
                    "recurso": str(row['Recurso']),
                    "profesor": str(row['Profesor']),
                    "curso": str(row['Curso']),
                    "color": get_color_from_string(str(row['Profesor']))
                })
            
            tv_data_json = json.dumps(tv_data)
            
            # Código HTML/JS con Carrusel Automático y Botón FullScreen
            tv_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
            <style>
                @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;700;900&display=swap');
                body {{ margin: 0; padding: 0; background-color: #0f172a; font-family: 'Inter', sans-serif; color: white; overflow: hidden; }}
                .tv-wrapper {{ display: flex; flex-direction: column; height: 100vh; width: 100vw; background: radial-gradient(circle at top right, #1e293b, #0f172a); }}
                
                .header {{ display: flex; justify-content: space-between; align-items: center; padding: 25px 40px; background: rgba(15, 23, 42, 0.8); border-bottom: 2px solid #334155; box-shadow: 0 4px 15px rgba(0,0,0,0.5); }}
                .title {{ font-size: 32px; font-weight: 900; color: #38bdf8; letter-spacing: 2px; text-transform: uppercase; }}
                .fullscreen-btn {{ background: #0284c7; color: white; border: none; padding: 12px 24px; font-size: 18px; border-radius: 8px; cursor: pointer; font-weight: bold; transition: background 0.3s; }}
                .fullscreen-btn:hover {{ background: #38bdf8; color: #0f172a; }}
                
                .content {{ flex-grow: 1; padding: 40px; display: flex; flex-direction: column; justify-content: center; gap: 20px; }}
                
                .tv-card {{ background: rgba(255, 255, 255, 0.05); backdrop-filter: blur(10px); border-left: 10px solid; border-radius: 15px; padding: 30px; display: flex; align-items: center; box-shadow: 0 10px 25px rgba(0,0,0,0.3); opacity: 0; transform: translateY(30px); transition: all 0.6s cubic-bezier(0.4, 0, 0.2, 1); }}
                .tv-card.visible {{ opacity: 1; transform: translateY(0); }}
                
                .time-box {{ min-width: 280px; border-right: 3px solid rgba(255,255,255,0.1); margin-right: 40px; padding-right: 20px; }}
                .date-text {{ font-size: 22px; color: #94a3b8; font-weight: 700; text-transform: uppercase; letter-spacing: 1px; margin-bottom: 8px; }}
                .time-text {{ font-size: 40px; font-weight: 900; color: #e2e8f0; }}
                
                .details-box {{ flex-grow: 1; }}
                .resource-text {{ font-size: 45px; font-weight: 900; color: #38bdf8; margin-bottom: 12px; text-shadow: 0 2px 5px rgba(0,0,0,0.5); }}
                .prof-text {{ font-size: 28px; color: #cbd5e1; font-weight: 400; }}
                
                .progress-bar {{ position: absolute; bottom: 0; left: 0; height: 6px; background: #38bdf8; width: 0%; transition: width 8s linear; }}
            </style>
            </head>
            <body>
                <div class="tv-wrapper" id="tv-wrapper">
                    <div class="header">
                        <div class="title">📡 HORARIO DE ENLACES Y RECURSOS</div>
                        <button class="fullscreen-btn" onclick="toggleFullScreen()">🔲 Pantalla Completa</button>
                    </div>
                    <div class="content" id="cards-container"></div>
                    <div class="progress-bar" id="progress"></div>
                </div>

                <script>
                    const data = {tv_data_json};
                    const container = document.getElementById('cards-container');
                    const progress = document.getElementById('progress');
                    let currentIndex = 0;
                    const itemsPerPage = 4; // Muestra 4 bloques por pantalla
                    const slideDuration = 8000; // 8 segundos por pantalla

                    function toggleFullScreen() {{
                        const elem = document.documentElement;
                        if (!document.fullscreenElement) {{
                            elem.requestFullscreen().catch(err => console.log(err));
                        }} else {{
                            if (document.exitFullscreen) document.exitFullscreen();
                        }}
                    }}

                    function renderCards() {{
                        container.innerHTML = '';
                        
                        // Reiniciar barra de progreso
                        progress.style.transition = 'none';
                        progress.style.width = '0%';
                        setTimeout(() => {{
                            progress.style.transition = `width ${{slideDuration}}ms linear`;
                            progress.style.width = '100%';
                        }}, 50);

                        // Seleccionar los datos de esta página
                        const pageData = [];
                        for(let i=0; i<itemsPerPage; i++) {{
                            if(data.length > 0) {{
                                pageData.push(data[(currentIndex + i) % data.length]);
                            }}
                        }}
                        
                        // Eliminar duplicados visuales si hay menos reservas que itemsPerPage
                        const uniquePageData = [...new Set(pageData)];

                        uniquePageData.forEach((item, index) => {{
                            const card = document.createElement('div');
                            card.className = 'tv-card';
                            card.style.borderLeftColor = item.color;
                            card.innerHTML = `
                                <div class="time-box">
                                    <div class="date-text">${{item.dia}} ${{item.fecha}}</div>
                                    <div class="time-text">🕒 ${{item.horario}}</div>
                                </div>
                                <div class="details-box">
                                    <div class="resource-text">${{item.recurso}}</div>
                                    <div class="prof-text">👨‍🏫 ${{item.profesor}} &nbsp;|&nbsp; 📚 ${{item.curso}}</div>
                                </div>
                            `;
                            container.appendChild(card);
                            
                            // Animación en cascada para entrar
                            setTimeout(() => {{ card.classList.add('visible'); }}, index * 200);
                        }});

                        // Avanzar el índice
                        if (data.length > itemsPerPage) {{
                            currentIndex = (currentIndex + itemsPerPage) % data.length;
                        }}
                    }}

                    renderCards();
                    if(data.length > itemsPerPage) {{
                        setInterval(renderCards, slideDuration);
                    }} else {{
                        // Si hay pocas reservas, dejar barra llena y no rotar
                        progress.style.transition = 'width 1s linear';
                        progress.style.width = '100%';
                    }}
                </script>
            </body>
            </html>
            """
            
            # Usamos una altura alta para visualizar bien antes del fullscreen
            components.html(tv_html, height=800)

    else:
        # ======= MODO TABLA NORMAL =======
        column_names = [f"{dias_es[d.weekday()]}<br><span style='font-size:0.8em; font-weight:normal;'>{d.strftime('%d/%m')}</span>" for d in week_days]
        
        base_horas = [
            "08:00 a 08:45", "08:00 a 09:30", "08:45 a 09:30", "09:45 a 10:30",
            "09:45 a 11:15", "10:30 a 11:15", "11:30 a 12:15", "11:30 a 13:00",
            "12:15 a 13:00", "14:00 a 14:45", "14:00 a 15:30", "14:00 a 16:30",
            "14:45 a 15:30", "14:45 a 16:30", "15:45 a 16:30", "17:00 a 18:30",
            "17:30 a 18:30"
        ]
        
        if not df.empty:
            df['bloque_hora'] = df['Hora inicio'].astype(str).str[:5] + " a " + df['Hora fin'].astype(str).str[:5]
            dynamic_horas = df['bloque_hora'].unique().tolist()
        else:
            dynamic_horas = []

        HORAS = sorted(list(set(base_horas + dynamic_horas)))
        schedule = pd.DataFrame(index=HORAS, columns=column_names).fillna('')

        if not df_week.empty:
            for _, row in df_week.iterrows():
                day_col = f"{dias_es[row['fecha_obj'].weekday()]}<br><span style='font-size:0.8em; font-weight:normal;'>{row['fecha_obj'].strftime('%d/%m')}</span>"
                bloque_actual = f"{str(row['Hora inicio'])[:5]} a {str(row['Hora fin'])[:5]}"
                
                if day_col in column_names:
                    prof_color = get_color_from_string(str(row['Profesor']))
                    observacion = str(row['Observaciones']) if pd.notna(row['Observaciones']) and str(row['Observaciones']).strip() != '' else ""
                    icon = " 📌" if observacion else ""

                    card_content = f"<div style='font-weight:bold; color:#1e293b; margin-bottom:4px;'>{row['Recurso']}{icon}</div><div style='color:#334155; margin-bottom:2px; font-size:0.95em;'>👨‍🏫 {row['Profesor']}</div><div style='color:#64748b; font-style:italic; font-size:0.9em;'>📚 {row['Curso']}</div>"
                    
                    safe_observacion = html_sanitizer.escape(observacion)
                    card_html = f"<div style='background-color:#ffffff; padding:12px; margin-bottom:8px; border-radius:6px; border-left: 5px solid {prof_color}; box-shadow: 0 2px 5px rgba(0,0,0,0.08); font-size:0.85em; text-align:left; cursor:help;' title='{safe_observacion}'>{card_content}</div>"

                    if schedule.at[bloque_actual, day_col] == '': 
                        schedule.at[bloque_actual, day_col] = card_html
                    else: 
                        schedule.at[bloque_actual, day_col] += card_html

        st.markdown("""
            <style>
            .dataframe { width: 100%; border-collapse: separate; border-spacing: 0; font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; border-radius: 8px; overflow: hidden; box-shadow: 0 4px 6px rgba(0,0,0,0.05); }
            .dataframe th { background: linear-gradient(135deg, #4A90E2 0%, #357ABD 100%); color: white; padding: 15px 10px; text-align: center; font-weight: 600; text-transform: uppercase; letter-spacing: 0.5px; border-right: 1px solid rgba(255,255,255,0.2); }
            .dataframe th:last-child { border-right: none; }
            .dataframe td { border-bottom: 1px solid #E2E8F0; border-right: 1px solid #E2E8F0; padding: 12px; vertical-align: top; min-width: 170px; background-color: #F8FAFC; }
            .dataframe td:last-child { border-right: none; }
            .dataframe tbody tr:hover td { background-color: #F1F5F9; }
            .dataframe th[scope="row"] { background: #EDF2F7; color: #2D3748; font-weight: bold; text-align: center; width: 110px; border-right: 1px solid #CBD5E1; }
            </style>
        """, unsafe_allow_html=True)
        
        st.markdown(schedule.to_html(escape=False), unsafe_allow_html=True)
        
# --- PANEL PRINCIPAL (DASHBOARD) ---
elif page == "Dashboard":
    import tempfile
    import unicodedata
    import os
    import matplotlib.pyplot as plt
    import plotly.express as px
    
    st.header("📊 Panel de Supervisión General")

    # Definir colores institucionales
    COLOR_PRIMARIO = '#1E3A8A' 
    COLOR_SECUNDARIO = '#10B981' 

    # --- 1. FUNCIÓN MAESTRA: GENERAR PDF PROFESIONAL CON GRÁFICOS ---
    def generar_pdf_profesional_con_graficos(df_datos):
        from fpdf import FPDF
        
        # Plantilla del Documento Estilizada
        class PDF(FPDF):
            def header(self):
                self.set_fill_color(30, 58, 138)
                self.rect(0, 0, 210, 15, 'F')
                self.set_y(15)
                self.set_font('Arial', 'B', 20)
                self.set_text_color(255, 255, 255)
                self.cell(0, 15, 'REPORTE EJECUTIVO DE GESTION', 0, 1, 'C')
                self.set_text_color(31, 41, 55)
                self.ln(10)

            def footer(self):
                self.set_y(-20)
                self.set_font('Arial', 'I', 9)
                self.set_text_color(100, 100, 100)
                hoy_str = dt.date.today().strftime('%d/%m/%Y')
                self.cell(0, 10, f'Generado el: {hoy_str} | Sistema de Reservas Enlaces', 0, 0, 'L')
                self.cell(0, 10, f'Pagina {self.page_no()}', 0, 0, 'R')

            def section_title(self, label):
                self.set_font('Arial', 'B', 14)
                self.set_fill_color(243, 244, 246)
                self.set_text_color(30, 58, 138)
                self.cell(0, 10, f"  {label}", 0, 1, 'L', 1)
                self.set_text_color(31, 41, 55)
                self.ln(4)

        # Limpiar acentos para evitar errores en PDF
        def s(texto):
            if not texto: return "N/A"
            texto_str = str(texto)
            nfkd = unicodedata.normalize('NFKD', texto_str)
            return u"".join([c for c in nfkd if not unicodedata.combining(c)])

        pdf = PDF()
        pdf.set_auto_page_break(auto=True, margin=25)
        pdf.add_page()
        pdf.set_font('Arial', '', 12)

        if df_datos.empty:
            pdf.cell(0, 10, s('No hay datos registrados para analizar.'), 0, 1)
        else:
            C_FECHA = 'Fecha'
            C_REC = 'Recurso'
            C_PROF = 'Profesor'

            # Cálculos rápidos
            total_res = len(df_datos)
            df_datos['fecha_obj'] = pd.to_datetime(df_datos[C_FECHA]).dt.date
            hoy_res = len(df_datos[df_datos['fecha_obj'] == dt.date.today()])
            recurso_top = df_datos[C_REC].mode()[0] if not df_datos[C_REC].empty else "N/A"
            profesor_top = df_datos[C_PROF].mode()[0] if not df_datos[C_PROF].empty else "N/A"

            # Calcular rango de fechas para el subtítulo
            fecha_min = df_datos['fecha_obj'].min().strftime('%d/%m/%Y')
            fecha_max = df_datos['fecha_obj'].max().strftime('%d/%m/%Y')

            # --- TÍTULO Y RANGO DE FECHAS EN EL DOCUMENTO ---
            pdf.set_font('Arial', 'B', 16)
            pdf.set_text_color(30, 58, 138) # Azul institucional
            pdf.cell(0, 10, s('INFORME ESTADISTICO DE USO DE RECURSOS'), 0, 1, 'C')
            
            pdf.set_font('Arial', 'I', 11)
            pdf.set_text_color(100, 100, 100) # Gris oscuro
            pdf.cell(0, 6, s(f'Periodo analizado: {fecha_min} al {fecha_max}'), 0, 1, 'C')
            pdf.ln(8)
            pdf.set_text_color(31, 41, 55) # Resetear a color texto normal

            # =========================================================
            # SECCIÓN 1: RESUMEN EJECUTIVO
            # =========================================================
            pdf.section_title('1. Resumen de Metricas Clave')
            
            pdf.set_fill_color(255, 255, 255)
            pdf.set_font('Arial', 'B', 12)
            pdf.cell(95, 10, s('Concepto'), 1, 0, 'C', 1)
            pdf.cell(95, 10, s('Valor'), 1, 1, 'C', 1)
            
            pdf.set_font('Arial', '', 11)
            pdf.cell(95, 9, s('Total de Reservas Historicas'), 1, 0, 'L')
            pdf.cell(95, 9, f"{total_res}", 1, 1, 'R')
            pdf.cell(95, 9, s('Reservas Agendadas para Hoy'), 1, 0, 'L')
            pdf.cell(95, 9, f"{hoy_res}", 1, 1, 'R')
            
            pdf.set_font('Arial', 'B', 11)
            pdf.cell(95, 9, s('Recurso Mas Ocupado (Estrella)'), 1, 0, 'L')
            pdf.cell(95, 9, s(recurso_top), 1, 1, 'R')
            pdf.cell(95, 9, s('Profesor con Mas Solicitudes'), 1, 0, 'L')
            pdf.cell(95, 9, s(profesor_top), 1, 1, 'R')
            
            pdf.ln(12)

            # =========================================================
            # SECCIÓN 2: GRÁFICOS CON ARCHIVOS TEMPORALES
            # =========================================================
            pdf.section_title('2. Analisis Visual de Ocupacion')
            pdf.set_font('Arial', '', 11)
            pdf.multi_cell(0, 7, s('A continuacion se presentan los graficos detallados de uso de recursos y distribucion temporal de las reservas.'), 0, 'L')
            pdf.ln(5)

            # --- GRÁFICO A: TOP 5 RECURSOS ---
            pdf.set_font('Arial', 'I', 10)
            pdf.cell(0, 8, s('Grafico A: Top 5 Recursos Mas Utilizados'), 0, 1, 'C')
            
            plt.figure(figsize=(8, 4))
            top5_data = df_datos[C_REC].value_counts().nlargest(5).sort_values(ascending=True)
            names_clean = [s(name) for name in top5_data.index]
            
            bars = plt.barh(names_clean, top5_data.values, color='#3B82F6', edgecolor='#1E3A8A', height=0.7)
            plt.xlabel('Cantidad de Reservas', fontsize=10, fontweight='bold', color='#4B5563')
            plt.title('Top 5 Recursos', fontsize=12, fontweight='bold', color='#1E3A8A')
            plt.gca().spines['top'].set_visible(False)
            plt.gca().spines['right'].set_visible(False)
            plt.grid(axis='x', linestyle='--', alpha=0.5)
            
            for bar in bars:
                width = bar.get_width()
                plt.text(width + (max(top5_data.values)*0.01), bar.get_y() + bar.get_height()/2, 
                         f'{int(width)}', va='center', fontsize=9, fontweight='bold', color='#1E3A8A')
            plt.tight_layout()
            
            # Guardar en archivo temporal físico
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_a:
                plt.savefig(tmp_a.name, format='png', dpi=150, bbox_inches='tight')
                img_path_a = tmp_a.name
            plt.close()
            
            # Insertar en PDF y eliminar archivo
            current_y = pdf.get_y()
            pdf.image(img_path_a, x=25, y=current_y, w=160)
            os.remove(img_path_a)
            
            # --- CORRECCIÓN VISUAL: FORZAR PÁGINA 2 PARA EL GRÁFICO DE PASTEL ---
            pdf.add_page() 
            
            # --- GRÁFICO B: DISTRIBUCIÓN DE USO ---
            pdf.set_font('Arial', 'I', 10)
            pdf.cell(0, 8, s('Grafico B: Distribucion de Uso por Tipo de Recurso (Total)'), 0, 1, 'C')
            
            plt.figure(figsize=(6, 6))
            uso_total = df_datos[C_REC].value_counts()
            colores_pie = ['#1E3A8A', '#10B981', '#F59E0B', '#EF4444', '#8B5CF6', '#EC4899', '#6B7280']
            labels_clean = [s(name) for name in uso_total.index]
            
            patches, texts, autotexts = plt.pie(uso_total.values, labels=labels_clean, autopct='%1.1f%%', 
                                              startangle=140, colors=colores_pie, shadow=False,
                                              wedgeprops={'edgecolor': 'white', 'linewidth': 2})
            
            plt.title('Uso General de Recursos', fontsize=12, fontweight='bold', color='#1E3A8A')
            for text in texts: text.set_color('#4B5563'); text.set_fontsize(9)
            for autotext in autotexts: autotext.set_color('white'); autotext.set_fontweight('bold'); autotext.set_fontsize(9)
            
            plt.axis('equal')
            plt.tight_layout()
            
            # Guardar en archivo temporal físico
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_b:
                plt.savefig(tmp_b.name, format='png', dpi=150, bbox_inches='tight')
                img_path_b = tmp_b.name
            plt.close()
            
            # Insertar en PDF (ahora en la parte superior de la página 2)
            current_y_b = pdf.get_y()
            pdf.image(img_path_b, x=55, y=current_y_b, w=100)
            os.remove(img_path_b)

        return pdf.output(dest='S').encode('latin-1')

    # --- 2. BOTÓN DE DESCARGA PDF ---
    col_titulo, col_boton = st.columns([2, 1])
    with col_boton:
        if not df.empty:
            with st.spinner("Preparando reporte PDF..."):
                try:
                    pdf_data = generar_pdf_profesional_con_graficos(df)
                    st.download_button(
                        label="📄 Descargar Reporte PDF",
                        data=pdf_data,
                        file_name=f"Reporte_Gestión_Enlaces_{dt.date.today().strftime('%d_%m_%Y')}.pdf",
                        mime="application/pdf",
                        type="primary",
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"⚠️ Error generando PDF: {e}")

    # --- 3. MÉTRICAS Y GRÁFICOS EN PANTALLA ---
    if not df.empty:
        C_FECHA = 'Fecha'
        C_REC = 'Recurso'
        C_PROF = 'Profesor'
        C_HORA_I = 'Hora inicio'
        C_HORA_F = 'Hora fin'
        
        df['fecha_obj'] = pd.to_datetime(df[C_FECHA]).dt.date
        hoy = dt.date.today()
        reservas_hoy = df[df['fecha_obj'] == hoy]
        
        c1, c2, c3, c4 = st.columns(4)
        style_kpi = "background:white; border-radius:12px; padding: 25px; box-shadow: 0 10px 15px -3px rgba(0,0,0,0.1); border-top: 5px solid;"

        with c1:
            st.markdown(f"""
                <div style="{style_kpi} border-top-color: #3B82F6;">
                    <div style="font-size:1em; color:#64748b; font-weight:bold;">TOTAL RESERVAS</div>
                    <div style="font-size:3em; font-weight:900; color:#1E3A8A;">{len(df)}</div>
                </div>
            """, unsafe_allow_html=True)
        with c2:
            st.markdown(f"""
                <div style="{style_kpi} border-top-color: #EF4444;">
                    <div style="font-size:1em; color:#64748b; font-weight:bold;">RESERVAS HOY</div>
                    <div style="font-size:3em; font-weight:900; color:#B91C1C;">{len(reservas_hoy)}</div>
                </div>
            """, unsafe_allow_html=True)
        with c3:
            recurso_top = df[C_REC].mode()[0] if not df[C_REC].empty else "N/A"
            st.markdown(f"""
                <div style="{style_kpi} border-top-color: #F59E0B;">
                    <div style="font-size:1em; color:#64748b; font-weight:bold;">RECURSO ESTRELLA</div>
                    <div style="font-size:1.8em; font-weight:800; color:#B45309;">{recurso_top}</div>
                </div>
            """, unsafe_allow_html=True)
        with c4:
            profesor_top = df[C_PROF].mode()[0] if not df[C_PROF].empty else "N/A"
            prof_display = str(profesor_top)[:16] + ".." if len(str(profesor_top)) > 18 else str(profesor_top)
            st.markdown(f"""
                <div style="{style_kpi} border-top-color: #8B5CF6;">
                    <div style="font-size:1em; color:#64748b; font-weight:bold;">PROF. FRECUENTE</div>
                    <div style="font-size:1.8em; font-weight:800; color:#5B21B6;">{prof_display}</div>
                </div>
            """, unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        col_graf_1, col_graf_2 = st.columns([2, 1])
        
        with col_graf_1:
            st.markdown("### 🗓️ Mapa de Calor: Ocupación Semanal")
            df['fecha_dt'] = pd.to_datetime(df[C_FECHA])
            dias_map = {0: '1-Lunes', 1: '2-Martes', 2: '3-Miércoles', 3: '4-Jueves', 4: '5-Viernes', 5: '6-Sábado', 6: '7-Domingo'}
            df['Dia_Semana'] = df['fecha_dt'].dt.dayofweek.map(dias_map)
            df['Bloque_Ordenado'] = df[C_HORA_I].astype(str).str[:5] + " - " + df[C_HORA_F].astype(str).str[:5]
            
            df_heatmap_filter = df[df['fecha_dt'].dt.dayofweek < 5]
            heatmap_data = df_heatmap_filter.groupby(['Bloque_Ordenado', 'Dia_Semana']).size().reset_index(name='Cantidad')
            
            if not heatmap_data.empty:
                fig_heat = px.density_heatmap(heatmap_data, x='Dia_Semana', y='Bloque_Ordenado', z='Cantidad',
                                             text_auto=True, color_continuous_scale='Viridis')
                fig_heat.update_layout(xaxis_nticks=5, yaxis={'categoryorder':'category ascending'}, height=450)
                fig_heat.update_xaxes(tickformat="%A", labelalias={d: d.split('-')[1] for d in dias_map.values()})
                st.plotly_chart(fig_heat, use_container_width=True, config={'displayModeBar': False})
            else:
                st.info("Datos insuficientes para el mapa de calor.")

        with col_graf_2:
            st.markdown("### 🏆 Top 5 Recursos")
            uso_recursos = df[C_REC].value_counts().reset_index()
            uso_recursos.columns = ['Recurso', 'Cantidad']
            top5_recursos = uso_recursos.nlargest(5, 'Cantidad').sort_values(by='Cantidad', ascending=True)
            
            if not top5_recursos.empty:
                fig_top5 = px.bar(top5_recursos, x='Cantidad', y='Recurso', orientation='h', text='Cantidad',
                                 color='Cantidad', color_continuous_scale='Blues')
                fig_top5.update_layout(showlegend=False, coloraxis_showscale=False, height=450)
                fig_top5.update_traces(textposition='outside')
                st.plotly_chart(fig_top5, use_container_width=True, config={'displayModeBar': False})

        st.markdown("<br>### 📊 Distribución Global de Recursos", unsafe_allow_html=True)
        fig_full = px.bar(uso_recursos, x='Recurso', y='Cantidad', color='Recurso', text_auto=True)
        fig_full.update_layout(showlegend=False, height=350, xaxis_tickangle=-45)
        st.plotly_chart(fig_full, use_container_width=True)
        
    else:
        st.info("No hay reservas registradas en el sistema todavía.")
# ==============================================================================
# --- SECCIÓN TÉCNICOS (TICKETS, BAJAS INDEPENDIENTES Y CÓDIGOS QR) ---
# ==============================================================================
elif page == "Técnicos":
    import io
    from io import BytesIO
    import qrcode
    try:
        import docx
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        st.error("🚨 Falta la librería 'python-docx'. Instálala con: pip install python-docx")
        st.stop()

    st.header("🛠️ Panel de Soporte Técnico")

   
    
    opciones_modulo_tecnico = [
        "🎫 Tickets",
        "🗑️ Baja de Equipos",
        "📋 Generador QR",
    ]

    modulo_tecnico_pendiente = st.session_state.pop(
        "_pending_tecnico_modulo",
        None,
    )

    if modulo_tecnico_pendiente in opciones_modulo_tecnico:
        st.session_state.tecnico_modulo = modulo_tecnico_pendiente

    if (
        "tecnico_modulo" not in st.session_state
        or st.session_state.tecnico_modulo not in opciones_modulo_tecnico
    ):
        st.session_state.tecnico_modulo = "🎫 Tickets"

    modulo_tec = st.radio(
        "Selecciona el módulo de trabajo:",
        opciones_modulo_tecnico,
        horizontal=True,
        key="tecnico_modulo",
    )
    st.markdown("---")

    # ---------------------------------------------------------
    # MÓDULO 1: TICKETS (GESTIÓN DE FALLAS)
    # ---------------------------------------------------------
    if modulo_tec == "🎫 Tickets":
        st.subheader("Gestión de Tickets ingresados vía QR")
        try:
            mant_data = supabase.table("mantenimientos").select("*, recursos(nombre)").order("fecha", desc=True).execute().data
        except Exception as e:
            st.error(f"Error consultando mantenimientos: {e}")
            mant_data = []

        if mant_data:
            df_mant = pd.DataFrame(mant_data)
            df_mant['Recurso'] = df_mant['recursos'].apply(lambda x: x.get('nombre', 'Desconocido') if isinstance(x, dict) else "Desconocido")
            
            if 'notas_tecnico' not in df_mant.columns: 
                df_mant['notas_tecnico'] = ""
            else: 
                df_mant['notas_tecnico'] = df_mant['notas_tecnico'].fillna("")
            
            # Métricas
            pendientes = len(df_mant[df_mant['estado'] == 'Reportado (Vía QR)'])
            en_revision = len(df_mant[df_mant['estado'] == 'En Revisión'])
            resueltos = len(df_mant[df_mant['estado'] == 'Resuelto'])
            
            c1, c2, c3 = st.columns(3)
            estilo_metrica = "background:white; border-radius:12px; padding: 15px; box-shadow: 0 4px 6px rgba(0,0,0,0.05); border-left: 5px solid"
            with c1: st.markdown(f'<div style="{estilo_metrica} #EF4444;"><div style="color:#64748b; font-size:0.85em; font-weight:bold;">🔴 PENDIENTES</div><div style="font-size:2.2em; font-weight:900; color:#B91C1C;">{pendientes}</div></div>', unsafe_allow_html=True)
            with c2: st.markdown(f'<div style="{estilo_metrica} #F59E0B;"><div style="color:#64748b; font-size:0.85em; font-weight:bold;">🟡 EN REVISIÓN</div><div style="font-size:2.2em; font-weight:900; color:#D97706;">{en_revision}</div></div>', unsafe_allow_html=True)
            with c3: st.markdown(f'<div style="{estilo_metrica} #10B981;"><div style="color:#64748b; font-size:0.85em; font-weight:bold;">🟢 RESUELTOS</div><div style="font-size:2.2em; font-weight:900; color:#047857;">{resueltos}</div></div>', unsafe_allow_html=True)
            
            st.markdown("<br>", unsafe_allow_html=True)
            
            t_pendientes, t_revision, t_resueltos = st.tabs(["🔴 Pendientes", "🟡 En Revisión", "🟢 Resueltos"])
            
            def renderizar_tickets(df_filtrado, color_icon, estados_destino):
                if df_filtrado.empty:
                    st.info("✨ No hay tickets en esta categoría actualmente.")
                    return
                for _, row in df_filtrado.iterrows():
                    with st.expander(f"{color_icon} Ticket #{row['id']} | {row['Recurso']} | Fecha: {row['fecha']}"):
                        quien_reporto = row.get('reportado_por', 'No registrado')
                        st.markdown(f"**👤 Reportado por:** {quien_reporto}")
                        st.markdown(f"**📝 Falla:**\n> {row['descripcion']}")
                        
                        # --- INTEGRACIÓN GEMINI EN TICKETS ---
                        if st.button("🤖 Obtener Diagnóstico Sugerido (IA)", key=f"ai_btn_{row['id']}"):
                            with st.spinner("Gemini está analizando la falla..."):
                                prompt_tecnico = f"Actúa como un experto en soporte técnico informático. El equipo '{row['Recurso']}' tiene este problema: '{row['descripcion']}'. Dame una solución técnica breve en 3 puntos."
                                sugerencia = consultar_gemini(prompt_tecnico)
                                st.info(f"**💡 Sugerencia de Gemini:**\n\n{sugerencia}")
                        
                        if row['notas_tecnico']: 
                            st.markdown(f"**🛠️ Notas Previas:**\n> {row['notas_tecnico']}")
                        
                        st.markdown("---")
                        col_a, col_b = st.columns([1, 2])
                        with col_a: 
                            nuevo_est = st.selectbox("Cambiar estado a:", estados_destino, key=f"est_{row['id']}")
                        with col_b: 
                            nueva_nota = st.text_area("Notas de Reparación:", value=row['notas_tecnico'], key=f"not_{row['id']}", height=68)

                        if st.button("💾 Guardar Cambios", key=f"btn_{row['id']}", type="primary", use_container_width=True):
                            try:
                                supabase.table("mantenimientos").update({
                                    "estado": nuevo_est, 
                                    "notas_tecnico": html_sanitizer.escape(nueva_nota).strip()
                                }).eq("id", row['id']).execute()
                                st.success("Ticket actualizado.")
                                time.sleep(0.25); st.rerun()
                            except Exception as e: 
                                st.error(f"Error técnico: {e}")

            with t_pendientes: renderizar_tickets(df_mant[df_mant['estado'] == 'Reportado (Vía QR)'], "🔴", ["En Revisión", "Resuelto", "Reportado (Vía QR)"])
            with t_revision: renderizar_tickets(df_mant[df_mant['estado'] == 'En Revisión'], "🟡", ["Resuelto", "En Revisión", "Reportado (Vía QR)"])
            with t_resueltos: renderizar_tickets(df_mant[df_mant['estado'] == 'Resuelto'], "🟢", ["Resuelto", "En Revisión", "Reportado (Vía QR)"])
        else: 
            st.info("No hay reportes de mantenimiento.")

    # ---------------------------------------------------------
    # MÓDULO 2: BAJA DE EQUIPOS
    # ---------------------------------------------------------
    elif modulo_tec == "🗑️ Baja de Equipos":
        from PIL import Image, ImageOps
        from docx.shared import Cm
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        st.subheader("🗑️ Procesar Baja y Generar Informe")
        st.caption("Genera un informe institucional en Word y adjunta todas las fotografías necesarias.")

        COLOR_BURDEO = RGBColor(128, 0, 32)
        RUTA_LOGO_BAJA = Path(__file__).parent / "logocav.png"

        def texto_seguro(valor, por_defecto=""):
            if valor is None:
                return por_defecto
            texto = str(valor).strip()
            return texto if texto else por_defecto

        def nombre_archivo_seguro(texto):
            limpio = re.sub(r"[^A-Za-z0-9ÁÉÍÓÚÜÑáéíóúüñ_-]+", "_", texto_seguro(texto, "Equipo"))
            return limpio.strip("_") or "Equipo"

        def quitar_bordes_tabla(tabla):
            """Quita los bordes visibles de una tabla Word."""
            tbl = tabla._tbl
            tbl_pr = tbl.tblPr
            borders = tbl_pr.first_child_found_in("w:tblBorders")
            if borders is None:
                borders = OxmlElement("w:tblBorders")
                tbl_pr.append(borders)

            for borde in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = "w:" + borde
                elemento = borders.find(qn(tag))
                if elemento is None:
                    elemento = OxmlElement(tag)
                    borders.append(elemento)
                elemento.set(qn("w:val"), "nil")

        def sombrear_celda(celda, color_hex):
            tc_pr = celda._tc.get_or_add_tcPr()
            shd = tc_pr.find(qn("w:shd"))
            if shd is None:
                shd = OxmlElement("w:shd")
                tc_pr.append(shd)
            shd.set(qn("w:fill"), color_hex)

        def configurar_fuente_run(run, tamano=11, negrita=False, color=None):
            run.font.name = "Arial"
            run.font.size = Pt(tamano)
            run.bold = negrita
            if color:
                run.font.color.rgb = color
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

        def agregar_titulo_seccion(documento, texto):
            p = documento.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(5)
            r = p.add_run(texto)
            configurar_fuente_run(r, tamano=12, negrita=True, color=COLOR_BURDEO)
            return p

        def preparar_foto_para_word(foto_bytes):
            """
            Corrige orientación, reduce imágenes demasiado grandes y devuelve
            un BytesIO compatible con Word.
            """
            entrada = BytesIO(foto_bytes)
            imagen = Image.open(entrada)
            imagen = ImageOps.exif_transpose(imagen)

            if imagen.mode not in ("RGB", "L"):
                fondo = Image.new("RGB", imagen.size, "white")
                if "A" in imagen.getbands():
                    fondo.paste(imagen, mask=imagen.getchannel("A"))
                else:
                    fondo.paste(imagen)
                imagen = fondo
            elif imagen.mode == "L":
                imagen = imagen.convert("RGB")

            imagen.thumbnail((1800, 1800))
            salida = BytesIO()
            imagen.save(salida, format="JPEG", quality=86, optimize=True)
            salida.seek(0)
            return salida

        def generar_docx_baja(datos, fotos=None):
            """
            Genera el Word con formato institucional.
            fotos: lista de bytes de imágenes JPG/PNG.
            """
            try:
                fotos = fotos or []
                document = Document()

                # Márgenes y fuente general
                section = document.sections[0]
                section.top_margin = Cm(1.6)
                section.bottom_margin = Cm(1.6)
                section.left_margin = Cm(1.8)
                section.right_margin = Cm(1.8)

                estilo_normal = document.styles["Normal"]
                estilo_normal.font.name = "Arial"
                estilo_normal.font.size = Pt(11)
                estilo_normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

                # ==========================================================
                # ENCABEZADO INSTITUCIONAL
                # ==========================================================
                header_table = document.add_table(rows=1, cols=2)
                header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                header_table.autofit = False
                header_table.columns[0].width = Cm(3.2)
                header_table.columns[1].width = Cm(13.5)
                quitar_bordes_tabla(header_table)

                celda_logo = header_table.cell(0, 0)
                celda_logo.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p_logo = celda_logo.paragraphs[0]
                p_logo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if RUTA_LOGO_BAJA.exists():
                    p_logo.add_run().add_picture(str(RUTA_LOGO_BAJA), width=Cm(2.8))

                celda_titulo = header_table.cell(0, 1)
                celda_titulo.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p_titulo = celda_titulo.paragraphs[0]
                p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                r_titulo = p_titulo.add_run("INFORME TÉCNICO DE\nBAJA DE EQUIPOS")
                configurar_fuente_run(r_titulo, tamano=18, negrita=True, color=COLOR_BURDEO)

                p_linea = document.add_paragraph()
                p_linea.paragraph_format.space_after = Pt(8)
                r_linea = p_linea.add_run("_" * 82)
                configurar_fuente_run(r_linea, tamano=8, color=RGBColor(80, 80, 80))

                # DE / FECHA / MATERIA
                meta = document.add_table(rows=3, cols=2)
                meta.alignment = WD_TABLE_ALIGNMENT.CENTER
                meta.autofit = False
                meta.columns[0].width = Cm(4)
                meta.columns[1].width = Cm(10)
                quitar_bordes_tabla(meta)

                fecha_doc = texto_seguro(datos.get("fecha_baja"))
                try:
                    fecha_obj = dt_datetime.strptime(fecha_doc, "%Y-%m-%d").date()
                    fecha_doc = format_date_es(fecha_obj)
                except Exception:
                    pass

                metadatos = [
                    ("DE:", "Departamento de Informática."),
                    ("FECHA:", fecha_doc),
                    ("MATERIA:", texto_seguro(datos.get("materia"), "Baja técnica")),
                ]
                for fila, (etiqueta, valor) in enumerate(metadatos):
                    p_et = meta.cell(fila, 0).paragraphs[0]
                    p_et.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    r_et = p_et.add_run(etiqueta)
                    configurar_fuente_run(r_et, tamano=11, negrita=True, color=COLOR_BURDEO)

                    p_val = meta.cell(fila, 1).paragraphs[0]
                    r_val = p_val.add_run(valor)
                    configurar_fuente_run(r_val, tamano=11)

                document.add_paragraph()

                # ==========================================================
                # 1. ANTECEDENTES Y CAUSAS
                # ==========================================================
                agregar_titulo_seccion(document, "1. ANTECEDENTES Y CAUSAS")
                p_ant = document.add_paragraph()
                p_ant.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p_ant.paragraph_format.line_spacing = 1.15
                antecedentes = texto_seguro(
                    datos.get("antecedentes") or datos.get("justificacion"),
                    "No se registraron antecedentes adicionales."
                )
                r_ant = p_ant.add_run(antecedentes)
                configurar_fuente_run(r_ant, tamano=11)

                # ==========================================================
                # 2. DIAGNÓSTICO
                # ==========================================================
                agregar_titulo_seccion(document, "2. DIAGNÓSTICO")
                p_diag = document.add_paragraph()
                p_diag.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p_diag.paragraph_format.line_spacing = 1.15
                r_diag = p_diag.add_run(
                    texto_seguro(datos.get("diagnosis"), "Sin diagnóstico técnico registrado.")
                )
                configurar_fuente_run(r_diag, tamano=11)

                # ==========================================================
                # 3. DETALLE DE EQUIPOS PARA BAJA
                # ==========================================================
                agregar_titulo_seccion(document, "3. DETALLE DE EQUIPOS PARA BAJA")
                tabla = document.add_table(rows=2, cols=5)
                tabla.style = "Table Grid"
                tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
                tabla.autofit = True

                encabezados = ["Cantidad", "Equipo", "Marca", "Modelo", "Estado"]
                for i, titulo_columna in enumerate(encabezados):
                    celda = tabla.cell(0, i)
                    sombrear_celda(celda, "800020")
                    celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    p = celda.paragraphs[0]
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    r = p.add_run(titulo_columna)
                    configurar_fuente_run(r, tamano=9, negrita=True, color=RGBColor(255, 255, 255))

                valores = [
                    str(datos.get("cantidad_baja", 1)),
                    texto_seguro(datos.get("recurso_nombre"), "Equipo"),
                    texto_seguro(datos.get("marca"), "S/M"),
                    texto_seguro(datos.get("modelo"), "S/M"),
                    texto_seguro(datos.get("estado_equipo") or datos.get("diagnosis"), "Para baja"),
                ]
                for i, valor in enumerate(valores):
                    celda = tabla.cell(1, i)
                    celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    p = celda.paragraphs[0]
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if i == 0 else WD_PARAGRAPH_ALIGNMENT.LEFT
                    r = p.add_run(valor)
                    configurar_fuente_run(r, tamano=9)

                # Datos complementarios
                complementos = []
                if texto_seguro(datos.get("serie")):
                    complementos.append(f"N.° de serie / inventario: {texto_seguro(datos.get('serie'))}")
                if texto_seguro(datos.get("ubicacion")):
                    complementos.append(f"Ubicación habitual: {texto_seguro(datos.get('ubicacion'))}")
                if texto_seguro(datos.get("fecha_adquisicion")):
                    complementos.append(f"Fecha o año de adquisición: {texto_seguro(datos.get('fecha_adquisicion'))}")

                if complementos:
                    p_comp = document.add_paragraph()
                    p_comp.paragraph_format.space_before = Pt(6)
                    for indice, complemento in enumerate(complementos):
                        r = p_comp.add_run(("• " if indice == 0 else "\n• ") + complemento)
                        configurar_fuente_run(r, tamano=10)

                recomendacion = texto_seguro(datos.get("recomendacion"))
                if recomendacion:
                    agregar_titulo_seccion(document, "4. RECOMENDACIÓN TÉCNICA")
                    p_rec = document.add_paragraph()
                    p_rec.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    r_rec = p_rec.add_run(recomendacion)
                    configurar_fuente_run(r_rec, tamano=11)

                # Firma responsable al cierre del informe
                document.add_paragraph()
                firma = document.add_table(rows=2, cols=2)
                firma.alignment = WD_TABLE_ALIGNMENT.CENTER
                quitar_bordes_tabla(firma)

                etiquetas_firma = [
                    ("Técnico responsable:", texto_seguro(datos.get("tecnico_responsable"), "No indicado")),
                    ("Departamento:", "Departamento de Informática"),
                ]
                for fila, (etiqueta, valor) in enumerate(etiquetas_firma):
                    p1 = firma.cell(fila, 0).paragraphs[0]
                    p1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    r1 = p1.add_run(etiqueta)
                    configurar_fuente_run(r1, tamano=10, negrita=True, color=COLOR_BURDEO)

                    p2 = firma.cell(fila, 1).paragraphs[0]
                    r2 = p2.add_run(valor)
                    configurar_fuente_run(r2, tamano=10)

                # ==========================================================
                # ANEXO FOTOGRÁFICO
                # ==========================================================
                document.add_page_break()
                p_anexo = document.add_paragraph()
                p_anexo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                r_anexo = p_anexo.add_run("ANEXO FOTOGRÁFICO DE SEGURIDAD")
                configurar_fuente_run(r_anexo, tamano=15, negrita=True, color=COLOR_BURDEO)

                document.add_paragraph()

                if fotos:
                    filas = (len(fotos) + 1) // 2
                    tabla_fotos = document.add_table(rows=filas, cols=2)
                    tabla_fotos.alignment = WD_TABLE_ALIGNMENT.CENTER
                    tabla_fotos.autofit = False
                    quitar_bordes_tabla(tabla_fotos)

                    for indice, foto_bytes in enumerate(fotos):
                        fila = indice // 2
                        columna = indice % 2
                        celda = tabla_fotos.cell(fila, columna)
                        celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                        p_img = celda.paragraphs[0]
                        p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        imagen_lista = preparar_foto_para_word(foto_bytes)
                        p_img.add_run().add_picture(imagen_lista, width=Cm(7.3))

                        p_cap = celda.add_paragraph()
                        p_cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        r_cap = p_cap.add_run(f"Fotografía {indice + 1}")
                        configurar_fuente_run(r_cap, tamano=8, negrita=True, color=RGBColor(90, 90, 90))
                else:
                    p_sin_fotos = document.add_paragraph()
                    p_sin_fotos.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    r_sin = p_sin_fotos.add_run("No se adjuntaron fotografías a este informe.")
                    configurar_fuente_run(r_sin, tamano=11, color=RGBColor(100, 100, 100))

                document.add_paragraph()
                p_emitido = document.add_paragraph()
                p_emitido.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                r_emitido = p_emitido.add_run(
                    "Emitido por:\n"
                    "Departamento de Informática\n"
                    "Liceo Bicentenario de Excelencia Colegio Antonio Varas"
                )
                configurar_fuente_run(r_emitido, tamano=10)

                docx_buf = BytesIO()
                document.save(docx_buf)
                docx_buf.seek(0)
                return docx_buf.read()

            except Exception as e:
                st.error(f"Error generando Word: {e}")
                return None

        tab_baja, tab_historial = st.tabs(["🆕 Procesar Nueva Baja", "📋 Ver Historial"])

        # Estados para conservar la descarga después del envío
        if "baja_lista" not in st.session_state:
            st.session_state.baja_lista = False
        if "baja_docx_data" not in st.session_state:
            st.session_state.baja_docx_data = None
        if "baja_filename" not in st.session_state:
            st.session_state.baja_filename = ""

        with tab_baja:
            try:
                res_data_raw = supabase.table("recursos").select("id, nombre").order("nombre").execute().data or []
                nombres_recursos = sorted({
                    texto_seguro(r.get("nombre"), "Sin nombre")
                    for r in res_data_raw
                })
            except Exception as e:
                st.error(f"No fue posible cargar los recursos: {e}")
                nombres_recursos = []

            if not nombres_recursos:
                st.warning("No hay recursos disponibles en la tabla 'recursos'.")
            else:
                st.markdown("""
                    <div style="
                        background:#EFF6FF;
                        border:1px solid #BFDBFE;
                        border-left:6px solid #1E3A8A;
                        border-radius:12px;
                        padding:14px 16px;
                        margin-bottom:16px;">
                        <b>📄 Informe institucional</b><br>
                        Completa los antecedentes, diagnóstico y detalle del equipo.
                        Puedes adjuntar varias fotografías JPG o PNG.
                    </div>
                """, unsafe_allow_html=True)

                with st.form("form_baja_profesional", clear_on_submit=False):
                    st.markdown("### 🏷️ Identificación del equipo")
                    col_a, col_b, col_c = st.columns(3)

                    with col_a:
                        recurso_cat_nom = st.selectbox("Equipo / categoría *", nombres_recursos)
                        cantidad_baja = st.number_input("Cantidad *", min_value=1, value=1, step=1)
                        num_serie = st.text_input("N.° de serie / inventario")

                    with col_b:
                        marca = st.text_input("Marca", placeholder="Ej. Sony")
                        modelo = st.text_input("Modelo", placeholder="Ej. E PZ 16–50 mm")
                        estado_equipo = st.text_input(
                            "Estado resumido",
                            placeholder="Ej. Falla mecánica / extraviado / irreparable"
                        )

                    with col_c:
                        materia = st.text_input("Materia", value="Baja técnica")
                        fecha_baja_ui = st.date_input(
                            "Fecha del informe *",
                            value=dt.date.today(),
                            format="DD/MM/YYYY"
                        )
                        tecnico = st.text_input(
                            "Técnico responsable *",
                            value=st.session_state.get("profesor_name") or ""
                        )

                    st.markdown("### 📝 Informe técnico")
                    antecedentes = st.text_area(
                        "1. Antecedentes y causas *",
                        height=160,
                        placeholder=(
                            "Describe cómo se recibió el equipo, desde cuándo está bajo custodia, "
                            "qué ocurrió y las causas conocidas de la baja."
                        )
                    )

                    diagnosis = st.text_area(
                        "2. Diagnóstico técnico *",
                        height=150,
                        placeholder="Describe la falla constatada y cómo afecta el funcionamiento."
                    )

                    recomendacion = st.text_area(
                        "Recomendación técnica",
                        height=90,
                        placeholder="Ej. Dar de baja, reemplazar, enviar a reciclaje electrónico..."
                    )

                    with st.expander("➕ Datos complementarios"):
                        col_d, col_e = st.columns(2)
                        ubicacion = col_d.text_input("Ubicación habitual")
                        fecha_adq = col_e.text_input("Fecha o año de adquisición")

                    uploaded_files = st.file_uploader(
                        "📷 Adjuntar fotografías",
                        type=["png", "jpg", "jpeg"],
                        accept_multiple_files=True,
                        help="Puedes seleccionar varias imágenes en una sola carga."
                    )

                    submit_baja = st.form_submit_button(
                        "🚫 Registrar Baja y Generar Informe Word",
                        type="primary",
                        use_container_width=True
                    )

                    if submit_baja:
                        faltantes = []
                        if not antecedentes.strip():
                            faltantes.append("Antecedentes y causas")
                        if not diagnosis.strip():
                            faltantes.append("Diagnóstico técnico")
                        if not tecnico.strip():
                            faltantes.append("Técnico responsable")

                        if faltantes:
                            st.warning("⚠️ Completa los campos obligatorios: " + ", ".join(faltantes) + ".")
                        else:
                            datos_bd = {
                                "recurso_nombre": recurso_cat_nom,
                                "marca": marca.strip(),
                                "modelo": modelo.strip(),
                                "cantidad_baja": int(cantidad_baja),
                                "serie": num_serie.strip(),
                                "diagnosis": diagnosis.strip(),
                                # Se conserva la compatibilidad con tu tabla actual:
                                # los antecedentes se almacenan en 'justificacion'.
                                "justificacion": antecedentes.strip(),
                                "recomendacion": recomendacion.strip(),
                                "tecnico_responsable": tecnico.strip(),
                                "fecha_baja": fecha_baja_ui.isoformat(),
                            }

                            datos_documento = {
                                **datos_bd,
                                "materia": materia.strip() or "Baja técnica",
                                "antecedentes": antecedentes.strip(),
                                "estado_equipo": estado_equipo.strip(),
                                "ubicacion": ubicacion.strip(),
                                "fecha_adquisicion": fecha_adq.strip(),
                            }

                            fotos_bytes = [archivo.getvalue() for archivo in (uploaded_files or [])]

                            try:
                                with st.spinner("Guardando la baja y preparando el informe..."):
                                    supabase.table("equipos").insert(datos_bd).execute()
                                    docx_bytes = generar_docx_baja(datos_documento, fotos_bytes)

                                if docx_bytes:
                                    fecha_nombre = fecha_baja_ui.strftime("%Y-%m-%d")
                                    st.session_state.baja_docx_data = docx_bytes
                                    st.session_state.baja_filename = (
                                        f"Informe_Baja_{nombre_archivo_seguro(recurso_cat_nom)}_{fecha_nombre}.docx"
                                    )
                                    st.session_state.baja_lista = True

                                st.success(
                                    f"✅ Baja registrada. Se incorporaron {len(fotos_bytes)} fotografía(s) al informe."
                                )
                                st.balloons()

                            except Exception as e:
                                st.error(f"Error al guardar la baja en Supabase: {e}")

                # Vista previa de imágenes, fuera del formulario
                if uploaded_files:
                    st.markdown("### 🖼️ Fotografías seleccionadas")
                    columnas_preview = st.columns(min(4, len(uploaded_files)))
                    for indice, archivo in enumerate(uploaded_files):
                        with columnas_preview[indice % len(columnas_preview)]:
                            st.image(
                                archivo.getvalue(),
                                caption=f"Foto {indice + 1}",
                                use_container_width=True
                            )

                if st.session_state.baja_lista and st.session_state.baja_docx_data:
                    st.markdown("### 📥 Informe listo")
                    st.download_button(
                        label="⬇️ Descargar Informe de Baja (.docx)",
                        data=st.session_state.baja_docx_data,
                        file_name=st.session_state.baja_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        type="primary"
                    )

                    if st.button("🧹 Limpiar y registrar otra baja", use_container_width=True):
                        st.session_state.baja_lista = False
                        st.session_state.baja_docx_data = None
                        st.session_state.baja_filename = ""
                        st.rerun()

        with tab_historial:
            st.subheader("📋 Historial de Bajas")
            st.write("Consulta y vuelve a generar los informes guardados en la tabla `equipos`.")

            try:
                bajas_db = (
                    supabase.table("equipos")
                    .select("*")
                    .order("fecha_baja", desc=True)
                    .execute()
                    .data or []
                )

                if bajas_db:
                    df_bajas = pd.DataFrame(bajas_db)
                    columnas_disponibles = [
                        col for col in [
                            "fecha_baja",
                            "recurso_nombre",
                            "cantidad_baja",
                            "marca",
                            "modelo",
                            "serie",
                            "tecnico_responsable",
                        ]
                        if col in df_bajas.columns
                    ]
                    df_mostrar = df_bajas[columnas_disponibles].copy()
                    df_mostrar.rename(columns={
                        "fecha_baja": "Fecha Baja",
                        "recurso_nombre": "Equipo / Recurso",
                        "cantidad_baja": "Cant.",
                        "marca": "Marca",
                        "modelo": "Modelo",
                        "serie": "N.° Serie",
                        "tecnico_responsable": "Técnico Responsable",
                    }, inplace=True)

                    st.dataframe(df_mostrar, use_container_width=True, hide_index=True)

                    st.info(
                        "ℹ️ Los informes regenerados desde el historial contienen los datos técnicos, "
                        "pero no las fotografías originales, porque tu tabla actual no almacena archivos."
                    )

                    st.markdown("### 🔍 Regenerar informe")
                    opciones_historial = [
                        (
                            f"{texto_seguro(b.get('fecha_baja'), 'Sin fecha')} - "
                            f"{texto_seguro(b.get('recurso_nombre'), 'Equipo')} "
                            f"(Serie: {texto_seguro(b.get('serie'), 'S/N')})"
                        )
                        for b in bajas_db
                    ]

                    seleccion_h = st.selectbox(
                        "Selecciona un registro:",
                        opciones_historial
                    )

                    if seleccion_h:
                        idx_sel = opciones_historial.index(seleccion_h)
                        datos_registro = bajas_db[idx_sel]
                        datos_registro["materia"] = "Baja técnica"
                        datos_registro["antecedentes"] = texto_seguro(
                            datos_registro.get("justificacion")
                        )
                        datos_registro["estado_equipo"] = texto_seguro(
                            datos_registro.get("diagnosis")
                        )

                        word_bytes_historial = generar_docx_baja(datos_registro, [])

                        if word_bytes_historial:
                            nombre_archivo_h = (
                                f"Informe_Baja_"
                                f"{nombre_archivo_seguro(datos_registro.get('recurso_nombre'))}_"
                                f"{texto_seguro(datos_registro.get('fecha_baja'), str(dt.date.today()))}.docx"
                            )
                            st.download_button(
                                label=f"⬇️ Descargar Word de {texto_seguro(datos_registro.get('recurso_nombre'), 'Equipo')}",
                                data=word_bytes_historial,
                                file_name=nombre_archivo_h,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"btn_hist_{datos_registro.get('id', idx_sel)}",
                                use_container_width=True
                            )
                else:
                    st.info("📂 No se han registrado bajas de equipos en el sistema aún.")

            except Exception as e:
                st.error(f"Error al cargar el historial desde Supabase: {e}")

    # ---------------------------------------------------------
    # MÓDULO 3: GENERADOR QR MAESTRO
    # ---------------------------------------------------------
    elif modulo_tec == "📋 Generador QR":
        st.subheader("📱 Generador de QR Maestro de Reportes")
        st.markdown("""
        Este código QR es **único para todo el colegio**. 
        Imprímelo y pégalo en lugares estratégicos (Sala de profesores, Inspectoría, pasillos). 
        Al escanearlo, el usuario abrirá el **Formulario Centralizado de Reportes** donde podrá elegir qué equipo o sala está fallando.
        """)
        
        # URL de tu app para la página de reporte maestro (Sin ID)
        base_url = "https://enlaces.streamlit.app/" 
        final_url = f"{base_url}?page=reporte"
        
        col1, col2 = st.columns([1, 2])
        with col1:
            qr = qrcode.make(final_url)
            buf = BytesIO()
            qr.save(buf, format="PNG")
            
            st.image(buf.getvalue(), width=250, caption="Código QR Maestro")
            
        with col2:
            st.info("💡 **Instrucciones de uso:**\n1. Descarga la imagen.\n2. Imprímela en tamaño carta o póster.\n3. Pégala en el colegio.\n\nOlvídate de generar un QR por cada equipo nuevo que compres.")
            st.download_button("⬇️ Descargar Código QR Maestro", data=buf.getvalue(), file_name="QR_Maestro_Fallas.png", mime="image/png", type="primary")
            st.code(final_url, language="html")
# ------------------------------------------------------------------
# INVENTARIO TÉCNICO
# ------------------------------------------------------------------
if page == "Inventario":
    st.title("💻 Inventario Técnico")
    st.caption("Pasaporte digital de equipos, garantías, ubicación y estado.")

    tab_nuevo_inv, tab_listado_inv = st.tabs(["➕ Registrar equipo", "📋 Inventario"])

    with tab_nuevo_inv:
        with st.form("form_inventario"):
            c1, c2, c3 = st.columns(3)
            with c1:
                nombre_inv = st.text_input("Nombre del equipo *")
                categoria_inv = st.selectbox("Categoría", ["Computador", "Notebook", "Proyector", "Impresora", "Tablet", "Cámara", "Audio", "Red", "Otro"])
                codigo_inv = st.text_input("Código patrimonial")
                serie_inv = st.text_input("Número de serie")
            with c2:
                marca_inv = st.text_input("Marca")
                modelo_inv = st.text_input("Modelo")
                ubicacion_inv = st.text_input("Ubicación")
                responsable_inv = st.text_input("Responsable")
            with c3:
                fecha_compra_inv = st.date_input("Fecha de compra", value=dt.date.today(), format="DD/MM/YYYY")
                garantia_inv = st.date_input("Fin de garantía", value=dt.date.today(), format="DD/MM/YYYY")
                valor_inv = st.number_input("Valor de compra (CLP)", min_value=0, value=0, step=1000)
                estado_inv = st.selectbox("Estado", ["Operativo", "En revisión", "En reparación", "Prestado", "Dado de baja"])

            observaciones_inv = st.text_area("Observaciones")
            if st.form_submit_button("💾 Guardar equipo", type="primary", use_container_width=True):
                if not nombre_inv.strip():
                    st.warning("El nombre del equipo es obligatorio.")
                else:
                    datos_inv = {
                        "nombre": nombre_inv.strip(),
                        "categoria": categoria_inv,
                        "codigo_patrimonial": codigo_inv.strip(),
                        "serie": serie_inv.strip(),
                        "marca": marca_inv.strip(),
                        "modelo": modelo_inv.strip(),
                        "ubicacion": ubicacion_inv.strip(),
                        "responsable": responsable_inv.strip(),
                        "fecha_compra": fecha_compra_inv.isoformat(),
                        "garantia_hasta": garantia_inv.isoformat(),
                        "valor_compra": int(valor_inv),
                        "estado": estado_inv,
                        "observaciones": observaciones_inv.strip(),
                    }
                    try:
                        res = supabase.table("inventario").insert(datos_inv).execute()
                        registrar_auditoria("crear", "inventario", detalle=datos_inv)
                        st.success("✅ Equipo incorporado al inventario.")
                        st.balloons()
                    except Exception as e:
                        registrar_error("crear_inventario", e)
                        st.error(f"No fue posible guardar el equipo: {e}")

    with tab_listado_inv:
        try:
            inventario_db = select_paginado("inventario", "*", orden="nombre")
        except Exception as e:
            registrar_error("listar_inventario", e)
            inventario_db = []

        if inventario_db:
            df_inv = pd.DataFrame(inventario_db)
            filtro_estado_inv = st.multiselect("Filtrar por estado", sorted(df_inv["estado"].dropna().unique().tolist()))
            filtro_texto_inv = st.text_input("🔎 Buscar por nombre, serie, código o ubicación")
            df_filtrado_inv = df_inv.copy()
            if filtro_estado_inv:
                df_filtrado_inv = df_filtrado_inv[df_filtrado_inv["estado"].isin(filtro_estado_inv)]
            if filtro_texto_inv.strip():
                patron = filtro_texto_inv.strip().lower()
                columnas_busqueda = [c for c in ["nombre", "serie", "codigo_patrimonial", "ubicacion", "marca", "modelo"] if c in df_filtrado_inv.columns]
                mascara = pd.Series(False, index=df_filtrado_inv.index)
                for col in columnas_busqueda:
                    mascara |= df_filtrado_inv[col].fillna("").astype(str).str.lower().str.contains(patron, regex=False)
                df_filtrado_inv = df_filtrado_inv[mascara]

            columnas_vista = [c for c in ["nombre", "categoria", "marca", "modelo", "serie", "ubicacion", "responsable", "estado", "garantia_hasta"] if c in df_filtrado_inv.columns]
            st.dataframe(df_filtrado_inv[columnas_vista], use_container_width=True, hide_index=True)

            st.markdown("### 🪪 Pasaporte digital")
            opciones_inv = {
                f"{r.get('nombre', 'Equipo')} · {r.get('serie') or 'S/N'}": r
                for r in inventario_db
            }
            sel_inv = st.selectbox("Selecciona un equipo", list(opciones_inv.keys()))
            equipo_inv = opciones_inv[sel_inv]
            with st.container(border=True):
                c1, c2 = st.columns([2, 1])
                with c1:
                    st.markdown(f"## {equipo_inv.get('nombre', 'Equipo')}")
                    st.write(f"**Categoría:** {equipo_inv.get('categoria') or '—'}")
                    st.write(f"**Marca / Modelo:** {equipo_inv.get('marca') or '—'} / {equipo_inv.get('modelo') or '—'}")
                    st.write(f"**Serie:** {equipo_inv.get('serie') or 'S/N'}")
                    st.write(f"**Ubicación:** {equipo_inv.get('ubicacion') or '—'}")
                    st.write(f"**Responsable:** {equipo_inv.get('responsable') or '—'}")
                with c2:
                    st.metric("Estado", equipo_inv.get("estado") or "—")
                    st.metric("Garantía hasta", equipo_inv.get("garantia_hasta") or "—")
                    st.metric("Valor", f"${int(equipo_inv.get('valor_compra') or 0):,}".replace(",", "."))
        else:
            st.info("No hay equipos en la tabla `inventario`.")

# ------------------------------------------------------------------
# MANTENCIÓN PREVENTIVA
# ------------------------------------------------------------------
elif page == "Mantención preventiva":
    st.title("🧰 Mantención Preventiva")
    st.caption("Programa, ejecuta y documenta revisiones periódicas.")

    try:
        inventario_mp = select_paginado("inventario", "id,nombre,serie,estado", orden="nombre")
    except Exception as e:
        registrar_error("inventario_preventivo", e)
        inventario_mp = []

    tab_plan_mp, tab_hist_mp = st.tabs(["📅 Programar", "📋 Plan e historial"])

    with tab_plan_mp:
        if not inventario_mp:
            st.warning("Primero registra equipos en el módulo Inventario.")
        else:
            opciones_mp = {
                f"{r['nombre']} · {r.get('serie') or 'S/N'}": r["id"]
                for r in inventario_mp
            }
            with st.form("form_mant_preventiva"):
                equipo_label_mp = st.selectbox("Equipo", list(opciones_mp.keys()))
                c1, c2, c3 = st.columns(3)
                fecha_prog_mp = c1.date_input("Fecha programada", value=dt.date.today(), format="DD/MM/YYYY")
                frecuencia_mp = c2.selectbox("Frecuencia", ["Única", "Mensual", "Trimestral", "Semestral", "Anual"])
                prioridad_mp = c3.selectbox("Prioridad", ["Baja", "Media", "Alta"])
                tarea_mp = st.text_area("Tareas a realizar *", placeholder="Limpieza, actualización, revisión de cables, pruebas...")
                responsable_mp = st.text_input("Responsable", value=st.session_state.get("profesor_name") or "")
                if st.form_submit_button("💾 Programar mantención", type="primary", use_container_width=True):
                    if not tarea_mp.strip():
                        st.warning("Describe las tareas a realizar.")
                    else:
                        datos_mp = {
                            "inventario_id": opciones_mp[equipo_label_mp],
                            "fecha_programada": fecha_prog_mp.isoformat(),
                            "frecuencia": frecuencia_mp,
                            "prioridad": prioridad_mp,
                            "tareas": tarea_mp.strip(),
                            "responsable": responsable_mp.strip(),
                            "estado": "Pendiente",
                        }
                        try:
                            supabase.table("mantenciones_preventivas").insert(datos_mp).execute()
                            registrar_auditoria("crear", "mantenciones_preventivas", detalle=datos_mp)
                            st.success("✅ Mantención programada.")
                        except Exception as e:
                            registrar_error("crear_mantencion_preventiva", e)
                            st.error(f"No fue posible programar: {e}")

    with tab_hist_mp:
        try:
            planes_mp = select_paginado(
                "mantenciones_preventivas",
                "*, inventario(nombre,serie)",
                orden="fecha_programada"
            )
        except Exception as e:
            registrar_error("listar_mantenciones_preventivas", e)
            planes_mp = []

        if planes_mp:
            for mp in planes_mp:
                inv = mp.get("inventario") or {}
                etiqueta = f"{mp.get('fecha_programada')} · {inv.get('nombre', 'Equipo')} · {mp.get('estado')}"
                with st.expander(etiqueta):
                    st.write(f"**Tareas:** {mp.get('tareas')}")
                    st.write(f"**Responsable:** {mp.get('responsable') or '—'}")
                    st.write(f"**Frecuencia:** {mp.get('frecuencia') or '—'}")
                    st.write(f"**Prioridad:** {mp.get('prioridad') or '—'}")
                    nuevo_estado_mp = st.selectbox(
                        "Estado",
                        ["Pendiente", "En proceso", "Completada", "Postergada"],
                        index=["Pendiente", "En proceso", "Completada", "Postergada"].index(mp.get("estado") if mp.get("estado") in ["Pendiente", "En proceso", "Completada", "Postergada"] else "Pendiente"),
                        key=f"estado_mp_{mp.get('id')}"
                    )
                    resultado_mp = st.text_area("Resultado / observaciones", value=mp.get("resultado") or "", key=f"resultado_mp_{mp.get('id')}")
                    if st.button("💾 Actualizar", key=f"guardar_mp_{mp.get('id')}", use_container_width=True):
                        cambios_mp = {
                            "estado": nuevo_estado_mp,
                            "resultado": resultado_mp.strip(),
                            "fecha_realizada": dt.date.today().isoformat() if nuevo_estado_mp == "Completada" else mp.get("fecha_realizada"),
                        }
                        supabase.table("mantenciones_preventivas").update(cambios_mp).eq("id", mp.get("id")).execute()
                        registrar_auditoria("actualizar", "mantenciones_preventivas", mp.get("id"), cambios_mp)
                        st.success("Actualizado.")
                        st.rerun()
        else:
            st.info("No hay mantenciones preventivas programadas.")

# ------------------------------------------------------------------
# AUDITORÍA
# ------------------------------------------------------------------
elif page == "Auditoría":
    st.title("🧾 Registro de Auditoría")
    st.caption("Trazabilidad de acciones administrativas y técnicas.")
    try:
        auditoria_db = select_paginado("auditoria", "*", orden="fecha", desc=True)
    except Exception as e:
        registrar_error("listar_auditoria", e)
        auditoria_db = []

    if auditoria_db:
        df_aud = pd.DataFrame(auditoria_db)
        c1, c2 = st.columns(2)
        usuarios_aud = sorted(df_aud["usuario"].dropna().unique().tolist()) if "usuario" in df_aud.columns else []
        modulos_aud = sorted(df_aud["modulo"].dropna().unique().tolist()) if "modulo" in df_aud.columns else []
        filtro_usuario_aud = c1.multiselect("Usuario", usuarios_aud)
        filtro_modulo_aud = c2.multiselect("Módulo", modulos_aud)
        if filtro_usuario_aud:
            df_aud = df_aud[df_aud["usuario"].isin(filtro_usuario_aud)]
        if filtro_modulo_aud:
            df_aud = df_aud[df_aud["modulo"].isin(filtro_modulo_aud)]
        st.dataframe(df_aud, use_container_width=True, hide_index=True)
    else:
        st.info("No hay registros o falta ejecutar la migración de auditoría.")


elif page == "Diplomas":
    from PIL import Image, ImageDraw, ImageFont, ImageOps, ImageChops
    from email.message import EmailMessage
    from email.utils import formataddr
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.lib.colors import HexColor
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfbase.pdfmetrics import stringWidth
    import fitz
    import unicodedata
    import ssl

    st.title("🎓 Diplomas Digitales CAV · V16.3 · Render digital definitivo")

    st.caption(
        "Reconocimientos digitales con registro oficial, enlace personal, "
        "PDF de respaldo e impacto ambiental acumulado."
    )

    try:
        filtro_creador = (
            None
            if st.session_state.get("role") == "admin"
            else st.session_state.get("profesor_name")
        )
        registros_diplomas = cargar_diplomas_registrados(filtro_creador)
        urls_corregidas = reparar_urls_publicas_diplomas(registros_diplomas)
        if urls_corregidas:
            st.info(
                f"🔗 Se corrigieron automáticamente {urls_corregidas} "
                "enlaces digitales antiguos."
            )
    except Exception as error:
        registros_diplomas = []
        st.warning(
            "El registro oficial todavía no está disponible. "
            "Ejecuta la migración SQL V15 en Supabase."
        )
        registrar_error("dashboard_diplomas", error)

    total_diplomas = len(registros_diplomas)
    enviados_diplomas = sum(
        1 for d in registros_diplomas if d.get("estado") == "ENVIADO"
    )
    pendientes_diplomas = sum(
        1 for d in registros_diplomas if d.get("estado") == "GENERADO"
    )
    errores_diplomas = sum(
        1 for d in registros_diplomas if d.get("estado") == "ERROR"
    )

    hojas_opalina = enviados_diplomas
    hojas_normales = enviados_diplomas * EQUIVALENCIA_HOJAS_NORMALES
    toner_min = enviados_diplomas * TONER_MIN_GRAMOS_POR_DIPLOMA
    toner_max = enviados_diplomas * TONER_MAX_GRAMOS_POR_DIPLOMA

    if st.session_state.get("role") == "admin":
        st.markdown("## 📊 Resumen institucional")
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Diplomas registrados", total_diplomas)
        m2.metric("Enviados", enviados_diplomas)
        m3.metric("Pendientes", pendientes_diplomas)
        m4.metric("Errores", errores_diplomas)

        e1, e2, e3 = st.columns(3)
        e1.metric("Hojas opalina evitadas", f"{hojas_opalina:,.0f}")
        e2.metric("Hojas normales equivalentes", f"{hojas_normales:,.1f}")
        e3.metric(
            "Tóner estimado evitado",
            f"{toner_min:,.1f}–{toner_max:,.1f} g",
        )

        with st.expander("📚 Historial y distribución por área", expanded=False):
            if registros_diplomas:
                df_historial = pd.DataFrame(registros_diplomas)
                columnas_historial = [
                    c for c in [
                        "codigo", "nombre", "curso", "area", "estado",
                        "creado_por", "creado_en", "enviado_en"
                    ] if c in df_historial.columns
                ]
                st.dataframe(
                    df_historial[columnas_historial],
                    use_container_width=True,
                    hide_index=True,
                )

                if "area" in df_historial.columns:
                    df_area = (
                        df_historial[df_historial["estado"] == "ENVIADO"]
                        .groupby("area")
                        .size()
                        .reset_index(name="Diplomas")
                        .sort_values("Diplomas", ascending=False)
                    )
                    if not df_area.empty:
                        fig_area = px.bar(
                            df_area,
                            x="area",
                            y="Diplomas",
                            title="Diplomas enviados por área",
                        )
                        st.plotly_chart(fig_area, use_container_width=True)
            else:
                st.info("Todavía no existen diplomas registrados.")
    else:
        st.markdown("## 🌱 Mis reconocimientos digitales")
        p1, p2, p3 = st.columns(3)
        p1.metric("Generados por mí", total_diplomas)
        p2.metric("Enviados", enviados_diplomas)
        p3.metric("Ahorro confirmado", f"{hojas_opalina} hojas")

    st.markdown("---")
    st.markdown("## ✨ Crear diploma digital")

    st.caption(
        "Genera diplomas digitales con mejor jerarquía visual, versión pública animada, "
        "elementos gráficos por área y respaldo PDF ordenado."
    )

    BASE_DIPLOMAS = Path(__file__).parent
    RUTA_LOGO_DIPLOMA = BASE_DIPLOMAS / "logocav.png"
    RUTA_FIRMA_DIRECTOR = BASE_DIPLOMAS / "firma_director.png"

    ESTILOS_DIPLOMA = {
        "Institucional general": {
            "primario": "#800020",
            "secundario": "#C7A44A",
            "fondo": "#FFFDF8",
            "suave": "#F5E9E2",
            "acento": "#5E0017",
            "motivo": "institucional",
        },
        "Lenguaje": {
            "primario": "#7A1734",
            "secundario": "#D7A642",
            "fondo": "#FFF9F1",
            "suave": "#F6E5DA",
            "acento": "#4F0E22",
            "motivo": "lenguaje",
        },
        "Matemática": {
            "primario": "#173B73",
            "secundario": "#D3A62E",
            "fondo": "#F8FBFF",
            "suave": "#E7EFFA",
            "acento": "#0B2851",
            "motivo": "matematica",
        },
        "Ciencias": {
            "primario": "#126B70",
            "secundario": "#A8B83D",
            "fondo": "#F7FFFC",
            "suave": "#DFF3EC",
            "acento": "#084C50",
            "motivo": "ciencias",
        },
        "Artes": {
            "primario": "#A23B72",
            "secundario": "#E39A37",
            "fondo": "#FFF9FC",
            "suave": "#F8E4EF",
            "acento": "#6F1E4A",
            "motivo": "artes",
        },
        "Música": {
            "primario": "#5D3A8E",
            "secundario": "#D6A83B",
            "fondo": "#FCF9FF",
            "suave": "#EEE6FA",
            "acento": "#3A1E65",
            "motivo": "musica",
        },
        "Tecnología": {
            "primario": "#1769AA",
            "secundario": "#31A8A0",
            "fondo": "#F7FBFF",
            "suave": "#E3F1FA",
            "acento": "#0E4778",
            "motivo": "tecnologia",
        },
        "Educación Física": {
            "primario": "#237A57",
            "secundario": "#E1A833",
            "fondo": "#F8FFF9",
            "suave": "#E2F3E8",
            "acento": "#14533B",
            "motivo": "educacion_fisica",
        },
        "Reconocimiento especial": {
            "primario": "#6F1027",
            "secundario": "#B9912E",
            "fondo": "#FFFDF5",
            "suave": "#F3E8D0",
            "acento": "#430817",
            "motivo": "especial",
        },
    }

    TEXTOS_RECONOCIMIENTO = {
        "Diploma de reconocimiento": (
            "Por su destacada participación, compromiso y dedicación, "
            "demostrando valores que contribuyen positivamente a nuestra comunidad educativa."
        ),
        "Mérito académico": (
            "Por su sobresaliente desempeño académico, responsabilidad y constancia "
            "durante el proceso educativo."
        ),
        "Participación destacada": (
            "Por su valiosa y entusiasta participación, demostrando compromiso, "
            "responsabilidad y espíritu de colaboración."
        ),
        "Convivencia escolar": (
            "Por promover una sana convivencia, el respeto y el compañerismo, "
            "siendo un aporte significativo para su comunidad escolar."
        ),
        "Superación personal": (
            "Por su esfuerzo, perseverancia y permanente espíritu de superación, "
            "constituyéndose en un ejemplo para la comunidad educativa."
        ),
        "Texto personalizado": "",
    }

    def color_rgb(hex_color):
        valor = hex_color.lstrip("#")
        return tuple(int(valor[i:i + 2], 16) for i in (0, 2, 4))

    def archivo_seguro_diploma(texto):
        texto = unicodedata.normalize("NFKD", str(texto or "Diploma"))
        texto = "".join(c for c in texto if not unicodedata.combining(c))
        texto = re.sub(r"[^A-Za-z0-9_-]+", "_", texto)
        return texto.strip("_") or "Diploma"

    def buscar_fuente_diploma(negrita=False, cursiva=False):
        # Priorizamos DejaVu porque soporta bien tildes, ñ y símbolos comunes.
        if negrita and cursiva:
            candidatos = [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf",
                "DejaVuSans-BoldOblique.ttf",
            ]
        elif negrita:
            candidatos = [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
                "DejaVuSans-Bold.ttf",
            ]
        elif cursiva:
            candidatos = [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf",
                "DejaVuSans-Oblique.ttf",
            ]
        else:
            candidatos = [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                "DejaVuSans.ttf",
            ]

        for fuente in candidatos:
            try:
                ImageFont.truetype(fuente, 40)
                return fuente
            except Exception:
                pass
        return None

    def fuente_diploma(tamano, negrita=False, cursiva=False):
        ruta = buscar_fuente_diploma(negrita=negrita, cursiva=cursiva)
        if ruta:
            try:
                return ImageFont.truetype(ruta, tamano)
            except Exception:
                pass

        # Fallback robusto para evitar texto diminuto.
        try:
            return ImageFont.load_default(size=tamano)
        except TypeError:
            return ImageFont.load_default()


    def limpiar_texto_diploma(texto, mayusculas=False):
        if texto is None:
            texto = ""
        texto = str(texto)

        # Corrige secuencias frecuentes de mojibake y normaliza Unicode.
        reemplazos = {
            "Ã¡": "á", "Ã©": "é", "Ã­": "í", "Ã³": "ó", "Ãº": "ú",
            "Ã": "Á", "Ã‰": "É", "Ã": "Í", "Ã“": "Ó", "Ãš": "Ú",
            "Ã±": "ñ", "Ã‘": "Ñ", "Â°": "°", "Âº": "º", "â€“": "–",
            "â€”": "—", "â€˜": "‘", "â€™": "’", "â€œ": "“", "â€": "”",
            "�": "",
        }
        for a, b in reemplazos.items():
            texto = texto.replace(a, b)

        texto = unicodedata.normalize("NFKC", texto)
        texto = texto.replace("\r", " ").replace("\n", " ")
        texto = " ".join(texto.split())

        if mayusculas:
            texto = texto.upper()
        return texto

    def medir_texto(draw, texto, font):
        caja = draw.textbbox((0, 0), texto, font=font)
        return caja[2] - caja[0], caja[3] - caja[1]

    def ajustar_fuente(draw, texto, max_ancho, tamano_max, tamano_min=28, negrita=False):
        for tamano in range(tamano_max, tamano_min - 1, -2):
            font = fuente_diploma(tamano, negrita=negrita)
            ancho, _ = medir_texto(draw, texto, font)
            if ancho <= max_ancho:
                return font
        return fuente_diploma(tamano_min, negrita=negrita)

    def envolver_por_ancho(draw, texto, font, max_ancho):
        parrafos = str(texto or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
        lineas_finales = []

        for parrafo in parrafos:
            parrafo = parrafo.strip()
            if not parrafo:
                lineas_finales.append("")
                continue

            palabras = parrafo.split()
            linea = ""

            for palabra in palabras:
                prueba = palabra if not linea else f"{linea} {palabra}"
                ancho, _ = medir_texto(draw, prueba, font)
                if ancho <= max_ancho:
                    linea = prueba
                else:
                    if linea:
                        lineas_finales.append(linea)
                    linea = palabra

            if linea:
                lineas_finales.append(linea)

        return lineas_finales

    def dibujar_texto_centrado(draw, texto, y, font, fill, ancho_total):
        ancho, alto = medir_texto(draw, texto, font)
        draw.text(
            ((ancho_total - ancho) / 2, y),
            texto,
            font=font,
            fill=fill,
        )
        return y + alto

    def dibujar_multilinea_centrada(
        draw,
        texto,
        y,
        font,
        fill,
        ancho_total,
        max_ancho,
        separacion=18,
    ):
        lineas = envolver_por_ancho(draw, texto, font, max_ancho)
        y_actual = y

        for linea in lineas:
            if not linea:
                y_actual += separacion
                continue

            ancho, alto = medir_texto(draw, linea, font)
            draw.text(
                ((ancho_total - ancho) / 2, y_actual),
                linea,
                font=font,
                fill=fill,
            )
            y_actual += alto + separacion

        return y_actual

    def abrir_imagen_diploma(origen):
        if origen is None:
            return None
        try:
            if isinstance(origen, (str, Path)):
                imagen = Image.open(origen)
            else:
                imagen = Image.open(BytesIO(origen))
            return ImageOps.exif_transpose(imagen).convert("RGBA")
        except Exception:
            return None

    def eliminar_fondo_blanco_firma(imagen):
        if imagen is None:
            return None

        rgba = imagen.convert("RGBA")
        pixeles = rgba.load()

        for y in range(rgba.height):
            for x in range(rgba.width):
                r, g, b, a = pixeles[x, y]
                promedio = (r + g + b) / 3
                if promedio > 247 and max(r, g, b) - min(r, g, b) < 12:
                    pixeles[x, y] = (255, 255, 255, 0)

        return rgba

    def pegar_imagen_contenida(canvas, imagen, caja, limpiar_blanco=False):
        if imagen is None:
            return

        if limpiar_blanco:
            imagen = eliminar_fondo_blanco_firma(imagen)

        x1, y1, x2, y2 = caja
        max_w = max(1, x2 - x1)
        max_h = max(1, y2 - y1)

        copia = imagen.copy()
        copia.thumbnail((max_w, max_h), Image.Resampling.LANCZOS)

        x = x1 + (max_w - copia.width) // 2
        y = y1 + (max_h - copia.height) // 2
        canvas.alpha_composite(copia, (x, y))

    def dibujar_decoracion_area(draw, estilo, w, h):
        # Decoración minimalista y limpia.
        primario = color_rgb(estilo["primario"])
        secundario = color_rgb(estilo["secundario"])
        suave = color_rgb(estilo["suave"])
        motivo = estilo["motivo"]

        draw.line((720, 700, w - 720, 700), fill=secundario, width=4)
        draw.line((820, 1790, w - 820, 1790), fill=suave, width=4)

        if motivo == "matematica":
            for x in (860, w - 860):
                draw.ellipse((x - 40, 1715, x + 40, 1795), outline=secundario, width=4)
                draw.line((x - 28, 1755, x + 28, 1755), fill=primario, width=4)
                draw.line((x, 1727, x, 1783), fill=primario, width=4)

        elif motivo == "ciencias":
            for x in (860, w - 860):
                p1, p2, p3 = (x - 34, 1785), (x + 28, 1738), (x + 48, 1792)
                for px, py in (p1, p2, p3):
                    draw.ellipse((px - 7, py - 7, px + 7, py + 7), fill=secundario)
                draw.line((*p1, *p2), fill=primario, width=3)
                draw.line((*p2, *p3), fill=primario, width=3)
                draw.line((*p3, *p1), fill=primario, width=3)

        elif motivo == "musica":
            for x in (860, w - 860):
                draw.line((x, 1718, x, 1792), fill=primario, width=5)
                draw.line((x, 1718, x + 45, 1700), fill=primario, width=5)
                draw.ellipse((x - 22, 1774, x + 8, 1804), fill=secundario)
                draw.ellipse((x + 26, 1760, x + 56, 1790), fill=secundario)

        elif motivo == "tecnologia":
            for x in (790, w - 900):
                draw.rounded_rectangle((x, 1712, x + 86, 1794), radius=12, outline=primario, width=3)
                draw.line((x + 14, 1735, x + 72, 1735), fill=secundario, width=3)
                draw.line((x + 14, 1758, x + 72, 1758), fill=secundario, width=3)
                draw.line((x + 14, 1780, x + 48, 1780), fill=secundario, width=3)

        elif motivo == "educacion_fisica":
            for x in (860, w - 860):
                draw.arc((x - 40, 1715, x + 40, 1795), 20, 340, fill=primario, width=4)
                draw.arc((x - 24, 1731, x + 24, 1779), 210, 520, fill=secundario, width=3)

        elif motivo == "artes":
            for x in (860, w - 860):
                draw.ellipse((x - 36, 1720, x + 36, 1792), outline=primario, width=4)
                draw.ellipse((x - 20, 1736, x + 20, 1776), outline=secundario, width=3)

        elif motivo == "lenguaje":
            for x in (800, w - 910):
                draw.rounded_rectangle((x, 1712, x + 86, 1796), radius=12, outline=primario, width=3)
                draw.line((x + 43, 1718, x + 43, 1790), fill=secundario, width=3)
                for y in (1740, 1762, 1783):
                    draw.line((x + 12, y, x + 31, y), fill=suave, width=4)
                    draw.line((x + 55, y, x + 74, y), fill=suave, width=4)

    DIPLOMA_RENDER_VERSION = "V16.3-STHTML-INSTITUTIONAL-2026-08-03"

    def preparar_imagen_para_pdf(origen, quitar_blanco=False):
        """Convierte logo o firma a PNG en memoria, recortado y transparente."""
        imagen = abrir_imagen_diploma(origen)
        if imagen is None:
            return None

        imagen = imagen.convert("RGBA")

        if quitar_blanco:
            datos = []
            for r, g, b, a in imagen.getdata():
                if r > 244 and g > 244 and b > 244:
                    datos.append((255, 255, 255, 0))
                else:
                    datos.append((r, g, b, a))
            imagen.putdata(datos)

        alpha = imagen.getchannel("A")
        bbox = alpha.getbbox()
        if bbox:
            imagen = imagen.crop(bbox)

        salida = BytesIO()
        imagen.save(salida, format="PNG", optimize=True)
        return salida.getvalue()

    def ajustar_tamano_pdf(texto, fuente, maximo, minimo, ancho_maximo):
        texto = limpiar_texto_diploma(texto)
        tamano = float(maximo)
        while tamano > minimo and stringWidth(texto, fuente, tamano) > ancho_maximo:
            tamano -= 1
        return max(tamano, minimo)

    def envolver_texto_pdf(texto, fuente, tamano, ancho_maximo):
        texto = limpiar_texto_diploma(texto)
        palabras = texto.split()
        lineas = []
        linea = ""

        for palabra in palabras:
            prueba = palabra if not linea else f"{linea} {palabra}"
            if stringWidth(prueba, fuente, tamano) <= ancho_maximo:
                linea = prueba
            else:
                if linea:
                    lineas.append(linea)
                linea = palabra

        if linea:
            lineas.append(linea)

        return lineas or [""]

    def dibujar_texto_centrado_pdf(pdf, texto, y, fuente, tamano, color):
        texto = limpiar_texto_diploma(texto)
        pdf.setFillColor(color)
        pdf.setFont(fuente, tamano)
        pdf.drawCentredString(396, y, texto)

    def dibujar_imagen_contenida_pdf(pdf, imagen_bytes, x, y, ancho, alto):
        if not imagen_bytes:
            return

        imagen = Image.open(BytesIO(imagen_bytes)).convert("RGBA")
        iw, ih = imagen.size
        proporcion = min(ancho / iw, alto / ih)
        nuevo_ancho = iw * proporcion
        nuevo_alto = ih * proporcion
        pos_x = x + (ancho - nuevo_ancho) / 2
        pos_y = y + (alto - nuevo_alto) / 2

        pdf.drawImage(
            ImageReader(BytesIO(imagen_bytes)),
            pos_x,
            pos_y,
            width=nuevo_ancho,
            height=nuevo_alto,
            preserveAspectRatio=True,
            mask="auto",
        )

    def crear_qr_diploma(url):
        if not url:
            return None
        qr = qrcode.QRCode(
            version=None,
            error_correction=qrcode.constants.ERROR_CORRECT_M,
            box_size=8,
            border=2,
        )
        qr.add_data(url)
        qr.make(fit=True)
        imagen = qr.make_image(fill_color="#800020", back_color="white").convert("RGB")
        salida = BytesIO()
        imagen.save(salida, format="PNG", optimize=True)
        return salida.getvalue()


    def dibujar_estrella_pdf(pdf, cx, cy, radio, color):
        import math
        puntos = []
        for i in range(10):
            angulo = math.radians(90 + i * 36)
            r = radio if i % 2 == 0 else radio * 0.42
            puntos.append(
                (cx + r * math.cos(angulo), cy + r * math.sin(angulo))
            )
        ruta = pdf.beginPath()
        ruta.moveTo(*puntos[0])
        for punto in puntos[1:]:
            ruta.lineTo(*punto)
        ruta.close()
        pdf.setFillColor(color)
        pdf.setStrokeColor(color)
        pdf.drawPath(ruta, fill=1, stroke=0)

    def dibujar_motivo_area_vectorial(pdf, tema, cx, cy, escala, primario, secundario):
        """
        Dibuja iconografía vectorial relacionada con el área.
        No usa emojis, porque Helvetica no los representa y producía cuadrados.
        """
        tema_normalizado = normalizar_texto_basico(tema)
        pdf.saveState()
        pdf.setLineCap(1)
        pdf.setLineJoin(1)
        pdf.setStrokeColor(primario)
        pdf.setFillColor(secundario)
        pdf.setLineWidth(max(0.8, 1.3 * escala))

        if "arte" in tema_normalizado:
            # Paleta y pincel.
            pdf.ellipse(
                cx - 20 * escala, cy - 15 * escala,
                cx + 19 * escala, cy + 16 * escala,
                stroke=1, fill=0,
            )
            pdf.circle(cx + 10 * escala, cy - 4 * escala, 6 * escala, stroke=1, fill=0)
            for dx, dy in [(-11, 6), (-2, 11), (8, 8), (-12, -5)]:
                pdf.circle(cx + dx * escala, cy + dy * escala, 2.6 * escala, stroke=1, fill=1)
            pdf.setStrokeColor(secundario)
            pdf.setLineWidth(3 * escala)
            pdf.line(
                cx + 18 * escala, cy + 17 * escala,
                cx + 31 * escala, cy + 31 * escala,
            )
            pdf.setFillColor(primario)
            dibujar_estrella_pdf(pdf, cx - 26 * escala, cy + 24 * escala, 5 * escala, secundario)

        elif "lenguaje" in tema_normalizado or "debate" in tema_normalizado:
            # Libro abierto y pluma.
            pdf.roundRect(
                cx - 25 * escala, cy - 13 * escala,
                24 * escala, 27 * escala, 3 * escala,
                fill=0, stroke=1,
            )
            pdf.roundRect(
                cx + 1 * escala, cy - 13 * escala,
                24 * escala, 27 * escala, 3 * escala,
                fill=0, stroke=1,
            )
            pdf.line(cx, cy - 13 * escala, cx, cy + 14 * escala)
            for dy in (-6, 1, 8):
                pdf.line(cx - 20 * escala, cy + dy * escala, cx - 6 * escala, cy + dy * escala)
                pdf.line(cx + 6 * escala, cy + dy * escala, cx + 20 * escala, cy + dy * escala)
            pdf.setStrokeColor(secundario)
            pdf.setLineWidth(2 * escala)
            pdf.line(cx + 21 * escala, cy + 18 * escala, cx + 31 * escala, cy + 31 * escala)

        elif "matemat" in tema_normalizado:
            # Geometría y operaciones.
            pdf.circle(cx - 13 * escala, cy + 4 * escala, 11 * escala, stroke=1, fill=0)
            ruta = pdf.beginPath()
            ruta.moveTo(cx + 2 * escala, cy - 12 * escala)
            ruta.lineTo(cx + 26 * escala, cy - 12 * escala)
            ruta.lineTo(cx + 14 * escala, cy + 13 * escala)
            ruta.close()
            pdf.drawPath(ruta, fill=0, stroke=1)
            pdf.setStrokeColor(secundario)
            pdf.line(cx - 25 * escala, cy - 20 * escala, cx + 25 * escala, cy + 20 * escala)
            pdf.circle(cx + 27 * escala, cy + 22 * escala, 3.2 * escala, stroke=1, fill=1)

        elif "ciencia" in tema_normalizado:
            # Átomo.
            pdf.circle(cx, cy, 3.5 * escala, stroke=1, fill=1)
            pdf.ellipse(
                cx - 27 * escala, cy - 9 * escala,
                cx + 27 * escala, cy + 9 * escala,
                stroke=1, fill=0,
            )
            pdf.ellipse(
                cx - 9 * escala, cy - 27 * escala,
                cx + 9 * escala, cy + 27 * escala,
                stroke=1, fill=0,
            )
            pdf.saveState()
            pdf.translate(cx, cy)
            pdf.rotate(45)
            pdf.ellipse(
                -27 * escala, -9 * escala,
                27 * escala, 9 * escala,
                stroke=1, fill=0,
            )
            pdf.restoreState()
            pdf.setFillColor(secundario)
            pdf.circle(cx + 24 * escala, cy, 3 * escala, stroke=0, fill=1)

        elif "musica" in tema_normalizado:
            # Notas musicales vectoriales.
            pdf.setLineWidth(2.2 * escala)
            pdf.line(cx - 8 * escala, cy - 16 * escala, cx - 8 * escala, cy + 17 * escala)
            pdf.line(cx + 14 * escala, cy - 10 * escala, cx + 14 * escala, cy + 23 * escala)
            pdf.line(cx - 8 * escala, cy + 17 * escala, cx + 14 * escala, cy + 23 * escala)
            pdf.setFillColor(secundario)
            pdf.ellipse(
                cx - 17 * escala, cy - 20 * escala,
                cx - 5 * escala, cy - 10 * escala,
                stroke=0, fill=1,
            )
            pdf.ellipse(
                cx + 5 * escala, cy - 14 * escala,
                cx + 17 * escala, cy - 4 * escala,
                stroke=0, fill=1,
            )

        elif "tecnolog" in tema_normalizado or "robot" in tema_normalizado:
            # Microchip.
            pdf.roundRect(
                cx - 18 * escala, cy - 18 * escala,
                36 * escala, 36 * escala, 4 * escala,
                stroke=1, fill=0,
            )
            pdf.roundRect(
                cx - 9 * escala, cy - 9 * escala,
                18 * escala, 18 * escala, 2 * escala,
                stroke=1, fill=0,
            )
            for offset in (-12, -4, 4, 12):
                pdf.line(cx - 26 * escala, cy + offset * escala, cx - 18 * escala, cy + offset * escala)
                pdf.line(cx + 18 * escala, cy + offset * escala, cx + 26 * escala, cy + offset * escala)
                pdf.line(cx + offset * escala, cy - 26 * escala, cx + offset * escala, cy - 18 * escala)
                pdf.line(cx + offset * escala, cy + 18 * escala, cx + offset * escala, cy + 26 * escala)
            pdf.setFillColor(secundario)
            pdf.circle(cx, cy, 3.5 * escala, stroke=0, fill=1)

        elif "fisica" in tema_normalizado or "deporte" in tema_normalizado:
            # Medalla y pelota.
            pdf.circle(cx - 8 * escala, cy + 2 * escala, 15 * escala, stroke=1, fill=0)
            pdf.line(cx - 16 * escala, cy + 15 * escala, cx - 21 * escala, cy + 28 * escala)
            pdf.line(cx, cy + 15 * escala, cx + 5 * escala, cy + 28 * escala)
            dibujar_estrella_pdf(pdf, cx - 8 * escala, cy + 2 * escala, 7 * escala, secundario)
            pdf.circle(cx + 20 * escala, cy - 11 * escala, 10 * escala, stroke=1, fill=0)
            pdf.line(cx + 10 * escala, cy - 11 * escala, cx + 30 * escala, cy - 11 * escala)

        elif "teatro" in tema_normalizado:
            # Máscaras simplificadas.
            pdf.ellipse(
                cx - 27 * escala, cy - 14 * escala,
                cx - 2 * escala, cy + 16 * escala,
                stroke=1, fill=0,
            )
            pdf.ellipse(
                cx + 2 * escala, cy - 16 * escala,
                cx + 27 * escala, cy + 14 * escala,
                stroke=1, fill=0,
            )
            pdf.circle(cx - 19 * escala, cy + 4 * escala, 1.8 * escala, stroke=1, fill=1)
            pdf.circle(cx - 10 * escala, cy + 4 * escala, 1.8 * escala, stroke=1, fill=1)
            pdf.circle(cx + 10 * escala, cy + 2 * escala, 1.8 * escala, stroke=1, fill=1)
            pdf.circle(cx + 19 * escala, cy + 2 * escala, 1.8 * escala, stroke=1, fill=1)
            pdf.arc(
                cx - 21 * escala, cy - 10 * escala,
                cx - 8 * escala, cy + 2 * escala,
                200, 140,
            )
            pdf.arc(
                cx + 8 * escala, cy - 5 * escala,
                cx + 21 * escala, cy + 7 * escala,
                20, 140,
            )

        elif "ambient" in tema_normalizado or "ecolog" in tema_normalizado:
            # Hojas.
            pdf.bezier(
                cx, cy - 25 * escala,
                cx - 24 * escala, cy - 9 * escala,
                cx - 20 * escala, cy + 23 * escala,
                cx, cy + 20 * escala,
            )
            pdf.bezier(
                cx, cy - 25 * escala,
                cx + 24 * escala, cy - 9 * escala,
                cx + 20 * escala, cy + 23 * escala,
                cx, cy + 20 * escala,
            )
            pdf.line(cx, cy - 25 * escala, cx, cy + 20 * escala)
            pdf.setFillColor(secundario)
            pdf.circle(cx, cy - 26 * escala, 3 * escala, stroke=0, fill=1)

        else:
            # Institucional: laurel y estrella.
            dibujar_estrella_pdf(pdf, cx, cy + 4 * escala, 12 * escala, secundario)
            for lado in (-1, 1):
                pdf.setStrokeColor(primario)
                pdf.setLineWidth(1.2 * escala)
                pdf.arc(
                    cx - 34 * escala, cy - 25 * escala,
                    cx + 34 * escala, cy + 25 * escala,
                    90 if lado < 0 else 270,
                    90,
                )

        pdf.restoreState()

    def crear_diploma_pdf(
        datos,
        logo_bytes=None,
        firma_director_bytes=None,
        firma_profesor_bytes=None,
    ):
        """
        Diploma PDF V16.1:
        - sin emojis ni cuadrados;
        - separación vertical real entre bloques;
        - decoraciones vectoriales relacionadas con el área;
        - firmas y QR sin superponerse.
        """
        salida = BytesIO()
        pdf = rl_canvas.Canvas(salida, pagesize=landscape(letter))
        page_w, page_h = landscape(letter)

        estilo = ESTILOS_DIPLOMA[datos["estilo"]]
        tema = obtener_tema_visual_diploma(datos.get("area") or datos["estilo"])

        primario = HexColor(estilo["primario"])
        secundario = HexColor(estilo["secundario"])
        fondo = HexColor(estilo["fondo"])
        suave = HexColor(estilo["suave"])
        acento = HexColor(estilo["acento"])
        gris = HexColor("#3E4653")
        gris_suave = HexColor("#747B86")
        blanco = HexColor("#FFFFFF")
        borde_claro = HexColor("#E5D9BC")

        titulo = limpiar_texto_diploma(datos["titulo"], mayusculas=True)
        estudiante = limpiar_texto_diploma(datos["estudiante"], mayusculas=True)
        curso = limpiar_texto_diploma(datos["curso"])
        area = limpiar_texto_diploma(datos["area"])
        motivo = limpiar_texto_diploma(datos["motivo"])
        profesor = limpiar_texto_diploma(datos["profesor"])
        cargo_profesor = limpiar_texto_diploma(datos["cargo_profesor"])
        director = limpiar_texto_diploma(datos["director"])
        fecha_texto = limpiar_texto_diploma(format_date_es(datos["fecha"]))
        codigo_diploma = limpiar_texto_diploma(datos.get("codigo", ""))
        url_publica_diploma = str(datos.get("url_publica", "") or "").strip()

        # Fondo y marco.
        pdf.setFillColor(fondo)
        pdf.rect(0, 0, page_w, page_h, fill=1, stroke=0)
        pdf.setFillColor(primario)
        pdf.rect(0, page_h - 22, page_w, 22, fill=1, stroke=0)
        pdf.rect(0, 0, page_w, 14, fill=1, stroke=0)

        pdf.setStrokeColor(secundario)
        pdf.setLineWidth(3)
        pdf.rect(14, 14, page_w - 28, page_h - 28, fill=0, stroke=1)
        pdf.setStrokeColor(primario)
        pdf.setLineWidth(1.1)
        pdf.rect(24, 24, page_w - 48, page_h - 48, fill=0, stroke=1)
        pdf.setStrokeColor(secundario)
        pdf.setLineWidth(0.65)
        pdf.roundRect(36, 36, page_w - 72, page_h - 72, 10, fill=0, stroke=1)

        # Logo.
        logo_preparado = preparar_imagen_para_pdf(logo_bytes)
        if logo_preparado is None and RUTA_LOGO_DIPLOMA.exists():
            logo_preparado = preparar_imagen_para_pdf(RUTA_LOGO_DIPLOMA)
        dibujar_imagen_contenida_pdf(pdf, logo_preparado, 361, 506, 70, 70)

        # Encabezado.
        dibujar_texto_centrado_pdf(
            pdf, "LICEO BICENTENARIO DE EXCELENCIA",
            490, "Helvetica-Bold", 10.6, primario,
        )
        dibujar_texto_centrado_pdf(
            pdf, "COLEGIO ANTONIO VARAS",
            473, "Helvetica-Bold", 17.8, primario,
        )
        pdf.setStrokeColor(secundario)
        pdf.setLineWidth(1.25)
        pdf.line(220, 459, 572, 459)

        # Etiqueta del área, separada del título.
        etiqueta_area = f"ÁREA · {area.upper()}"
        tam_area = ajustar_tamano_pdf(
            etiqueta_area, "Helvetica-Bold", 9.2, 7.6, 235
        )
        pdf.setFillColor(suave)
        pdf.setStrokeColor(borde_claro)
        pdf.roundRect(278, 426, 236, 21, 10, fill=1, stroke=1)
        dibujar_texto_centrado_pdf(
            pdf, etiqueta_area, 433,
            "Helvetica-Bold", tam_area, acento,
        )

        # Decoraciones temáticas laterales.
        dibujar_motivo_area_vectorial(
            pdf, tema["nombre"], 89, 327, 0.9, primario, secundario
        )
        dibujar_motivo_area_vectorial(
            pdf, tema["nombre"], 703, 327, 0.9, primario, secundario
        )

        # Título: una o dos líneas, nunca superpuesto con la etiqueta.
        tam_titulo = 27.5
        lineas_titulo = envolver_texto_pdf(
            titulo, "Helvetica-Bold", tam_titulo, 560
        )
        while len(lineas_titulo) > 2 and tam_titulo > 20:
            tam_titulo -= 1
            lineas_titulo = envolver_texto_pdf(
                titulo, "Helvetica-Bold", tam_titulo, 560
            )

        alto_titulo = tam_titulo + 4
        y_titulo = 395 if len(lineas_titulo) == 1 else 405
        pdf.setFillColor(primario)
        pdf.setFont("Helvetica-Bold", tam_titulo)
        for indice, linea in enumerate(lineas_titulo[:2]):
            pdf.drawCentredString(
                page_w / 2,
                y_titulo - indice * alto_titulo,
                linea,
            )

        y_intro = 364 if len(lineas_titulo) == 1 else 345
        dibujar_texto_centrado_pdf(
            pdf,
            "Se otorga el presente diploma a",
            y_intro,
            "Helvetica-Oblique",
            11.2,
            gris_suave,
        )

        # Nombre: máximo dos líneas y con espacio suficiente.
        tam_nombre = 33
        lineas_nombre = envolver_texto_pdf(
            estudiante, "Helvetica-Bold", tam_nombre, 570
        )
        while len(lineas_nombre) > 2 and tam_nombre > 22:
            tam_nombre -= 1
            lineas_nombre = envolver_texto_pdf(
                estudiante, "Helvetica-Bold", tam_nombre, 570
            )

        y_nombre = y_intro - 39
        pdf.setFillColor(acento)
        pdf.setFont("Helvetica-Bold", tam_nombre)
        inter_nombre = tam_nombre + 4
        for indice, linea in enumerate(lineas_nombre[:2]):
            pdf.drawCentredString(
                page_w / 2,
                y_nombre - indice * inter_nombre,
                linea,
            )

        y_separador = (
            y_nombre
            - (len(lineas_nombre[:2]) - 1) * inter_nombre
            - 14
        )
        ancho_nombre = min(
            max(
                max(
                    stringWidth(linea, "Helvetica-Bold", tam_nombre)
                    for linea in lineas_nombre[:2]
                ) + 36,
                260,
            ),
            600,
        )
        pdf.setStrokeColor(secundario)
        pdf.setLineWidth(1.1)
        pdf.line(
            (page_w - ancho_nombre) / 2,
            y_separador,
            (page_w + ancho_nombre) / 2,
            y_separador,
        )

        # Curso y área separados.
        y_curso = y_separador - 24
        tam_curso = ajustar_tamano_pdf(
            curso, "Helvetica-Bold", 13.5, 10, 420
        )
        dibujar_texto_centrado_pdf(
            pdf, curso, y_curso,
            "Helvetica-Bold", tam_curso, primario,
        )
        dibujar_texto_centrado_pdf(
            pdf, area, y_curso - 17,
            "Helvetica", 11.3, gris,
        )

        # Motivo con altura dinámica.
        caja_x = 124
        caja_w = 544
        tam_motivo = 13.2
        lineas_motivo = envolver_texto_pdf(
            motivo, "Helvetica", tam_motivo, caja_w - 54
        )
        while len(lineas_motivo) > 4 and tam_motivo > 10.2:
            tam_motivo -= 0.5
            lineas_motivo = envolver_texto_pdf(
                motivo, "Helvetica", tam_motivo, caja_w - 54
            )

        inter_motivo = tam_motivo + 4
        caja_h = max(68, len(lineas_motivo[:4]) * inter_motivo + 28)
        caja_y = 164

        pdf.setFillColor(blanco)
        pdf.setStrokeColor(borde_claro)
        pdf.setLineWidth(0.85)
        pdf.roundRect(
            caja_x, caja_y, caja_w, caja_h,
            12, fill=1, stroke=1,
        )

        bloque_h = len(lineas_motivo[:4]) * inter_motivo
        y_linea = caja_y + (caja_h + bloque_h) / 2 - inter_motivo + 2
        pdf.setFillColor(gris)
        pdf.setFont("Helvetica", tam_motivo)
        for linea in lineas_motivo[:4]:
            pdf.drawCentredString(page_w / 2, y_linea, linea)
            y_linea -= inter_motivo

        # Fecha.
        dibujar_texto_centrado_pdf(
            pdf,
            f"Vicuña, {fecha_texto}",
            148,
            "Helvetica-Oblique",
            10.2,
            gris,
        )

        # Firmas y QR en tres columnas.
        firma_prof_preparada = preparar_imagen_para_pdf(
            firma_profesor_bytes, quitar_blanco=True
        )
        firma_dir_preparada = preparar_imagen_para_pdf(
            firma_director_bytes, quitar_blanco=True
        )
        if firma_dir_preparada is None and RUTA_FIRMA_DIRECTOR.exists():
            firma_dir_preparada = preparar_imagen_para_pdf(
                RUTA_FIRMA_DIRECTOR, quitar_blanco=True
            )

        dibujar_imagen_contenida_pdf(
            pdf, firma_prof_preparada, 86, 68, 205, 57
        )
        dibujar_imagen_contenida_pdf(
            pdf, firma_dir_preparada, 500, 67, 178, 62
        )

        qr_bytes = crear_qr_diploma(url_publica_diploma)
        if qr_bytes:
            pdf.setFillColor(blanco)
            pdf.setStrokeColor(borde_claro)
            pdf.roundRect(343, 54, 106, 84, 10, fill=1, stroke=1)
            dibujar_imagen_contenida_pdf(
                pdf, qr_bytes, 368, 75, 56, 56
            )
            pdf.setFillColor(primario)
            pdf.setFont("Helvetica-Bold", 6.6)
            pdf.drawCentredString(396, 68, codigo_diploma)
            pdf.setFillColor(gris_suave)
            pdf.setFont("Helvetica", 6.0)
            pdf.drawCentredString(
                396, 59, "Abrir diploma digital"
            )

        pdf.setStrokeColor(gris)
        pdf.setLineWidth(0.9)
        pdf.line(86, 61, 291, 61)
        pdf.line(500, 61, 678, 61)

        tam_prof = ajustar_tamano_pdf(
            profesor, "Helvetica-Bold", 10.2, 8, 200
        )
        pdf.setFillColor(primario)
        pdf.setFont("Helvetica-Bold", tam_prof)
        pdf.drawCentredString(188.5, 48, profesor)
        pdf.setFillColor(gris_suave)
        pdf.setFont("Helvetica", 7.8)
        pdf.drawCentredString(188.5, 37, cargo_profesor)

        tam_dir = ajustar_tamano_pdf(
            director, "Helvetica-Bold", 10.2, 8, 170
        )
        pdf.setFillColor(primario)
        pdf.setFont("Helvetica-Bold", tam_dir)
        pdf.drawCentredString(589, 48, director)
        pdf.setFillColor(gris_suave)
        pdf.setFont("Helvetica", 7.8)
        pdf.drawCentredString(589, 37, "Director")

        pdf.setFillColor(blanco)
        pdf.setFont("Helvetica", 6.2)
        pdf.drawCentredString(
            page_w / 2, 4.5,
            "Diploma digital emitido por el Sistema Institucional CAV",
        )

        pdf.setTitle(f"Diploma - {estudiante}")
        pdf.setAuthor("Colegio Antonio Varas")
        pdf.setSubject(titulo)
        pdf.showPage()
        pdf.save()
        return salida.getvalue()

    def crear_preview_desde_pdf(pdf_bytes):
        """Renderiza el mismo PDF como PNG; la vista previa y el archivo coinciden."""
        documento = fitz.open(stream=pdf_bytes, filetype="pdf")
        pagina = documento.load_page(0)
        matriz = fitz.Matrix(2.2, 2.2)
        pixmap = pagina.get_pixmap(matrix=matriz, alpha=False)
        png = pixmap.tobytes("png")
        documento.close()
        return png

    def exportar_diploma(imagen):
        png_buffer = BytesIO()
        imagen.save(png_buffer, format="PNG", optimize=True)
        png_bytes = png_buffer.getvalue()

        pdf_buffer = BytesIO()
        imagen.save(
            pdf_buffer,
            format="PDF",
            resolution=300.0,
            quality=95,
        )
        pdf_bytes = pdf_buffer.getvalue()

        return png_bytes, pdf_bytes

    try:
        CONFIG_DIPLOMAS = st.secrets["diplomas"]
    except Exception:
        CONFIG_DIPLOMAS = {}

    DOMINIO_INSTITUCIONAL = str(
        CONFIG_DIPLOMAS.get(
            "institutional_domain",
            "colegioantoniovaras.cl",
        )
    ).strip().lower()

    def validar_correo_institucional(correo):
        correo = str(correo or "").strip().lower()
        patron = r"^[A-Za-z0-9.!#$%&'*+/=?^_`{|}~-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$"

        if not re.fullmatch(patron, correo):
            return False, "El formato del correo no es válido."

        dominio = correo.rsplit("@", 1)[-1]
        if dominio != DOMINIO_INSTITUCIONAL:
            return (
                False,
                "El destinatario debe utilizar el dominio institucional "
                f"@{DOMINIO_INSTITUCIONAL}.",
            )

        return True, ""

    def correo_enmascarado(correo):
        correo = str(correo or "")
        if "@" not in correo:
            return "***"
        usuario, dominio = correo.split("@", 1)
        visible = usuario[:2] if len(usuario) >= 2 else usuario[:1]
        return f"{visible}***@{dominio}"

    def enviar_diploma_workspace(
        destinatario,
        datos,
        pdf_bytes,
        png_bytes,
        nombre_base,
        incluir_png=False,
    ):
        try:
            creds = st.secrets["email_credentials"]
        except Exception as exc:
            raise RuntimeError(
                "No existe la sección [email_credentials] en Streamlit Secrets."
            ) from exc

        smtp_server = str(creds.get("smtp_server", "smtp.gmail.com")).strip()
        smtp_port = int(creds.get("smtp_port", 587))
        smtp_username = str(creds.get("smtp_username", "")).strip()
        smtp_password = str(creds.get("smtp_password", "")).replace(" ", "")
        sender_email = str(creds.get("sender_email", smtp_username)).strip()
        sender_name = str(
            creds.get(
                "sender_name",
                "Liceo Bicentenario de Excelencia Colegio Antonio Varas",
            )
        ).strip()
        reply_to = str(creds.get("reply_to", sender_email)).strip()
        usar_tls = bool(creds.get("use_tls", smtp_port != 465))
        usar_ssl = bool(creds.get("use_ssl", smtp_port == 465))

        if not smtp_server:
            raise RuntimeError("Falta smtp_server en [email_credentials].")
        if not sender_email:
            raise RuntimeError(
                "Falta sender_email o smtp_username en [email_credentials]."
            )
        if smtp_server.lower() == "smtp.gmail.com":
            if not smtp_username or not smtp_password:
                raise RuntimeError(
                    "Para smtp.gmail.com debes configurar smtp_username y "
                    "smtp_password con una contraseña de aplicación."
                )

        estudiante = html_sanitizer.escape(str(datos["estudiante"]))
        titulo = html_sanitizer.escape(str(datos["titulo"]))
        curso = html_sanitizer.escape(str(datos["curso"]))
        area = html_sanitizer.escape(str(datos["area"]))
        fecha = html_sanitizer.escape(format_date_es(datos["fecha"]))
        codigo = html_sanitizer.escape(str(datos.get("codigo", "")))
        url_publica = str(datos.get("url_publica", "") or "").strip()

        asunto = f"Tu diploma digital CAV · {datos['estudiante']}"

        texto_plano = f"""Estimada/o {datos['estudiante']}:

El Liceo Bicentenario de Excelencia Colegio Antonio Varas hace entrega de tu
diploma digital: {datos['titulo']}.

Curso o unidad: {datos['curso']}
Área: {datos['area']}
Fecha: {format_date_es(datos['fecha'])}
Código: {datos.get('codigo', '')}

Abre tu diploma desde este enlace:
{url_publica}

También se adjunta una copia en PDF como respaldo.

Este reconocimiento digital contribuye a reducir el consumo de papel, tóner
y residuos de impresión.

Atentamente,
Colegio Antonio Varas
"""

        boton = ""
        if url_publica:
            boton = f"""
            <div style="text-align:center;margin:26px 0 22px;">
                <a href="{html_sanitizer.escape(url_publica)}"
                   style="display:inline-block;background:#800020;color:white;
                          padding:15px 25px;border-radius:12px;text-decoration:none;
                          font-weight:800;font-size:17px;">
                    🎓 Abrir mi diploma digital
                </a>
            </div>
            """

        html = f"""
        <div style="font-family:Arial,Helvetica,sans-serif;color:#20242b;line-height:1.6">
          <div style="max-width:680px;margin:auto;border:1px solid #e8dfc9;
                      border-radius:20px;overflow:hidden;background:#ffffff;">
            <div style="background:linear-gradient(135deg,#72001d,#970027);
                        color:white;padding:25px 30px;border-bottom:6px solid #c7a44a;">
              <div style="font-size:24px;font-weight:800;">Diploma Digital CAV</div>
              <div style="opacity:.92;">Liceo Bicentenario de Excelencia Colegio Antonio Varas</div>
            </div>

            <div style="padding:30px;">
              <p>Estimada/o <strong>{estudiante}</strong>:</p>
              <p>
                La comunidad educativa hace entrega de tu
                <strong>{titulo}</strong>.
              </p>

              <div style="background:#fff9ea;border:1px solid #ead8a7;
                          border-radius:14px;padding:16px 18px;margin:20px 0;">
                <strong>Curso o unidad:</strong> {curso}<br>
                <strong>Área:</strong> {area}<br>
                <strong>Fecha:</strong> {fecha}<br>
                <strong>Código oficial:</strong> {codigo}
              </div>

              {boton}

              <p style="font-size:14px;color:#5e6470;">
                Se adjunta una copia PDF como respaldo. No necesitas abrir
                archivos HTML ni instalar aplicaciones.
              </p>

              <div style="margin-top:22px;padding:16px 18px;border-radius:14px;
                          background:#eff9f1;border:1px solid #b9dec0;color:#174c29;">
                <strong>🌱 Acción ambiental institucional</strong><br>
                Este diploma digital evita una impresión en papel opalina y
                ayuda a reducir papel, tóner y residuos de impresión.
              </div>

              <p style="margin-top:26px;">
                Atentamente,<br>
                <strong>Colegio Antonio Varas</strong>
              </p>
            </div>
          </div>
        </div>
        """

        mensaje = EmailMessage()
        mensaje["Subject"] = asunto
        mensaje["From"] = formataddr((sender_name, sender_email))
        mensaje["To"] = destinatario
        if reply_to:
            mensaje["Reply-To"] = reply_to

        mensaje.set_content(texto_plano)
        mensaje.add_alternative(html, subtype="html")

        # Solo PDF como respaldo; no se adjunta HTML.
        mensaje.add_attachment(
            pdf_bytes,
            maintype="application",
            subtype="pdf",
            filename=f"{nombre_base}.pdf",
        )

        contexto_ssl = ssl.create_default_context()

        if usar_ssl or smtp_port == 465:
            with smtplib.SMTP_SSL(
                smtp_server,
                smtp_port,
                context=contexto_ssl,
                timeout=45,
            ) as servidor:
                if smtp_username and smtp_password:
                    servidor.login(smtp_username, smtp_password)
                servidor.send_message(mensaje)
        else:
            with smtplib.SMTP(smtp_server, smtp_port, timeout=45) as servidor:
                servidor.ehlo()
                if usar_tls:
                    servidor.starttls(context=contexto_ssl)
                    servidor.ehlo()
                if smtp_username and smtp_password:
                    servidor.login(smtp_username, smtp_password)
                servidor.send_message(mensaje)

        return True

    # Invalida automáticamente imágenes generadas por versiones anteriores.
    # Antes, Streamlit conservaba diploma_png en session_state y parecía que el
    # código nuevo no se aplicaba aunque el archivo sí se hubiera actualizado.
    if (
        st.session_state.get("diploma_render_version")
        != DIPLOMA_RENDER_VERSION
    ):
        st.session_state.diploma_png = None
        st.session_state.diploma_pdf = None
        st.session_state.diploma_datos = None
        st.session_state.diploma_nombre_archivo = "Diploma_CAV"
        st.session_state.diploma_ultimo_envio = None
        st.session_state.diploma_registro_id = None
        st.session_state.diploma_codigo = ""
        st.session_state.diploma_url_publica = ""
        st.session_state.diploma_render_version = DIPLOMA_RENDER_VERSION

    if "diploma_png" not in st.session_state:
        st.session_state.diploma_png = None
    if "diploma_pdf" not in st.session_state:
        st.session_state.diploma_pdf = None
    if "diploma_nombre_archivo" not in st.session_state:
        st.session_state.diploma_nombre_archivo = "Diploma_CAV"
    if "diploma_datos" not in st.session_state:
        st.session_state.diploma_datos = None
    if "diploma_correo" not in st.session_state:
        st.session_state.diploma_correo = ""
    if "diploma_ultimo_envio" not in st.session_state:
        st.session_state.diploma_ultimo_envio = None
    if "diploma_registro_id" not in st.session_state:
        st.session_state.diploma_registro_id = None
    if "diploma_codigo" not in st.session_state:
        st.session_state.diploma_codigo = ""
    if "diploma_url_publica" not in st.session_state:
        st.session_state.diploma_url_publica = ""

    col_formulario, col_info = st.columns([1.4, 0.8], gap="large")

    with col_info:
        st.markdown(
            """
            <div style="
                background:#FFF8E8;
                border:1px solid #E6C66B;
                border-left:6px solid #800020;
                border-radius:14px;
                padding:16px 18px;
                margin-bottom:16px;">
                <b>📄 Formato del diploma</b><br>
                Carta horizontal, alta resolución y listo para imprimir.
            </div>
            """,
            unsafe_allow_html=True,
        )

        if RUTA_LOGO_DIPLOMA.exists():
            st.success("✅ Logo institucional detectado.")
        else:
            st.warning("⚠️ No se encontró `logocav.png`.")

        if RUTA_FIRMA_DIRECTOR.exists():
            st.success("✅ Firma del director con sello detectada.")
            st.caption("Archivo utilizado: `firma_director.png`")
        else:
            st.warning(
                "⚠️ No se encontró `firma_director.png`. "
                "Puedes cargarla temporalmente en el formulario."
            )

        st.info(
            "La firma del director debe incluir el sello dentro de la misma imagen. "
            "No se utiliza un archivo de sello separado."
        )

        st.success(
            "✅ Motor activo: V16.3 · Render directo st.html + burdeo CAV #800020. "
            "Si no ves este mensaje, Streamlit todavía está ejecutando otro archivo."
        )

        url_base_detectada = obtener_url_base_aplicacion()
        if url_base_detectada:
            st.caption(f"🔗 URL pública detectada: {url_base_detectada}")
        else:
            st.warning(
                "No se pudo detectar la URL pública de la aplicación. "
                "Configura diplomas.public_app_url en Secrets."
            )

        try:
            correo_config = st.secrets["email_credentials"]
            servidor_config = correo_config.get("smtp_server", "smtp.gmail.com")
            st.success(f"✅ Correo configurado: {servidor_config}")
        except Exception:
            st.warning(
                "⚠️ Falta configurar [email_credentials] para enviar diplomas."
            )

        st.caption(
            f"Los destinatarios deben usar @{DOMINIO_INSTITUCIONAL}."
        )

        st.markdown("#### 🎨 Diseños disponibles")
        for nombre_estilo in ESTILOS_DIPLOMA:
            st.markdown(f"- {nombre_estilo}")

    with col_formulario:
        with st.form("form_generador_diploma", clear_on_submit=False):
            st.markdown("### 👤 Datos del reconocimiento")

            estudiante = st.text_input(
                "Nombre completo del estudiante *",
                placeholder="Ej. Martina González Rojas",
                max_chars=120,
            )

            correo_destinatario = st.text_input(
                "Correo institucional del destinatario",
                value=st.session_state.get("diploma_correo", ""),
                placeholder=f"nombre@{DOMINIO_INSTITUCIONAL}",
                help=(
                    "Puede dejarse vacío para generar solamente el diploma. "
                    "Será obligatorio al momento de enviarlo."
                ),
                max_chars=180,
            )

            c1, c2 = st.columns(2)

            opciones_curso = list(CURSOS) if CURSOS else []
            opciones_curso = opciones_curso + ["OTRO / SIN CURSO"]
            curso_seleccionado = c1.selectbox(
                "Curso *",
                opciones_curso,
                index=0 if opciones_curso else None,
            )

            curso_personalizado = ""
            if curso_seleccionado == "OTRO / SIN CURSO":
                curso_personalizado = c1.text_input(
                    "Escribe el curso",
                    placeholder="Ej. Taller de Robótica",
                )

            area = c2.selectbox(
                "Área o asignatura *",
                [
                    "Reconocimiento institucional",
                    "Lenguaje",
                    "Matemática",
                    "Ciencias",
                    "Artes",
                    "Música",
                    "Tecnología",
                    "Educación Física",
                    "Convivencia Escolar",
                    "Otra",
                ],
            )

            if area == "Otra":
                area = c2.text_input(
                    "Escribe el área",
                    placeholder="Ej. Taller de Debate",
                )

            tipo_reconocimiento = st.selectbox(
                "Tipo de reconocimiento",
                list(TEXTOS_RECONOCIMIENTO.keys()),
            )

            titulo_default = (
                "Diploma de Reconocimiento"
                if tipo_reconocimiento == "Texto personalizado"
                else tipo_reconocimiento
            )

            titulo_diploma = st.text_input(
                "Título del diploma *",
                value=titulo_default,
                max_chars=90,
            )

            motivo_default = TEXTOS_RECONOCIMIENTO[tipo_reconocimiento]
            motivo = st.text_area(
                "Texto del reconocimiento *",
                value=motivo_default,
                height=130,
                max_chars=650,
                placeholder="Describe el motivo del reconocimiento.",
            )

            c3, c4 = st.columns(2)
            fecha_diploma = c3.date_input(
                "Fecha del diploma",
                value=dt.date.today(),
                format="DD/MM/YYYY",
            )

            estilo_diploma = c4.selectbox(
                "Diseño visual",
                list(ESTILOS_DIPLOMA.keys()),
            )

            st.markdown("### ✍️ Firmas")

            c5, c6 = st.columns(2)
            profesor = c5.text_input(
                "Nombre del profesor responsable *",
                value=st.session_state.get("profesor_name") or "",
                max_chars=90,
            )
            cargo_profesor = c5.text_input(
                "Cargo del profesor",
                value="Profesor responsable",
                max_chars=70,
            )

            director = c6.text_input(
                "Nombre del director *",
                value="Director(a)",
                max_chars=90,
            )

            firma_profesor_archivo = st.file_uploader(
                "Firma del profesor (opcional)",
                type=["png", "jpg", "jpeg"],
                help="Preferentemente PNG con fondo transparente.",
                key="firma_profesor_diploma",
            )

            firma_director_archivo = st.file_uploader(
                "Firma del director con sello (opcional, reemplaza el archivo local)",
                type=["png", "jpg", "jpeg"],
                help=(
                    "Si no cargas una imagen, se utilizará automáticamente "
                    "`firma_director.png` cuando exista en el repositorio."
                ),
                key="firma_director_diploma",
            )

            logo_archivo = st.file_uploader(
                "Logo institucional alternativo (opcional)",
                type=["png", "jpg", "jpeg"],
                help=(
                    "Si no cargas un logo, se utilizará automáticamente "
                    "`logocav.png`."
                ),
                key="logo_diploma",
            )

            generar_diploma = st.form_submit_button(
                "✨ Generar vista previa y archivos",
                type="primary",
                use_container_width=True,
            )

            if generar_diploma:
                curso_final = (
                    curso_personalizado.strip()
                    if curso_seleccionado == "OTRO / SIN CURSO"
                    else curso_seleccionado
                )

                faltantes = []
                if not estudiante.strip():
                    faltantes.append("nombre del estudiante")
                if not curso_final:
                    faltantes.append("curso")
                if not str(area).strip():
                    faltantes.append("área")
                if not titulo_diploma.strip():
                    faltantes.append("título")
                if not motivo.strip():
                    faltantes.append("texto del reconocimiento")
                if not profesor.strip():
                    faltantes.append("profesor responsable")
                if not director.strip():
                    faltantes.append("director")

                firma_director_disponible = (
                    firma_director_archivo is not None
                    or RUTA_FIRMA_DIRECTOR.exists()
                )

                correo_valido = True
                mensaje_correo = ""
                if correo_destinatario.strip():
                    correo_valido, mensaje_correo = validar_correo_institucional(
                        correo_destinatario
                    )

                if faltantes:
                    st.warning(
                        "Completa los campos obligatorios: "
                        + ", ".join(faltantes)
                        + "."
                    )
                elif not correo_valido:
                    st.warning(mensaje_correo)
                elif not firma_director_disponible:
                    st.warning(
                        "Falta la firma del director con sello. "
                        "Carga la imagen o agrega `firma_director.png` "
                        "al repositorio."
                    )
                else:
                    datos_diploma = {
                        "estudiante": estudiante.strip(),
                        "curso": curso_final.strip(),
                        "area": str(area).strip(),
                        "titulo": titulo_diploma.strip(),
                        "motivo": motivo.strip(),
                        "fecha": fecha_diploma,
                        "estilo": estilo_diploma,
                        "profesor": profesor.strip(),
                        "cargo_profesor": cargo_profesor.strip() or "Profesor responsable",
                        "director": director.strip(),
                    }

                    logo_bytes = (
                        logo_archivo.getvalue()
                        if logo_archivo is not None
                        else None
                    )
                    firma_profesor_bytes = (
                        firma_profesor_archivo.getvalue()
                        if firma_profesor_archivo is not None
                        else None
                    )
                    firma_director_bytes = (
                        firma_director_archivo.getvalue()
                        if firma_director_archivo is not None
                        else None
                    )

                    registro_diploma = None
                    try:
                        with st.spinner(
                            "Creando registro oficial y diploma digital..."
                        ):
                            registro_diploma = crear_registro_diploma(
                                datos_diploma,
                                correo_destinatario,
                                origen="individual",
                            )

                            datos_diploma["codigo"] = registro_diploma["codigo"]
                            datos_diploma["public_token"] = registro_diploma["public_token"]
                            datos_diploma["url_publica"] = (
                                registro_diploma.get("url_publica") or ""
                            )

                            pdf_bytes = crear_diploma_pdf(
                                datos_diploma,
                                logo_bytes=logo_bytes,
                                firma_director_bytes=firma_director_bytes,
                                firma_profesor_bytes=firma_profesor_bytes,
                            )
                            png_bytes = crear_preview_desde_pdf(pdf_bytes)

                            carpeta = (
                                f"{dt.date.today().year}/"
                                f"{registro_diploma['codigo']}"
                            )
                            pdf_path = f"{carpeta}/diploma.pdf"
                            preview_path = f"{carpeta}/preview.png"

                            subir_archivo_diploma_storage(
                                pdf_path,
                                pdf_bytes,
                                "application/pdf",
                            )
                            subir_archivo_diploma_storage(
                                preview_path,
                                png_bytes,
                                "image/png",
                            )
                            actualizar_archivos_diploma(
                                registro_diploma["id"],
                                pdf_path,
                                preview_path,
                            )

                        nombre_base = (
                            "Diploma_"
                            + archivo_seguro_diploma(estudiante)
                            + "_"
                            + archivo_seguro_diploma(str(area))
                        )

                        st.session_state.diploma_png = png_bytes
                        st.session_state.diploma_pdf = pdf_bytes
                        st.session_state.diploma_nombre_archivo = nombre_base
                        st.session_state.diploma_datos = datos_diploma
                        st.session_state.diploma_correo = (
                            correo_destinatario.strip().lower()
                        )
                        st.session_state.diploma_ultimo_envio = None
                        st.session_state.diploma_registro_id = registro_diploma["id"]
                        st.session_state.diploma_codigo = registro_diploma["codigo"]
                        st.session_state.diploma_url_publica = (
                            registro_diploma.get("url_publica") or ""
                        )

                        registrar_auditoria(
                            "generó diploma digital",
                            "Diplomas",
                            registro_id=registro_diploma["id"],
                            detalle={
                                "codigo": registro_diploma["codigo"],
                                "curso": curso_final.strip(),
                                "area": str(area).strip(),
                                "estilo": estilo_diploma,
                            },
                        )

                        st.success(
                            "✅ Diploma digital registrado y generado correctamente."
                        )
                        if not datos_diploma.get("url_publica"):
                            st.warning(
                                "El PDF se generó, pero falta configurar "
                                "`diplomas.public_app_url` para crear el enlace digital."
                            )
                        st.balloons()

                    except Exception as e:
                        if registro_diploma:
                            marcar_error_diploma(registro_diploma["id"], e)
                        registrar_error("generar_diploma", e)
                        st.error(
                            "No fue posible generar el diploma digital. "
                            f"Detalle técnico: {e}"
                        )

    if st.session_state.diploma_png:
        st.markdown("---")
        st.markdown("## 👁️ Vista previa")

        st.image(
            st.session_state.diploma_png,
            caption=(
                "Vista previa renderizada desde el mismo PDF vectorial V13. "
                "Lo que ves coincide con el archivo descargado y enviado."
            ),
            use_container_width=True,
        )

        if st.session_state.get("diploma_url_publica"):
            st.link_button(
                "🌐 Abrir versión digital",
                st.session_state.diploma_url_publica,
                type="primary",
                use_container_width=True,
            )
            st.caption(
                f"Código oficial: {st.session_state.get('diploma_codigo', '')}"
            )

        col_pdf, col_png, col_limpiar = st.columns([1, 1, 0.75])

        with col_pdf:
            st.download_button(
                "📄 Descargar diploma PDF",
                data=st.session_state.diploma_pdf,
                file_name=(
                    st.session_state.diploma_nombre_archivo + ".pdf"
                ),
                mime="application/pdf",
                type="primary",
                use_container_width=True,
            )

        with col_png:
            st.download_button(
                "🖼️ Descargar diploma PNG",
                data=st.session_state.diploma_png,
                file_name=(
                    st.session_state.diploma_nombre_archivo + ".png"
                ),
                mime="image/png",
                use_container_width=True,
            )

        with col_limpiar:
            if st.button(
                "🧹 Limpiar",
                use_container_width=True,
                key="limpiar_diploma",
            ):
                st.session_state.diploma_png = None
                st.session_state.diploma_pdf = None
                st.session_state.diploma_datos = None
                st.session_state.diploma_correo = ""
                st.session_state.diploma_ultimo_envio = None
                st.session_state.diploma_nombre_archivo = "Diploma_CAV"
                st.session_state.diploma_registro_id = None
                st.session_state.diploma_codigo = ""
                st.session_state.diploma_url_publica = ""
                st.rerun()

        st.markdown("---")
        st.markdown("## 📧 Enviar por Google Workspace")
        st.caption(
            "El destinatario recibirá un botón directo para abrir el diploma digital "
            "y un PDF de respaldo. No se adjuntan archivos HTML."
        )

        datos_generados = st.session_state.get("diploma_datos") or {}

        with st.form("form_enviar_diploma_workspace", clear_on_submit=False):
            correo_envio = st.text_input(
                "Correo institucional del destinatario *",
                value=st.session_state.get("diploma_correo", ""),
                placeholder=f"nombre@{DOMINIO_INSTITUCIONAL}",
                max_chars=180,
            )

            st.info(
                "El correo incluirá un botón para abrir el diploma digital y "
                "una copia PDF de respaldo. No se adjuntan archivos HTML."
            )

            confirmar_envio = st.checkbox(
                "Confirmo que el nombre y el correo del destinatario son correctos.",
                value=False,
            )

            enviar_correo_diploma = st.form_submit_button(
                "📨 Enviar diploma al correo institucional",
                type="primary",
                use_container_width=True,
            )

            if enviar_correo_diploma:
                valido, mensaje_validacion = validar_correo_institucional(
                    correo_envio
                )

                if not datos_generados:
                    st.error(
                        "No se encontraron los datos del diploma. "
                        "Vuelve a generarlo antes de enviarlo."
                    )
                elif not valido:
                    st.warning(mensaje_validacion)
                elif not confirmar_envio:
                    st.warning(
                        "Marca la casilla de confirmación antes de enviar."
                    )
                else:
                    token_envio = hashlib.sha256(
                        st.session_state.diploma_pdf
                        + correo_envio.strip().lower().encode("utf-8")
                    ).hexdigest()

                    if (
                        st.session_state.diploma_ultimo_envio == token_envio
                    ):
                        st.warning(
                            "Este mismo diploma ya fue enviado a ese correo "
                            "durante la sesión. Genera nuevamente el diploma "
                            "o cambia el destinatario para realizar otro envío."
                        )
                    else:
                        try:
                            with st.spinner(
                                "Enviando diploma mediante Google Workspace..."
                            ):
                                if datos_generados.get("public_token"):
                                    url_actualizada = construir_url_publica_diploma(
                                        datos_generados["public_token"]
                                    )
                                    datos_generados["url_publica"] = url_actualizada
                                    st.session_state.diploma_url_publica = url_actualizada

                                    if st.session_state.get("diploma_registro_id"):
                                        supabase.table("diplomas_digitales").update({
                                            "url_publica": url_actualizada,
                                            "actualizado_en": dt_datetime.now().isoformat(),
                                        }).eq(
                                            "id",
                                            st.session_state.diploma_registro_id,
                                        ).execute()

                                enviar_diploma_workspace(
                                    destinatario=correo_envio.strip().lower(),
                                    datos=datos_generados,
                                    pdf_bytes=st.session_state.diploma_pdf,
                                    png_bytes=st.session_state.diploma_png,
                                    nombre_base=(
                                        st.session_state.diploma_nombre_archivo
                                    ),
                                    incluir_png=False,
                                )

                            st.session_state.diploma_correo = (
                                correo_envio.strip().lower()
                            )
                            st.session_state.diploma_ultimo_envio = token_envio

                            if st.session_state.get("diploma_registro_id"):
                                marcar_diploma_enviado(
                                    st.session_state.diploma_registro_id,
                                    correo_envio,
                                )

                            registrar_auditoria(
                                "envió diploma por correo",
                                "Diplomas",
                                detalle={
                                    "destinatario": correo_enmascarado(
                                        correo_envio
                                    ),
                                    "area": datos_generados.get("area"),
                                    "formato": "enlace_digital_mas_pdf",
                                    "codigo": st.session_state.get("diploma_codigo"),
                                },
                            )

                            st.success(
                                "✅ Diploma enviado correctamente a "
                                f"{correo_enmascarado(correo_envio)}."
                            )
                            st.balloons()

                        except smtplib.SMTPAuthenticationError:
                            st.error(
                                "Google rechazó la autenticación SMTP. "
                                "Verifica el correo emisor y utiliza una "
                                "contraseña de aplicación, no la contraseña "
                                "normal de la cuenta."
                            )
                        except smtplib.SMTPRecipientsRefused:
                            st.error(
                                "El servidor rechazó el correo destinatario. "
                                "Comprueba que la cuenta institucional exista."
                            )
                        except Exception as e:
                            if st.session_state.get("diploma_registro_id"):
                                marcar_error_diploma(
                                    st.session_state.diploma_registro_id,
                                    e,
                                )
                            registrar_error("enviar_diploma_workspace", e)
                            st.error(
                                "No fue posible enviar el diploma. "
                                f"Detalle técnico: {e}"
                            )


# ------------------------------------------------------------------
# SECCIÓN: CONFIGURACIÓN
# ------------------------------------------------------------------
if page == "Configuración":
    st.title("⚙️ Configuración del Sistema")
    st.write("Desde aquí puedes administrar los elementos centrales de la aplicación.")
    
    tab_prof, tab_cur, tab_rec, tab_estado = st.tabs(["Profesores", "Cursos", "Recursos", "Estado del sistema"])
    
    with tab_prof:
        st.write("### 👥 Administración de Profesores")
        col_add, col_list = st.columns([1, 2])
        with col_add:
            with st.form("form_add_prof"):
                nuevo_prof = st.text_input("Nombre y Apellidos")
                nuevo_email = st.text_input("Correo Electrónico (Opcional)")
                if st.form_submit_button("➕ Agregar Profesor", use_container_width=True):
                    if nuevo_prof.strip():
                        try:
                            supabase.table("profesores").insert({"nombre": nuevo_prof.strip().upper(), "email": nuevo_email.strip()}).execute()
                            st.success("¡Profesor agregado!")
                            st.cache_data.clear(); time.sleep(0.15); st.rerun()
                        except Exception as e: st.error(f"Error al agregar: {e}")
                    else: st.error("El nombre es obligatorio.")
        with col_list:
            prof_data = supabase.table("profesores").select("*").order("nombre").execute().data
            if prof_data:
                df_p = pd.DataFrame(prof_data)
                st.dataframe(df_p[['nombre', 'email']], use_container_width=True, hide_index=True)
                with st.expander("🗑️ Eliminar un Profesor"):
                    st.warning("⚠️ Nota: No puedes eliminar a un profesor si ya tiene reservas en el sistema.")
                    prof_borrar = st.selectbox("Selecciona el profesor a eliminar", df_p['nombre'].tolist(), key="del_prof")
                    if st.button("Eliminar Profesor Definitivamente", type="primary"):
                        try:
                            id_b = int(df_p[df_p['nombre'] == prof_borrar]['id'].values[0])
                            supabase.table("profesores").delete().eq("id", id_b).execute()
                            st.success(f"Profesor {prof_borrar} eliminado.")
                            st.cache_data.clear(); time.sleep(0.15); st.rerun()
                        except Exception as e: st.error("No se puede eliminar porque este profesor tiene reservas asociadas.")

    with tab_cur:
        st.write("### 📚 Administración de Cursos")
        col_add_c, col_list_c = st.columns([1, 2])
        with col_add_c:
            with st.form("form_add_cur"):
                nuevo_curso = st.text_input("Nombre del Curso (ej. 1° BÁSICO A)")
                if st.form_submit_button("➕ Agregar Curso", use_container_width=True):
                    if nuevo_curso.strip():
                        try:
                            supabase.table("cursos").insert({"nombre": nuevo_curso.strip().upper()}).execute()
                            st.success("¡Curso agregado!")
                            st.cache_data.clear(); time.sleep(0.15); st.rerun()
                        except Exception as e: st.error(f"Error al agregar: {e}")
                    else: st.error("El nombre es obligatorio.")
        with col_list_c:
            cur_data = supabase.table("cursos").select("*").order("nombre").execute().data
            if cur_data:
                df_c = pd.DataFrame(cur_data)
                st.dataframe(df_c[['nombre']], use_container_width=True, hide_index=True)
                with st.expander("🗑️ Eliminar un Curso"):
                    cur_borrar = st.selectbox("Selecciona el curso a eliminar", df_c['nombre'].tolist(), key="del_cur")
                    if st.button("Eliminar Curso Definitivamente", type="primary"):
                        try:
                            id_b = int(df_c[df_c['nombre'] == cur_borrar]['id'].values[0])
                            supabase.table("cursos").delete().eq("id", id_b).execute()
                            st.success(f"Curso {cur_borrar} eliminado.")
                            st.cache_data.clear(); time.sleep(0.15); st.rerun()
                        except Exception as e: st.error("No se puede eliminar porque este curso tiene reservas asociadas.")

    with tab_rec:
        st.write("### 💻 Administración de Recursos")
        col_add_r, col_list_r = st.columns([1, 2])
        with col_add_r:
            with st.form("form_add_rec"):
                nuevo_rec = st.text_input("Nombre del Recurso (ej. Proyector 5)")
                if st.form_submit_button("➕ Agregar Recurso", use_container_width=True):
                    if nuevo_rec.strip():
                        try:
                            supabase.table("recursos").insert({"nombre": nuevo_rec.strip().upper()}).execute()
                            st.success("¡Recurso agregado!")
                            st.cache_data.clear(); time.sleep(0.15); st.rerun()
                        except Exception as e: st.error(f"Error al agregar: {e}")
                    else: st.error("El nombre es obligatorio.")
        with col_list_r:
            rec_data = supabase.table("recursos").select("*").order("nombre").execute().data
            if rec_data:
                df_r = pd.DataFrame(rec_data)
                st.dataframe(df_r[['nombre']], use_container_width=True, hide_index=True)
                with st.expander("🗑️ Eliminar un Recurso"):
                    rec_borrar = st.selectbox("Selecciona el recurso a eliminar", df_r['nombre'].tolist(), key="del_rec")
                    if st.button("Eliminar Recurso Definitivamente", type="primary"):
                        try:
                            id_b = int(df_r[df_r['nombre'] == rec_borrar]['id'].values[0])
                            supabase.table("recursos").delete().eq("id", id_b).execute()
                            st.success(f"Recurso {rec_borrar} eliminado.")
                            st.cache_data.clear(); time.sleep(0.15); st.rerun()
                        except Exception as e: st.error("No se puede eliminar porque tiene reservas o reportes de mantenimiento asociados.")

    with tab_estado:
        st.write("### 🩺 Estado del sistema")
        st.caption("Diagnóstico rápido de servicios y tablas principales.")

        servicios = []
        try:
            prueba = supabase.table("recursos").select("id").limit(1).execute()
            servicios.append(("Supabase", "✅ Conectado", "Base de datos disponible"))
        except Exception as e:
            registrar_error("estado_supabase", e)
            servicios.append(("Supabase", "❌ Error", str(e)))

        try:
            _ = model
            servicios.append(("Gemini", "✅ Configurado", "Modelo cargado"))
        except Exception as e:
            servicios.append(("Gemini", "❌ Error", str(e)))

        try:
            creds = st.secrets["email_credentials"]
            servicios.append(("Correo SMTP", "✅ Configurado", creds.get("smtp_server", "Servidor definido")))
        except Exception:
            servicios.append(("Correo SMTP", "⚠️ Sin configurar", "Las notificaciones por correo no estarán disponibles"))

        st.dataframe(
            pd.DataFrame(servicios, columns=["Servicio", "Estado", "Detalle"]),
            use_container_width=True,
            hide_index=True
        )

        st.write("### 📊 Conteo de tablas")
        conteos = []
        for nombre_tabla in ["reservas", "recursos", "profesores", "cursos", "mantenimientos", "equipos", "eventos_tv", "anuncios_urgentes"]:
            try:
                respuesta = supabase.table(nombre_tabla).select("id", count="exact").limit(1).execute()
                conteos.append({"Tabla": nombre_tabla, "Registros": respuesta.count or 0, "Estado": "✅"})
            except Exception as e:
                registrar_error(f"conteo_{nombre_tabla}", e)
                conteos.append({"Tabla": nombre_tabla, "Registros": None, "Estado": "❌"})
        st.dataframe(pd.DataFrame(conteos), use_container_width=True, hide_index=True)

        st.write("### 🧾 Errores recientes")
        errores = st.session_state.get("errores_sistema", [])
        if errores:
            st.dataframe(pd.DataFrame(errores[::-1]), use_container_width=True, hide_index=True)
            if st.button("🧹 Limpiar registro de errores", use_container_width=True):
                st.session_state.errores_sistema = []
                st.rerun()
        else:
            st.success("No hay errores registrados en esta sesión.")

# ==============================================================================
# 📺 PÁGINA: GESTIÓN DE TV Y MENSAJERÍA
# ==============================================================================
elif page == "Modo TV":
    # ==============================================================================
    # 📺 1. VISUALIZACIÓN DE LA PANTALLA PÚBLICA (MODO KIOSCO)
    # ==============================================================================
    if st.session_state.get("ver_pantalla_tv", False):
        import streamlit.components.v1 as components
        import pandas as pd
        
        # Refresco automático de pantalla cada 20 segundos
        refresh_count = st_autorefresh(interval=30000, key="tv_refresh_timer")
        
        if "tv_scale_saved" not in st.session_state:
            st.session_state.tv_scale_saved = 100

        escala_pct = int(st.session_state.get("tv_scale_saved", 100))
        escala_pct = max(80, min(130, escala_pct))
        st.session_state.tv_scale_saved = escala_pct
        escala = escala_pct / 100.0

        def guardar_escala_tv_secundaria():
            valor = int(st.session_state.get("tv_scale_widget_secundario", 100))
            st.session_state.tv_scale_saved = max(80, min(130, valor))
        
        now_dt = dt_datetime.now()
        hoy_str = now_dt.strftime("%Y-%m-%d")
        hora_actual = now_dt.strftime("%H:%M")

        # --- 🚨 PRIORIDAD MULTIMEDIA: ALERTA ROJA GIGANTE ---
        try:
            alertas_criticas = (
                supabase.table("anuncios_urgentes")
                .select("*")
                .eq("is_active", True)
                .eq("prioridad", 999)
                .order("id", desc=True)
                .limit(100)
                .execute()
                .data or []
            )

            alerta = None
            ahora_tv = dt_datetime.now(TZ_CHILE)
            for candidata in alertas_criticas:
                _, _, estado_alerta = obtener_ventana_alerta(candidata, ahora_tv)
                if estado_alerta == "EN_CURSO" and alerta is None:
                    alerta = candidata
                elif estado_alerta in ("FINALIZADA", "INVALIDA"):
                    supabase.table("anuncios_urgentes").update(
                        {"is_active": False}
                    ).eq("id", candidata["id"]).execute()

            if alerta:
                mensaje_seguro = html_sanitizer.escape(
                    str(alerta.get("descripcion", ""))
                )
                st.markdown(f"""
                    <style>
                    html, body, .stApp {{
                        zoom: 1 !important;
                        transform: none !important;
                    }}
                    .stApp {{ background-color: #ff0000 !important; }}
                    header, [data-testid="stSidebar"] {{ display: none !important; }}
                    .block-container {{ max-width:none !important; padding:0 !important; }}
                    .alerta-total {{
                        position:fixed;
                        inset:0;
                        z-index:999999;
                        width:100vw;
                        height:100vh;
                        display:flex;
                        flex-direction:column;
                        align-items:center;
                        justify-content:center;
                        padding:4vh 5vw;
                        box-sizing:border-box;
                        color:white;
                        text-align:center;
                        background:#ff0000;
                    }}
                    .at-titulo {{
                        font-size:clamp(46px,7vw,105px);
                        font-weight:950;
                        animation:blink 1s infinite;
                    }}
                    .at-msg {{
                        margin-top:4vh;
                        width:min(92vw,1500px);
                        font-size:clamp(28px,4.2vw,64px);
                        line-height:1.12;
                        font-weight:750;
                        overflow-wrap:anywhere;
                    }}
                    @keyframes blink {{
                        0%,100% {{ opacity:1; }}
                        50% {{ opacity:.45; }}
                    }}
                    </style>
                    <div class="alerta-total">
                        <div class="at-titulo">⚠️ AVISO URGENTE ⚠️</div>
                        <div class="at-msg">{mensaje_seguro}</div>
                    </div>
                """, unsafe_allow_html=True)
                st.stop()
        except Exception as e:
            registrar_error("alerta_roja_tv_secundaria", e)

        # --- 🎨 ESTILOS INYECTADOS ANTIBUGS (FUERZA COLORES DE CONTROLES) ---
        st.markdown(f"""
        <style>
            @import url('https://unpkg.com/@phosphor-icons/web@2.1.1/src/fill/style.css');
            
            /* Forzar ocultamiento de cabeceras Streamlit */
            [data-testid="stHeader"], [data-testid="stSidebar"] {{ display: none !important; }}
            
            /* Encabezado Principal de la TV */
            .tv-header {{ background: #ffffff; color: #0f172a; padding: 15px 30px; border-radius: 20px; margin-bottom: 25px; display: flex; justify-content: space-between; align-items: center; border: 2px solid #cbd5e1; box-shadow: 0 4px 6px rgba(0,0,0,0.05); }}
            
            /* Bloques de Actividades */
            .card-evento {{ background: #ffffff; padding: 22px; border-radius: 15px; border-left: 10px solid #3b82f6; margin-bottom: 18px; color: #1e293b; border-top: 1px solid #e2e8f0; border-right: 1px solid #e2e8f0; border-bottom: 1px solid #e2e8f0; box-shadow: 0 2px 4px rgba(0,0,0,0.02); }}
            .card-reserva {{ border-left-color: #10b981 !important; }}
            .evento-titulo {{ font-weight: 800; font-size: calc(1.4rem * {escala}); color: #1e40af; }}
            
            /* Barra de progreso de refresco */
            .progress-bar {{ height: 6px; background: #3b82f6; width: 0%; animation: load 20s linear infinite; border-radius: 10px; margin-bottom: 20px; }}
            @keyframes load {{ 0% {{ width: 0%; }} 100% {{ width: 100%; }} }}
            
            /* 👇 PROTECCIÓN ABSOLUTA PARA CAJA DE CONFIGURACIÓN CONTRA TEXTO INVISIBLE 👇 */
            .contenedor-controles {{ background-color: #f8fafc !important; padding: 20px; border-radius: 15px; border: 2px solid #3b82f6 !important; color: #0f172a !important; }}
            .contenedor-controles * {{ color: #0f172a !important; font-weight: 600 !important; }}
            div[data-testid="stSelectbox"] label, div[data-testid="stSlider"] label {{ color: #0f172a !important; font-weight: 800 !important; font-size: 1.1rem !important; }}
            div[data-testid="stSelectbox"] div[data-baseweb="select"] {{ background-color: #ffffff !important; border: 1px solid #cbd5e1 !important; }}
        </style>
        """, unsafe_allow_html=True)

        clima = obtener_clima_vicuna()
        st.markdown(f"""
            <div class="tv-header">
                <div style="font-size: calc(1.8rem * {escala}); font-weight: 900; color: #1e40af;"><i class="ph-fill ph-monitor-play"></i> COLEGIO TV</div>
                <div style="font-size: calc(1.2rem * {escala}); font-weight: 700; color: #334155;">{now_dt.strftime("%d/%m/%Y")} | {clima} | <span style="color:#3b82f6;">{hora_actual}</span></div>
            </div><div class="progress-bar"></div>
        """, unsafe_allow_html=True)

        col_izq, col_der = st.columns([2.3, 1], gap="large")

        with col_izq:
            eventos_hoy = []
            if st.session_state.get('url_calendario_tv'):
                eventos_hoy.extend(obtener_eventos_google_calendar(st.session_state.url_calendario_tv))
            
            # --- CARGA DE DATOS DE SUPABASE (CON ALERTA DE ERROR VISIBLE SI FALLA) ---
            try:
                try:
                    ev_data = (
                        supabase.table("eventos_tv")
                        .select("*")
                        .eq("is_active", True)
                        .lte("fecha_inicio", hoy_str)
                        .gte("fecha_fin", hoy_str)
                        .order("hora_inicio")
                        .execute()
                        .data or []
                    )
                except Exception:
                    ev_data = (
                        supabase.table("eventos_tv")
                        .select("*")
                        .eq("fecha_evento", hoy_str)
                        .eq("is_active", True)
                        .order("hora_inicio")
                        .execute()
                        .data or []
                    )

                for e in ev_data:
                    if hora_actual <= str(e.get("hora_fin", "23:59")):
                        fecha_ini_ev = str(e.get("fecha_inicio") or e.get("fecha_evento") or hoy_str)
                        fecha_fin_ev = str(e.get("fecha_fin") or e.get("fecha_evento") or hoy_str)
                        desc_evento = str(e.get("descripcion", "") or "").strip()
                        if fecha_ini_ev != fecha_fin_ev:
                            try:
                                periodo = (
                                    f"Vigente del "
                                    f"{dt_datetime.strptime(fecha_ini_ev, '%Y-%m-%d').strftime('%d/%m')} "
                                    f"al {dt_datetime.strptime(fecha_fin_ev, '%Y-%m-%d').strftime('%d/%m')}"
                                )
                            except Exception:
                                periodo = f"Vigencia: {fecha_ini_ev} al {fecha_fin_ev}"
                            desc_evento = f"{desc_evento} · {periodo}" if desc_evento else periodo

                        eventos_hoy.append({
                            "hora": str(e.get("hora_inicio", "00:00"))[:5],
                            "titulo": e["titulo"],
                            "desc": desc_evento,
                            "tipo": "evento",
                        })
                
                # Se obtienen reservas para el perfil seleccionado
                perfil_actual = st.session_state.get("tv_profile", "")
                if "PROFESOR" in perfil_actual.upper() or "INSPECTOR" in perfil_actual.upper() or perfil_actual == "":
                    res_data = supabase.table("reservas").select("*, recursos(nombre), cursos(nombre), profesores(nombre)").eq("fecha", hoy_str).execute().data or []
                    for r in res_data:
                        if hora_actual <= str(r.get("hora_fin", "23:59")):
                            recurso_nom = r['recursos']['nombre'] if r.get('recursos') else "Recurso"
                            curso_nom = r['cursos']['nombre'] if r.get('cursos') else "Sin Curso"
                            prof_nom = r['profesores']['nombre'] if r.get('profesores') else "Docente"
                            eventos_hoy.append({
                                "hora": str(r.get("hora_inicio", "00:00"))[:5], 
                                "titulo": f"🔒 {recurso_nom} - {curso_nom}", 
                                "desc": f"Docente: {prof_nom}", 
                                "tipo": "reserva"
                            })
            except Exception as e:
                st.error(f"⚠️ Error cargando datos de la Base de Datos: {e}")

            eventos_hoy = sorted(eventos_hoy, key=lambda x: x['hora'])
            
            if not eventos_hoy:
                st.info("No quedan actividades o reservas programadas para el resto del día.")
            else:
                PAG_SIZE = 4
                total_pag = max(1, (len(eventos_hoy) + PAG_SIZE - 1) // PAG_SIZE)
                p_act = refresh_count % total_pag
                items = eventos_hoy[p_act*PAG_SIZE : (p_act+1)*PAG_SIZE]
                
                st.markdown(f"<h3 style='color:#1e3a8a;'>📅 Cronograma de Actividades ({p_act+1}/{total_pag})</h3>", unsafe_allow_html=True)
                for it in items:
                    clase = "card-evento card-reserva" if it['tipo'] == "reserva" else "card-evento"
                    st.markdown(f"""
                    <div class="{clase}">
                        <div style="display:flex; justify-content:space-between; align-items:center;">
                            <div class="evento-titulo">{it['titulo']}</div>
                            <div style="font-weight:800; background:#f1f5f9; color:#1e40af; padding:6px 12px; border-radius:8px; border:1px solid #cbd5e1;">⏱️ {it['hora']}</div>
                        </div>
                        <div style="margin-top:8px; color:#475569; font-size:calc(1.05rem * {escala}); font-weight:500;">{it['desc']}</div>
                    </div>
                    """, unsafe_allow_html=True)

        with col_der:
            # --- CONTENEDOR DE CONTROLES CON TIPOGRAFÍA PROTEGIDA EN NEGRO ---
            st.markdown('<div class="contenedor-controles">', unsafe_allow_html=True)
            st.markdown("<h4 style='margin-top:0; color:#0f172a; font-weight:900;'>⚙️ Panel de Pantalla</h4>", unsafe_allow_html=True)
            
            st.selectbox("👁️ Perfil Visual", ["Inspectoría / UTP", "Profesores / PIE", "Apoderados"], key="tv_profile")
            if "tv_scale_widget_secundario" not in st.session_state:
                st.session_state.tv_scale_widget_secundario = int(
                    st.session_state.get("tv_scale_saved", 100)
                )

            st.slider(
                "🔍 Tamaño Texto (%)",
                50,
                200,
                key="tv_scale_widget_secundario",
                step=5,
                on_change=guardar_escala_tv_secundaria,
            )
            
            st.markdown("</div>", unsafe_allow_html=True) # Cierre del contenedor estilizado
            
            st.write("") # Espaciador
            if st.button("🔙 VOLVER AL PANEL DE GESTIÓN", use_container_width=True, type="primary"):
                st.session_state.ver_pantalla_tv = False
                st.rerun()
            
            components.html("""
                <button onclick="const doc=window.parent.document; if(!doc.fullscreenElement){doc.documentElement.requestFullscreen();}else{doc.exitFullscreen();}" 
                style="width:100%; height:42px; background:#ffffff; border:2px solid #cbd5e1; border-radius:8px; font-family:sans-serif; font-weight:bold; color:#0f172a; cursor:pointer; box-shadow:0 2px 4px rgba(0,0,0,0.05);">
                🔲 Alternar Pantalla Completa</button>
            """, height=48)

            st.divider()
            
            # --- SECCIÓN AVISOS LATERALES ---
            st.markdown("<h3 style='color:#1e3a8a;'>🚨 Avisos Vigentes</h3>", unsafe_allow_html=True)
            try:
                avisos = (
                    supabase.table("anuncios_urgentes")
                    .select("*")
                    .eq("is_active", True)
                    .neq("prioridad", 999)
                    .execute()
                    .data or []
                )

                avisos_vivos = []
                ahora_secundario = dt_datetime.now(TZ_CHILE)

                for aviso in avisos:
                    _, _, estado_aviso = obtener_ventana_alerta(
                        aviso,
                        ahora_secundario,
                    )
                    if estado_aviso == "EN_CURSO":
                        avisos_vivos.append(aviso)
                if not avisos_vivos:
                    st.markdown("<p style='color:#64748b; font-style:italic;'>No hay avisos publicados para hoy.</p>", unsafe_allow_html=True)
                else:
                    for a in avisos_vivos[:3]:
                        color_borde = "#ef4444" if str(a['prioridad']) == "1" else "#f59e0b"
                        bg_color = "#fef2f2" if str(a['prioridad']) == "1" else "#fffbeb"
                        st.markdown(f"""
                        <div style='background:{bg_color}; padding:15px; border-radius:12px; border-left:6px solid {color_borde}; margin-bottom:12px; color:#1e293b; border-top:1px solid #e2e8f0; border-right:1px solid #e2e8f0; border-bottom:1px solid #e2e8f0;'>
                            <div style='font-weight:900; color:{color_borde}; font-size:calc(1.1rem * {escala});'>{a['titulo']}</div>
                            <div style='font-size:calc(0.95rem * {escala}); margin-top:5px; font-weight:500;'>{a['descripcion']}</div>
                        </div>
                        """, unsafe_allow_html=True)
            except Exception as e:
                st.text(f"Error Avisos: {e}")

        st.stop()

    # ==============================================================================
    # 💻 2. PANEL DE GESTIÓN (MENSAJERÍA INTERNA)
    # ==============================================================================
    import pandas as pd
    import datetime as dt
    from datetime import datetime as dt_datetime
    
    st.title("📺 Gestión de Pantalla y Mensajería")
    
    col1, col2 = st.columns([1, 1])
    with col1:
        with st.container(border=True):
            st.subheader("🖥️ Control de Proyección")
            if st.button("🚀 INICIAR PANTALLA PÚBLICA", type="primary", use_container_width=True):
                st.session_state.ver_pantalla_tv = True
                st.rerun()
    with col2:
        with st.container(border=True):
            st.subheader("🔗 Sincronización")
            url_cal = st.text_input("URL Google Calendar (.ics)", value=st.session_state.get('url_calendario_tv', ''), label_visibility="collapsed")
            if st.button("Sincronizar ahora"):
                st.session_state.url_calendario_tv = url_cal
                st.success("Calendario sincronizado")

    st.divider()
    
    tab1, tab2, tab3, tab4 = st.tabs(["🔴 Alerta Roja", "🗓️ Añadir Evento", "🔔 Añadir Aviso", "🗑️ Gestionar y Eliminar"])
    
    with tab1:
        st.warning(
            "Las alertas rojas interrumpen la pantalla completa. "
            "Puedes activarlas de inmediato o dejarlas programadas."
        )

        sub_alerta_ahora, sub_alerta_programar = st.tabs(
            ["🚨 Activar ahora", "🗓️ Programar alerta"]
        )

        with sub_alerta_ahora:
            with st.form("form_alerta_roja_inmediata"):
                msg_rojo_ahora = st.text_area(
                    "Mensaje urgente",
                    placeholder="Ej. Evacuar hacia la zona de seguridad.",
                    key="msg_rojo_ahora",
                )
                minutos = st.number_input(
                    "Duración (minutos)",
                    min_value=1,
                    max_value=240,
                    value=5,
                    step=1,
                )

                if st.form_submit_button(
                    "🚨 ACTIVAR ALERTA AHORA",
                    type="primary",
                    use_container_width=True,
                ):
                    if not msg_rojo_ahora.strip():
                        st.warning("Escribe el mensaje de la alerta.")
                    else:
                        try:
                            inicio = dt_datetime.now(TZ_CHILE)
                            fin = inicio + dt.timedelta(minutes=int(minutos))

                            datos_alerta = {
                                "titulo": "ALERTA",
                                "descripcion": msg_rojo_ahora.strip(),
                                "prioridad": 999,
                                "inicio_programado": inicio.isoformat(),
                                "fin_programado": fin.isoformat(),
                                # Se conserva para compatibilidad con el sistema anterior.
                                "expiracion": fin.isoformat(),
                                "is_active": True,
                            }

                            resultado = (
                                supabase.table("anuncios_urgentes")
                                .insert(datos_alerta)
                                .execute()
                            )

                            registro_id = (
                                resultado.data[0].get("id")
                                if resultado.data
                                else None
                            )
                            registrar_auditoria(
                                "activar alerta roja inmediata",
                                "Modo TV",
                                registro_id=registro_id,
                                detalle={
                                    "inicio": inicio.isoformat(),
                                    "fin": fin.isoformat(),
                                },
                            )

                            st.success(
                                f"Alerta activada hasta las {fin.strftime('%H:%M')}."
                            )
                            time.sleep(0.15)
                            st.rerun()

                        except Exception as e:
                            registrar_error("activar_alerta_roja", e)
                            st.error(f"No fue posible activar la alerta: {e}")

        with sub_alerta_programar:
            with st.form("form_alerta_roja_programada"):
                msg_rojo_programado = st.text_area(
                    "Mensaje de la alerta programada",
                    placeholder="Ej. Simulacro de evacuación.",
                    key="msg_rojo_programado",
                )

                hoy_chile = dt_datetime.now(TZ_CHILE).date()
                ahora_redondeada = (
                    dt_datetime.now(TZ_CHILE)
                    + dt.timedelta(minutes=5)
                ).replace(second=0, microsecond=0)
                fin_sugerido = ahora_redondeada + dt.timedelta(minutes=15)

                col_ai, col_ahi, col_af, col_ahf = st.columns(4)
                fecha_inicio_alerta = col_ai.date_input(
                    "Fecha de inicio",
                    value=hoy_chile,
                    format="DD/MM/YYYY",
                    key="fecha_inicio_alerta",
                )
                hora_inicio_alerta = col_ahi.time_input(
                    "Hora de inicio",
                    value=ahora_redondeada.time(),
                    key="hora_inicio_alerta",
                )
                fecha_fin_alerta = col_af.date_input(
                    "Fecha de término",
                    value=hoy_chile,
                    format="DD/MM/YYYY",
                    key="fecha_fin_alerta",
                )
                hora_fin_alerta = col_ahf.time_input(
                    "Hora de término",
                    value=fin_sugerido.time(),
                    key="hora_fin_alerta",
                )

                if st.form_submit_button(
                    "🗓️ GUARDAR PROGRAMACIÓN",
                    type="primary",
                    use_container_width=True,
                ):
                    inicio = combinar_fecha_hora_chile(
                        fecha_inicio_alerta,
                        hora_inicio_alerta,
                    )
                    fin = combinar_fecha_hora_chile(
                        fecha_fin_alerta,
                        hora_fin_alerta,
                    )

                    if not msg_rojo_programado.strip():
                        st.warning("Escribe el mensaje de la alerta.")
                    elif fin <= inicio:
                        st.warning(
                            "La fecha y hora de término deben ser posteriores "
                            "a la fecha y hora de inicio."
                        )
                    else:
                        try:
                            datos_alerta = {
                                "titulo": "ALERTA PROGRAMADA",
                                "descripcion": msg_rojo_programado.strip(),
                                "prioridad": 999,
                                "inicio_programado": inicio.isoformat(),
                                "fin_programado": fin.isoformat(),
                                "expiracion": fin.isoformat(),
                                "is_active": True,
                            }

                            resultado = (
                                supabase.table("anuncios_urgentes")
                                .insert(datos_alerta)
                                .execute()
                            )
                            registro_id = (
                                resultado.data[0].get("id")
                                if resultado.data
                                else None
                            )

                            registrar_auditoria(
                                "programar alerta roja",
                                "Modo TV",
                                registro_id=registro_id,
                                detalle={
                                    "inicio": inicio.isoformat(),
                                    "fin": fin.isoformat(),
                                },
                            )

                            st.success(
                                "Alerta programada desde "
                                f"{inicio.strftime('%d/%m/%Y %H:%M')} hasta "
                                f"{fin.strftime('%d/%m/%Y %H:%M')}."
                            )

                        except Exception as e:
                            registrar_error("programar_alerta_roja", e)
                            st.error(f"No fue posible programar la alerta: {e}")

        st.markdown("#### 🛑 Control de alertas en curso")
        if st.button(
            "Apagar únicamente las alertas que están activas ahora",
            use_container_width=True,
        ):
            try:
                ahora = dt_datetime.now(TZ_CHILE)
                alertas = (
                    supabase.table("anuncios_urgentes")
                    .select("*")
                    .eq("prioridad", 999)
                    .eq("is_active", True)
                    .execute()
                    .data or []
                )
                ids_apagados = []
                for alerta_actual in alertas:
                    _, _, estado = obtener_ventana_alerta(alerta_actual, ahora)
                    if estado == "EN_CURSO":
                        supabase.table("anuncios_urgentes").update(
                            {"is_active": False}
                        ).eq("id", alerta_actual["id"]).execute()
                        ids_apagados.append(alerta_actual["id"])

                registrar_auditoria(
                    "apagar alertas rojas en curso",
                    "Modo TV",
                    detalle={"ids": ids_apagados},
                )

                if ids_apagados:
                    st.success("Las alertas en curso fueron apagadas.")
                else:
                    st.info("No había alertas rojas activas en este momento.")
                st.rerun()

            except Exception as e:
                registrar_error("apagar_alertas_rojas", e)
                st.error(f"No fue posible apagar las alertas: {e}")

    with tab2:
        st.info(
            "El evento aparecerá cada día comprendido entre la fecha inicial "
            "y la fecha final, dentro del horario indicado."
        )

        with st.form("nuevo_evento_tv"):
            titulo_evento = st.text_input(
                "Título del evento",
                placeholder="Ej. Semana de la Educación Técnico Profesional",
            )
            descripcion_evento = st.text_area(
                "Descripción",
                placeholder="Información que se mostrará en la pantalla.",
            )

            hoy_chile = dt_datetime.now(TZ_CHILE).date()
            col_fi, col_ff, col_hi, col_hf = st.columns(4)
            fecha_inicio_evento = col_fi.date_input(
                "Fecha de inicio",
                value=hoy_chile,
                format="DD/MM/YYYY",
                key="fecha_inicio_evento_tv",
            )
            fecha_fin_evento = col_ff.date_input(
                "Fecha de término",
                value=hoy_chile,
                format="DD/MM/YYYY",
                key="fecha_fin_evento_tv",
            )
            hora_inicio_evento = col_hi.time_input(
                "Hora de inicio diaria",
                value=dt.time(8, 0),
                key="hora_inicio_evento_tv",
            )
            hora_fin_evento = col_hf.time_input(
                "Hora de término diaria",
                value=dt.time(17, 0),
                key="hora_fin_evento_tv",
            )

            if st.form_submit_button(
                "💾 Guardar evento",
                type="primary",
                use_container_width=True,
            ):
                if not titulo_evento.strip():
                    st.warning("Ingresa un título para el evento.")
                elif fecha_fin_evento < fecha_inicio_evento:
                    st.warning(
                        "La fecha de término no puede ser anterior a la fecha de inicio."
                    )
                elif hora_fin_evento <= hora_inicio_evento:
                    st.warning(
                        "La hora de término debe ser posterior a la hora de inicio."
                    )
                else:
                    try:
                        datos_evento = {
                            "titulo": titulo_evento.strip(),
                            "descripcion": descripcion_evento.strip(),
                            # Se conserva fecha_evento para compatibilidad.
                            "fecha_evento": fecha_inicio_evento.isoformat(),
                            "fecha_inicio": fecha_inicio_evento.isoformat(),
                            "fecha_fin": fecha_fin_evento.isoformat(),
                            "hora_inicio": hora_inicio_evento.strftime("%H:%M"),
                            "hora_fin": hora_fin_evento.strftime("%H:%M"),
                            "is_active": True,
                        }

                        resultado = (
                            supabase.table("eventos_tv")
                            .insert(datos_evento)
                            .execute()
                        )
                        registro_id = (
                            resultado.data[0].get("id")
                            if resultado.data
                            else None
                        )

                        registrar_auditoria(
                            "crear evento con rango",
                            "eventos_tv",
                            registro_id=registro_id,
                            detalle=datos_evento,
                        )

                        st.success(
                            "Evento guardado desde "
                            f"{fecha_inicio_evento.strftime('%d/%m/%Y')} hasta "
                            f"{fecha_fin_evento.strftime('%d/%m/%Y')}, "
                            f"de {hora_inicio_evento.strftime('%H:%M')} a "
                            f"{hora_fin_evento.strftime('%H:%M')}."
                        )

                    except Exception as e:
                        registrar_error("crear_evento_tv", e)
                        st.error(f"No fue posible guardar el evento: {e}")

    with tab3:
        st.subheader("🔔 Programar aviso lateral")
        st.caption(
            "El aviso aparecerá en la zona principal de la pantalla solamente "
            "durante el rango de fecha y hora seleccionado."
        )

        ahora_aviso = dt_datetime.now(TZ_CHILE).replace(second=0, microsecond=0)
        fin_aviso_default = ahora_aviso + dt.timedelta(hours=1)

        with st.form("nuevo_anuncio_programado", clear_on_submit=False):
            titulo_aviso = st.text_input(
                "Título del aviso *",
                placeholder="Ej. Reemplazos de profesores",
            )

            descripcion_aviso = st.text_area(
                "Descripción *",
                height=180,
                placeholder=(
                    "Puedes escribir cada información en una línea distinta.\n"
                    "08:00 - 09:30 Curso / Profesor\n"
                    "09:45 - 11:15 Curso / Profesor"
                ),
                help="Los saltos de línea se conservarán en la pantalla TV.",
            )

            prioridad_aviso = st.selectbox(
                "Prioridad",
                [1, 2],
                format_func=lambda valor: (
                    "🔴 Alta" if valor == 1 else "🟡 Media"
                ),
            )

            st.markdown("#### 📆 Vigencia del aviso")
            col_inicio_aviso, col_fin_aviso = st.columns(2)

            with col_inicio_aviso:
                fecha_inicio_aviso = st.date_input(
                    "Fecha de inicio *",
                    value=ahora_aviso.date(),
                    format="DD/MM/YYYY",
                    key="fecha_inicio_aviso_tv",
                )
                hora_inicio_aviso = st.time_input(
                    "Hora de inicio *",
                    value=ahora_aviso.time(),
                    step=dt.timedelta(minutes=5),
                    key="hora_inicio_aviso_tv",
                )

            with col_fin_aviso:
                fecha_fin_aviso = st.date_input(
                    "Fecha de término *",
                    value=fin_aviso_default.date(),
                    format="DD/MM/YYYY",
                    key="fecha_fin_aviso_tv",
                )
                hora_fin_aviso = st.time_input(
                    "Hora de término *",
                    value=fin_aviso_default.time(),
                    step=dt.timedelta(minutes=5),
                    key="hora_fin_aviso_tv",
                )

            publicar_aviso = st.form_submit_button(
                "📢 Guardar y programar aviso",
                type="primary",
                use_container_width=True,
            )

            if publicar_aviso:
                if not titulo_aviso.strip():
                    st.warning("Ingresa el título del aviso.")
                elif not descripcion_aviso.strip():
                    st.warning("Ingresa la descripción del aviso.")
                else:
                    inicio_aviso = combinar_fecha_hora_chile(
                        fecha_inicio_aviso,
                        hora_inicio_aviso,
                    )
                    fin_aviso = combinar_fecha_hora_chile(
                        fecha_fin_aviso,
                        hora_fin_aviso,
                    )

                    if fin_aviso <= inicio_aviso:
                        st.warning(
                            "La fecha y hora de término deben ser posteriores "
                            "a la fecha y hora de inicio."
                        )
                    else:
                        datos_aviso = {
                            "titulo": titulo_aviso.strip(),
                            "descripcion": descripcion_aviso.strip(),
                            "prioridad": int(prioridad_aviso),
                            "inicio_programado": inicio_aviso.isoformat(),
                            "fin_programado": fin_aviso.isoformat(),
                            # Se mantiene para compatibilidad con registros antiguos.
                            "expiracion": fin_aviso.isoformat(),
                            "is_active": True,
                        }

                        try:
                            resultado = (
                                supabase.table("anuncios_urgentes")
                                .insert(datos_aviso)
                                .execute()
                            )

                            registro_id = (
                                resultado.data[0].get("id")
                                if resultado.data
                                else None
                            )

                            registrar_auditoria(
                                "programar aviso",
                                "Modo TV",
                                registro_id=registro_id,
                                detalle={
                                    "titulo": titulo_aviso.strip(),
                                    "prioridad": int(prioridad_aviso),
                                    "inicio": inicio_aviso.isoformat(),
                                    "fin": fin_aviso.isoformat(),
                                },
                            )

                            estado_inicial = (
                                "publicado y visible"
                                if inicio_aviso <= dt_datetime.now(TZ_CHILE)
                                else "programado"
                            )

                            st.success(
                                f"✅ Aviso {estado_inicial} desde "
                                f"{inicio_aviso.strftime('%d/%m/%Y %H:%M')} "
                                f"hasta {fin_aviso.strftime('%d/%m/%Y %H:%M')}."
                            )
                            st.balloons()

                        except Exception as e:
                            registrar_error("crear_aviso_programado", e)

                            detalle_error = str(e)
                            error_columnas = any(
                                pista in detalle_error.lower()
                                for pista in [
                                    "inicio_programado",
                                    "fin_programado",
                                    "pgrst204",
                                    "schema cache",
                                ]
                            )

                            if error_columnas:
                                st.error(
                                    "Supabase todavía no reconoce las columnas "
                                    "`inicio_programado` y `fin_programado`. "
                                    "Ejecuta el archivo "
                                    "`migracion_avisos_programados.sql` en "
                                    "Supabase → SQL Editor y reinicia la app."
                                )
                            else:
                                st.error(
                                    "No fue posible guardar el aviso en Supabase. "
                                    "Revisa Manage app → Logs. Detalle técnico: "
                                    f"{detalle_error}"
                                )

    with tab4:
        st.subheader("🗑️ Gestionar contenido programado")
        col_del1, col_del2 = st.columns(2)

        with col_del1:
            st.markdown("#### Eventos")
            try:
                evs = (
                    supabase.table("eventos_tv")
                    .select(
                        "id, titulo, fecha_evento, fecha_inicio, fecha_fin, "
                        "hora_inicio, hora_fin"
                    )
                    .eq("is_active", True)
                    .order("fecha_inicio")
                    .execute()
                    .data or []
                )
            except Exception:
                evs = (
                    supabase.table("eventos_tv")
                    .select("id, titulo, fecha_evento, hora_inicio, hora_fin")
                    .eq("is_active", True)
                    .order("fecha_evento")
                    .execute()
                    .data or []
                )

            if evs:
                ev_dict = {}
                for evento in evs:
                    fecha_i = evento.get("fecha_inicio") or evento.get("fecha_evento")
                    fecha_f = evento.get("fecha_fin") or evento.get("fecha_evento")
                    etiqueta = (
                        f"{evento.get('titulo', 'Evento')} | "
                        f"{fecha_i} al {fecha_f} | "
                        f"{str(evento.get('hora_inicio', ''))[:5]}-"
                        f"{str(evento.get('hora_fin', ''))[:5]}"
                    )
                    ev_dict[etiqueta] = evento["id"]

                sel_ev = st.selectbox(
                    "Selecciona un evento:",
                    ["-- Seleccionar --"] + list(ev_dict.keys()),
                    key="eliminar_evento_programado",
                )
                if (
                    st.button(
                        "🗑️ Desactivar evento",
                        use_container_width=True,
                        key="btn_desactivar_evento",
                    )
                    and sel_ev != "-- Seleccionar --"
                ):
                    evento_id = ev_dict[sel_ev]
                    supabase.table("eventos_tv").update(
                        {"is_active": False}
                    ).eq("id", evento_id).execute()
                    registrar_auditoria(
                        "desactivar evento",
                        "eventos_tv",
                        registro_id=evento_id,
                    )
                    st.success("Evento desactivado.")
                    st.rerun()
            else:
                st.info("No hay eventos activos o programados.")

        with col_del2:
            st.markdown("#### Avisos y alertas")
            anns = (
                supabase.table("anuncios_urgentes")
                .select(
                    "id, titulo, descripcion, prioridad, expiracion, "
                    "inicio_programado, fin_programado"
                )
                .eq("is_active", True)
                .order("id", desc=True)
                .execute()
                .data or []
            )

            if anns:
                ahora = dt_datetime.now(TZ_CHILE)
                an_dict = {}

                for aviso in anns:
                    if str(aviso.get("prioridad")) == "999":
                        inicio, fin, estado = obtener_ventana_alerta(aviso, ahora)
                        icono_estado = {
                            "EN_CURSO": "🚨 EN CURSO",
                            "PROGRAMADA": "🗓️ PROGRAMADA",
                            "FINALIZADA": "⚫ FINALIZADA",
                            "INVALIDA": "⚠️ SIN FECHA",
                        }.get(estado, estado)
                        etiqueta = (
                            f"{icono_estado} | "
                            f"{inicio.strftime('%d/%m %H:%M') if inicio else 'Sin inicio'} "
                            f"→ {fin.strftime('%d/%m %H:%M') if fin else 'Sin fin'} | "
                            f"{str(aviso.get('descripcion', ''))[:45]}"
                        )
                    else:
                        inicio, fin, estado = obtener_ventana_alerta(aviso, ahora)
                        icono = (
                            "🔴"
                            if str(aviso.get("prioridad")) == "1"
                            else "🟡"
                        )
                        estado_legible = {
                            "EN_CURSO": "EN CURSO",
                            "PROGRAMADA": "PROGRAMADO",
                            "FINALIZADA": "FINALIZADO",
                            "INVALIDA": "SIN FECHA",
                        }.get(estado, estado)

                        etiqueta = (
                            f"{icono} {estado_legible} | "
                            f"{inicio.strftime('%d/%m %H:%M') if inicio else 'Inmediato'} "
                            f"→ {fin.strftime('%d/%m %H:%M') if fin else 'Sin término'} | "
                            f"{aviso.get('titulo', 'Aviso')}"
                        )

                    an_dict[etiqueta] = aviso["id"]

                sel_an = st.selectbox(
                    "Selecciona un aviso o alerta:",
                    ["-- Seleccionar --"] + list(an_dict.keys()),
                    key="eliminar_aviso_programado",
                )
                if (
                    st.button(
                        "🗑️ Desactivar aviso / alerta",
                        use_container_width=True,
                        key="btn_desactivar_aviso",
                    )
                    and sel_an != "-- Seleccionar --"
                ):
                    aviso_id = an_dict[sel_an]
                    supabase.table("anuncios_urgentes").update(
                        {"is_active": False}
                    ).eq("id", aviso_id).execute()
                    registrar_auditoria(
                        "desactivar aviso o alerta",
                        "Modo TV",
                        registro_id=aviso_id,
                    )
                    st.success("Contenido desactivado.")
                    st.rerun()
            else:
                st.info("No hay avisos ni alertas activas o programadas.")

        st.divider()
        st.subheader("📋 Calendario de publicaciones")

        try:
            eventos_resumen = (
                supabase.table("eventos_tv")
                .select(
                    "titulo, descripcion, fecha_inicio, fecha_fin, "
                    "hora_inicio, hora_fin, is_active"
                )
                .eq("is_active", True)
                .order("fecha_inicio")
                .execute()
                .data or []
            )
        except Exception:
            eventos_resumen = (
                supabase.table("eventos_tv")
                .select(
                    "titulo, descripcion, fecha_evento, "
                    "hora_inicio, hora_fin, is_active"
                )
                .eq("is_active", True)
                .order("fecha_evento")
                .execute()
                .data or []
            )

        if eventos_resumen:
            df_ev = pd.DataFrame(eventos_resumen)
            df_ev.rename(columns={
                "titulo": "Evento",
                "descripcion": "Descripción",
                "fecha_inicio": "Desde",
                "fecha_fin": "Hasta",
                "fecha_evento": "Fecha",
                "hora_inicio": "Hora inicio",
                "hora_fin": "Hora fin",
                "is_active": "Activo",
            }, inplace=True)
            st.write("**Eventos del cronograma:**")
            st.dataframe(df_ev, use_container_width=True, hide_index=True)
        else:
            st.info("No hay eventos activos.")

        alertas_resumen = (
            supabase.table("anuncios_urgentes")
            .select(
                "titulo, descripcion, prioridad, inicio_programado, "
                "fin_programado, expiracion, is_active"
            )
            .eq("is_active", True)
            .order("id", desc=True)
            .execute()
            .data or []
        )

        if alertas_resumen:
            filas_alertas = []
            ahora = dt_datetime.now(TZ_CHILE)
            for aviso in alertas_resumen:
                inicio, fin, estado = obtener_ventana_alerta(aviso, ahora)
                nivel = (
                    "🚨 Crítica"
                    if str(aviso.get("prioridad")) == "999"
                    else (
                        "🔴 Alta"
                        if str(aviso.get("prioridad")) == "1"
                        else "🟡 Media"
                    )
                )
                filas_alertas.append({
                    "Nivel": nivel,
                    "Título": aviso.get("titulo"),
                    "Mensaje": aviso.get("descripcion"),
                    "Inicio": inicio.strftime("%d/%m/%Y %H:%M") if inicio else "Inmediato",
                    "Fin": fin.strftime("%d/%m/%Y %H:%M") if fin else "Sin término",
                    "Estado": estado,
                })

            st.write("**Avisos y alertas:**")
            st.dataframe(
                pd.DataFrame(filas_alertas),
                use_container_width=True,
                hide_index=True,
            )
        else:
            st.info("No hay avisos o alertas activas/programadas.")

